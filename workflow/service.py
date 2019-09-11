# -*- coding: utf-8 -*-
from xml.etree import ElementTree

from workflow.models import FlowNode, FlowTrans
from docx import Document
from django.conf import settings
import os
import logging

logger = logging.getLogger(__name__)


def bpmn_parse(xml_text):
    """
    解析bpmn格式文件
    :param xml_text:
    :return:
    """
    ns = {
        'xsi': 'http://www.w3.org/2001/XMLSchema-instance',
        'bpmn': 'http://www.omg.org/spec/BPMN/20100524/MODEL',
        'bpmndi': 'http://www.omg.org/spec/BPMN/20100524/DI',
        'dc': 'http://www.omg.org/spec/DD/20100524/DC',
        'di': 'http://www.omg.org/spec/DD/20100524/DI'
    }
    nodes = []
    trans = []
    root = ElementTree.fromstring(xml_text)
    process = list(root)[0]
    # 环节
    tasks = process.findall('bpmn:task', ns)
    # 流转
    sequences = process.findall('bpmn:sequenceFlow', ns)

    for task in tasks:
        has_multi_instances = task.findall('bpmn:multiInstanceLoopCharacteristics', ns)
        node = task.attrib
        if len(has_multi_instances):
            node['is_parallel'] = True
        else:
            node['is_parallel'] = False
        nodes.append(node)

    for sequence in sequences:
        trans.append(sequence.attrib)

    return nodes, trans


def flow_nodes(flow_id):
    """
    获取流程环节信息
    :param flow_id:
    :return:
    """
    qs = FlowNode.objects.filter(flow_id=flow_id, del_flag=0).order_by('name')
    node_list = []

    for node in qs:
        is_start_node = FlowTrans.objects.filter(flow_id=flow_id, incoming__startswith='StartEvent',
                                                 outgoing=node.task_id).exists()
        if node.process:
            process = {
                'id': node.process.id, 'name': node.process.name, 'type': node.process.type,
                'image': node.process.image.url if node.process.image else None
            }
        else:
            process = None
        node_list.append({
            'id': node.id, 'name': node.name, 'look_on': node.look_on, 'step': node.step, 'task_id': node.task_id,
            'condition': node.condition, 'process': process, 'is_start_node': is_start_node, 'is_parallel_start_node': node.parallel_node_start
        })

    return node_list


def get_start_node(flow_id):
    """
    获取流程第一个环节
    :param flow_id:
    :return:
    """
    transition = FlowTrans.objects.filter(flow_id=flow_id, incoming__startswith='StartEvent').first()

    if transition:
        node = FlowNode.objects.filter(flow_id=flow_id, task_id=transition.outgoing).first()

        if node:
            return node.id
    return None

def get_end_node(flow_id):
    """
    获取流程第一个环节
    :param flow_id:
    :return:
    """
    transition = FlowTrans.objects.filter(flow_id=flow_id, outgoing__startswith='EndEvent').first()

    if transition:
        node = FlowNode.objects.filter(flow_id=flow_id, task_id=transition.incoming).first()

        if node:
            return node.id
    return None

def get_out_nodes(node):
    out_trans = FlowTrans.objects.filter(incoming=node.task_id)
    nodes = []
    for out_tran in out_trans:
        out_node = FlowNode.objects.filter(task_id=out_tran.outgoing).first()
        if out_node:
            nodes.append(out_node)
    return nodes

def get_incoming_nodes(node):
    incoming_trans = FlowTrans.objects.filter(outgoing=node.task_id)
    nodes = []
    for incoming_tran in incoming_trans:
        incoming_node = FlowNode.objects.filter(task_id=incoming_tran.incoming).first()
        if incoming_node:
            nodes.append(incoming_node)
    return nodes

stacks = []
merging_nodes = []

def get_end_parallel_node(node_id):
    global stacks
    global merging_nodes
    start_node = FlowNode.objects.filter(pk=node_id).first()
    out_nodes = get_out_nodes(start_node)
    stacks = [{
        'parent': start_node,
        'nodes': out_nodes
    }]
    merging_nodes = []
    get_end_parallel_node_stack()
    print stacks
    print merging_nodes
    return merging_nodes


def get_end_parallel_node_stack():
    global stacks
    global merging_nodes
    len_stacks = len(stacks)
    if len_stacks == 0:
        return True
    print stacks
    origin_stacks = stacks[:] # 초기 stacks
    origin_last_stack = origin_stacks[len_stacks-1] # 마지막 초기 stack
    origin_last_stack_nodes = origin_last_stack['nodes'][:] # 초기 마지막 stack의 nodes

    last_stack_nodes = stacks[len_stacks-1]['nodes'] # 변경할 마지막 stack의 nodes

    for origin_node in origin_last_stack_nodes:
        out_nodes = get_out_nodes(origin_node)
        if len(out_nodes) > 1:
            stacks.append({
                'parent': origin_node,
                'nodes': out_nodes
            })
        elif len(out_nodes) == 1:
            out_node = out_nodes[0]
            out_node_incoming_nodes = get_incoming_nodes(out_node)
            all_is_in = True
            for out_node_incoming_node in out_node_incoming_nodes:
                is_in = False
                for stack in stacks:
                    if out_node_incoming_node in stack['nodes']:
                        is_in = True
                        break
                if not is_in:
                    all_is_in = False
                    break
            if all_is_in:
                first_changed_stack = None
                for stack in stacks:
                    for out_node_incoming_node in out_node_incoming_nodes:
                        if out_node_incoming_node in stack['nodes']:
                            stack['nodes'].remove(out_node_incoming_node)
                            if not first_changed_stack:
                                first_changed_stack = stack
                dupStacks = stacks[:]

                if len(out_node_incoming_nodes) > 1 and first_changed_stack and first_changed_stack['parent'].parallel_node_start:
                    merging_nodes.append({
                        'start': first_changed_stack['parent'],
                        'end': out_node
                    })
                for dupStack in dupStacks:
                    if len(dupStack['nodes']) == 0:
                        for stack in stacks:
                            if dupStack['parent'] in stack['nodes']:
                                stack['nodes'].remove(dupStack['parent'])
                                stack['nodes'].append(out_node)

                        stacks.remove(dupStack)
                if first_changed_stack is not None and not out_node in first_changed_stack['nodes']:
                    first_changed_stack['nodes'].append(out_node)
                    # for stack in stacks:
                    #     if stack['parent'] == last_changed_stack['parent']:
                    #         stack = last_changed_stack
        else:
            continue

    get_end_parallel_node_stack()


def bpmn_color(xml, passed, current, mode, stop_node):
    """
    解析bpmn格式文件,对环节着色
    :param xml_text:
    :return:
    """
    xml = xml.replace('bpmn:definitions',
                      'bpmn:definitions xmlns:bioc="http://bpmn.io/schema/bpmn/biocolor/1.0"', 1)

    if len(passed) > 0:
        if mode == 0:
            last_task_id = passed[-1]
            for task_id in passed:
                old = 'bpmnElement="{}"'.format(task_id)
                if task_id == last_task_id:
                    new = 'bpmnElement="{}" bioc:stroke="#FE2E2E" bioc:fill="#F8E0E0"'.format(task_id)
                else:
                    new = 'bpmnElement="{}" bioc:stroke="#1E88E5" bioc:fill="#BBDEFB"'.format(task_id)
                xml = xml.replace(old, new, 1)
        else:
            for task_id in passed:
                old = 'bpmnElement="{}"'.format(task_id)
                new = 'bpmnElement="{}" bioc:stroke="#1E88E5" bioc:fill="#BBDEFB"'.format(task_id)
                xml = xml.replace(old, new, 1)
    if len(current) > 0 and mode == 1:
        for task_id in current:
            old = 'bpmnElement="{}"'.format(task_id)
            new = 'bpmnElement="{}" bioc:stroke="#FE2E2E" bioc:fill="#F8E0E0"'.format(task_id)
            xml = xml.replace(old, new, 1)
    if len(stop_node) > 0:
        for stop_task_id in stop_node:
            old = 'bpmnElement="{}"'.format(stop_task_id)
            new = 'bpmnElement="{}" bioc:stroke="#f18a00" bioc:fill="#ffe6af"'.format(stop_task_id)
            xml = xml.replace(old, new, 1)
    # print xml
    return xml


def flow_doc_save(flow_id, name, content):
    """
    保存应用模板生成文件
    :param content: 内容
    :return:
    """
    path = ''
    try:
        # 打开文档
        document = Document()
        # 添加文本
        # document.add_paragraph(content)   改成下面这样就可以保存中文字体了， fuck
        document.add_paragraph('%s' % content)
        # 保存文件
        media = settings.MEDIA_ROOT
        path = u'{}/workflow/{}'.format(media, flow_id)
        is_exists = os.path.exists(path)
        if not is_exists:
            os.makedirs(path)
        media_path = u'{}/{}'.format(path, name)
        logger.info(media_path)
        document.save(media_path)
        path = u'workflow/{}/{}'.format(flow_id, name)
    except Exception as e:
        logger.exception(u'flow_doc_save Exception:{}'.format(str(e)))
    return path

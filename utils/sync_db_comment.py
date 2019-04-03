# coding=utf-8
import MySQLdb
from django.db.models import ManyToManyRel, ManyToOneRel, ForeignKey

from lets2017 import settings
from django.apps import apps
import os, django

from docx import Document
from docx.oxml.ns import qn

from docx.shared import Pt, Inches


# Create your tests here.

def get_db_connection():
    """
    获取数据库连接
    :return:
    """
    conn = MySQLdb.connect(
        host=settings.HOST,
        port=3306,
        user=settings.DB_USER,
        passwd=settings.DB_PWD,
        db=settings.DB_NAME,
        charset="utf8"
    )
    return conn


def get_model_field(app_name):
    """
    处理一个app下所有model关联的表
    :param app_name:
    :return:
    """
    if app_name.__contains__('.'):
        names = app_name.split(".")
        app_name = names[names.__len__() - 1]
    # 获取一个app
    app = apps.get_app_config(app_name)
    # 获取app下的所有model
    models = app.get_models()
    connect = get_db_connection()
    cur = connect.cursor()

    for item in models:
        # print item
        table_name = str(item._meta.db_table)
        table_remark = str(item._meta.verbose_name)
        print "表名称==" + table_name
        print "表注释==" + table_remark
        sql = "alter table " + table_name + " comment '" + table_remark + "';"
        print sql
        # 修改表的备注
        cur.execute(sql)
        cur.close()
        connect.commit()

        # 查询出表的DDL：
        cur = connect.cursor()
        cur.execute("show create table " + table_name + ";")
        datas = cur.fetchall()
        # 表的构造sql
        table_create_sql = datas[0][1]
        left = table_create_sql.find('(') + 1
        right = table_create_sql.rfind(')') - 1
        lines = table_create_sql[left:right].split('\n')
        fileds = item._meta.get_fields()
        lines = lines[0:len(fileds) + 1]
        # print "截取结果=" + str(lines)
        for field in fileds:
            # 多对多类型和多对一类型没有对应model，无法读取verbose_name进行同步注释
            if not (isinstance(field, ManyToManyRel) or isinstance(field, ManyToOneRel)):
                # print field.verbose_name
                # print field.name
                col_name = str(field.name)
                col_remark = str(field.verbose_name)
                if isinstance(field, ForeignKey):
                    col_name = col_name + "_id"
                for line in lines:
                    line = str.lstrip(str(line))
                    if line.endswith(','):
                        line = str.rstrip(line, ',')
                    if "`" + col_name + "`" in line and line.startswith("`" + col_name + "`"):
                        sql = "alter table " + table_name + " modify column " + line + " comment '" + col_remark + "';"
                        print sql
                        cur.execute(sql)

    # 关闭数据库连接
    cur.close()
    connect.commit()
    connect.close()


def get_models():
    """
    读取应用设置中所有app下的model
    :return:
    """
    for item in settings.INSTALLED_APPS:
        get_model_field(item)


def get_data(sql):
    """获取sql查询结果"""
    conn = get_db_connection()
    cur = conn.cursor()
    cur.execute(sql)
    results = cur.fetchall()
    conn.commit()
    cur.close()
    conn.close()
    return results


def get_mysql_table_name():
    sql = """
        select table_name,table_comment from information_schema.tables where table_schema='%s' and table_type='base table';
        """ % settings.DB_NAME

    return get_data(sql)


def get_mysql_table_col(table_name):
    # format_type(a.atttypid, a.atttypmod) as data_type,   A.COLUMN_ID as "index",
    sql = """
        SELECT
        column_name,
        column_type,
        ifnull(CHARACTER_MAXIMUM_LENGTH,'') AS length,
        is_nullable,
        column_comment ,
        (
            CASE column_key
            WHEN 'PRI' THEN
                'Y'
            ELSE
                'N'
            END
        ) pri
    FROM
        information_schema.COLUMNS 
    WHERE
        table_schema = '%s' 
        AND table_name = '%s' 
    ORDER BY
        ordinal_position

    """ % (settings.DB_NAME, table_name)
    return get_data(sql)


def export_mysql_database_word(title):
    """
    导出数据库到word
    :param title: 文档标题
    :param table_schema: 数据库名称
    :return:
    """

    # 打开文档
    document = Document()
    # 加入不同等级的标题
    document.add_heading(title, 0)

    # 查询所有的表名称
    tables = get_mysql_table_name()

    for item in tables:
        # ('station_water_alarm', '水位告警')
        print("")

        document.add_heading(u'表名：%s  说明：%s' % (item[0], item[1]), 3)

        print(item)
        cols = get_mysql_table_col(item[0])
        col_count = 7
        table = document.add_table(rows=len(cols) + 1, cols=col_count, style="Table Grid")

        # 设置表格宽度
        table.columns[0].width = Inches(0.5)
        table.columns[1].width = Inches(1.5)
        table.columns[2].width = Inches(1.5)
        table.columns[3].width = Inches(0.5)
        table.columns[4].width = Inches(0.5)
        table.columns[5].width = Inches(1.5)
        table.columns[6].width = Inches(0.5)

        # 设置表格头
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = u"序号"
        hdr_cells[1].text = u"列名"
        hdr_cells[2].text = u"类型"
        hdr_cells[3].text = u"长度"
        hdr_cells[4].text = u"为空"
        hdr_cells[5].text = u"说明"
        hdr_cells[6].text = u"主键"

        # ('LEGAL_DEVICE_TYPE', 'NVARCHAR2', 64, 'Y', '设备类型', 'N')
        for i, col in enumerate(cols):
            print(col)
            hdr_cells = table.rows[i + 1].cells
            for j in range(col_count):
                if j == 0:
                    hdr_cells[0].text = str(i + 1)
                else:
                    hdr_cells[j].text = col[j - 1]

    # 保存文件
    document.save(u'项目数据库文档.docx')

    print(u"项目数据库文档导出成功！")


if __name__ == '__main__':
    # 加载当前项目的环境
    os.environ.setdefault("DJANGO_SETTINGS_MODULE", "lets2017.settings")
    django.setup()
    get_models()
    export_mysql_database_word(u"项目数据库文档")
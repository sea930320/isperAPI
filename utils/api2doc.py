# coding=utf-8
import json
from docx import Document

from docx.shared import Inches

# 文件名与对应模块
file_dict = {'account': u'用户模块', 'cms': u'消息模块', 'course': u'课程模块', 'experiment': u'实验模块',
             'project': u'项目模块', 'system': u'系统模块', 'team': u'小组模块', 'workflow': u'流程模块'}


def export_json_2_word(title):
    '''
    将 api json 文件转化为word文档形式
    :param title:  文档名称
    :return:
    '''
    # 打开文档
    document = Document()
    # 加入不同等级的标题
    document.add_heading(title, 1)
    # 二级标题索引
    index = 1

    for item in file_dict:
        print (u'写入模块 %s -ing' % file_dict.get(item))
        # 写入二级标题
        document.add_heading(u'%d %s %s' % (index, item, file_dict.get(item)), 2)
        path = u'../templates/api/%s.json' % item
        # 打开文件
        with open(path, 'r') as load_f:
            # 将文件数据加载为json
            file_json = json.load(load_f)
            api_json = file_json['apis']
            # 三级标题索引
            sub_index = 1
            for api in api_json:
                path = api['path']
                operations = api['operations'][0]
                method = operations['method']
                summary = operations['summary']
                print (u'写入接口 %s -ing' % summary)
                document.add_heading(u'%d.%d %s' % (index, sub_index, summary), 3)
                document.add_paragraph(u'请求方式： %s' % method)
                document.add_paragraph(u'请求地址： %s' % path)
                sub_index += 1
                if 'parameters' in operations:
                    parameters = operations['parameters']
                    if len(parameters) > 0:
                        document.add_paragraph(u'参数列表： ')
                        table = document.add_table(rows=len(parameters) + 1, cols=4, style="Table Grid")

                        # 设置表格宽度
                        table.columns[0].width = Inches(1)
                        table.columns[1].width = Inches(2)
                        table.columns[2].width = Inches(1.5)
                        table.columns[3].width = Inches(1)

                        # 设置表格头
                        hdr_cells = table.rows[0].cells
                        hdr_cells[0].text = u"参数名称"
                        hdr_cells[1].text = u"描述"
                        hdr_cells[2].text = u"参数类型"
                        hdr_cells[3].text = u"是否必填"

                        for i, parameter in enumerate(parameters):
                            name = parameter['name']
                            description = parameter['description']
                            type = parameter['type']
                            required = parameter['required']
                            print(parameter)
                            hdr_cells = table.rows[i + 1].cells
                            hdr_cells[0].text = name
                            hdr_cells[1].text = description
                            hdr_cells[2].text = type
                            hdr_cells[3].text = u'是' if required else u'否'

                    document.add_paragraph(u'')

        print item
        index += 1
    # 保存文件
    document.save(u'%s.docx' % title)

    print(u"%s导出成功！" % title)


if __name__ == '__main__':
    export_json_2_word(u'系统接口说明文档')



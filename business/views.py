#!/usr/bin/python
# -*- coding=utf-8 -*-

# from django.shortcuts import
import json
import logging
import random
from platform import node
import re
import xlrd
import xlwt
from django.utils.http import urlquote

from django.utils.timezone import now

from account.models import Tuser, TNotifications, TInnerPermission
from django.core.paginator import Paginator, EmptyPage
from django.db.models import Q, F, Count
from django.http import HttpResponse
from business.models import *
from business.service import *
from project.models import Project, ProjectRole, ProjectRoleAllocation, ProjectDoc, ProjectDocRole, ProjectUseLog
from team.models import Team, TeamMember
from utils import const, code, tools, easemob
from utils.request_auth import auth_check
from workflow.models import FlowNode, FlowAction, FlowRoleActionNew, FlowRolePosition, \
    FlowPosition, RoleImage, Flow, ProcessRoleActionNew, FlowDocs, FlowRole, FlowRoleAllocation, \
    FlowRoleAllocationAction, ProcessRoleAllocationAction, FlowNodeSelectDecide, SelectDecideItem
from workflow.service import get_start_node, bpmn_color
from datetime import datetime
from django.utils import timezone
import random
import string
from utils.public_fun import getProjectIDByGroupManager, \
    getProjectIDByCompanyManager, \
    getProjectIDByCompanyManagerAssistant, \
    getProjectIDByGroupManagerAssistant
from django.forms.models import model_to_dict
from socketio.socketIO_client import SocketIO, LoggingNamespace
from django.forms.models import model_to_dict
import codecs
import pypandoc
from system.models import UploadFile
import html2text
from account.service import get_client_ip
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import Mm
from docx.enum.text import WD_BREAK
import copy

logger = logging.getLogger(__name__)


def randomString(stringLength=10):
    """Generate a random string of fixed length"""
    letters = string.ascii_lowercase
    return ''.join(random.sample(letters, stringLength))


def api_business_create(request):
    resp = auth_check(request, "POST")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    try:
        project_id = request.POST.get("project_id")  # 项目ID
        use_to = request.POST.get("use_to")

        project = Project.objects.get(pk=project_id)
        loginType = request.session['login_type'] if 'login_type' in request.session else None
        # 判断项目是否存在
        if project:
            # 验证项目中是否有未配置的跳转项目
            if not check_jump_project(project):
                resp = code.get_msg(code.EXPERIMENT_JUMP_PROJECT_SETUP_ERROR)
                return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

            roles = ProjectRole.objects.filter(project_id=project_id)
            if roles.exists() is False:
                resp = code.get_msg(code.PROJECT_ROLE_NOT_EXIST)
                return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
            with transaction.atomic():
                if project.created_role_id in [3, 7]:
                    company_id = project.created_by.tcompanymanagers_set.get().tcompany.id if project.created_role_id == 3 else project.created_by.t_company_set_assistants.get().id if project.created_role_id == 7 else None
                target_company_id = use_to if project.created_role_id in [2,
                                                                          6] else company_id if project.created_role_id in [
                    3, 7] and project.use_to_id is None else None
                target_part_id = project.use_to_id if project.created_role_id in [3,
                                                                                  7] and project.use_to_id is not None else None

                business = Business.objects.create(
                    project_id=project_id,
                    name=project.name,
                    cur_project_id=project_id,
                    created_by=request.user,
                    officeItem=project.officeItem,
                    target_company_id=target_company_id,
                    target_part_id=target_part_id,
                )
                business_roles = []
                for item in roles:
                    business_roles.append(BusinessRole(business=business, name=item.name,
                                                       type=item.type, flow_role_id=item.flow_role_id,
                                                       project_role_id=item.id, project_id=project_id,
                                                       category=item.category, capacity=item.capacity,
                                                       job_type=item.job_type))
                BusinessRole.objects.bulk_create(business_roles)

                # 复制流程角色分配设置
                business_allocations = []
                allocations = ProjectRoleAllocation.objects.filter(project_id=project_id)
                for item in allocations:
                    # 将角色分配中的role_id设置为ProjectRole id
                    role = BusinessRole.objects.filter(business=business, project_role_id=item.role_id).first()
                    project = Project.objects.filter(pk=item.project_id).first()
                    projectRole = ProjectRole.objects.filter(pk=item.role_id).first()
                    if project is None or projectRole is None:
                        continue
                    flow_role_alloc = FlowRoleAllocation.objects.filter(flow_id=project.flow_id, node_id=item.node_id,
                                                                        role_id=projectRole.flow_role_id,
                                                                        no=item.no).first()
                    if flow_role_alloc is None:
                        continue
                    if role:
                        business_allocations.append(
                            BusinessRoleAllocation(business=business, node=FlowNode.objects.get(pk=item.node_id),
                                                   project_id=project_id,
                                                   project_role_alloc_id=item.id,
                                                   flow_role_alloc_id=flow_role_alloc.id,
                                                   role=role,
                                                   can_start=item.can_start,
                                                   can_terminate=item.can_terminate,
                                                   can_brought=item.can_brought,
                                                   can_take_in=item.can_take_in,
                                                   no=item.no))
                BusinessRoleAllocation.objects.bulk_create(business_allocations)

                teammates_configuration(business.id, [])
                target_company = TCompany.objects.get(
                    pk=target_company_id) if target_company_id else TParts.objects.get(
                    pk=target_part_id).company if target_part_id else None
                p_u_log = ProjectUseLog(user=request.user, role_id=loginType,
                                        group_id=target_company.group_id if target_company else None,
                                        company=target_company, ip=get_client_ip(request),
                                        request_url=request.path, project=project, business=business)
                p_u_log.save()
                resp = code.get_msg(code.SUCCESS)
                resp['d'] = {
                    'id': business.id, 'name': u'{0} {1}'.format(business.id, business.name),
                    'project_id': business.project_id,
                    'show_nickname': business.show_nickname, 'start_time': business.start_time,
                    'end_time': business.end_time, 'status': business.status,
                    'created_by': user_simple_info(business.created_by.id) if business.created_by else '',
                    'node_id': business.node.id if business.node else '',
                    'create_time': business.create_time.strftime('%Y-%m-%d')
                }
        else:
            resp = code.get_msg(code.PROJECT_NOT_EXIST)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    except Exception as e:
        logger.exception('api_business_create Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


def api_business_remove(request):
    resp = auth_check(request, "GET")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    try:
        business_id = request.GET.get("business_id", None)  # 项目ID
        if business_id is None:
            resp = code.get_msg(code.PARAMETER_ERROR)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
        business = Business.objects.filter(pk=business_id).first()
        if business is None:
            resp = code.get_msg(code.BUSINESS_NOT_EXIST)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
        if business.status != 1:
            resp = code.get_msg(code.BUSINESS_HAS_STARTED)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
        business.del_flag = True
        business.save()
        resp = code.get_msg(code.SUCCESS)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
    except Exception as e:
        logger.exception('api_business_create Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


def teammates_configuration(business_id, seted_users_fromInnerPermission):
    # check team counts

    business = Business.objects.get(id=business_id)
    business_team_counts = list(BusinessRole.objects.filter(
        Q(business_id=business_id, project_id=business.cur_project_id) & ~Q(job_type_id=None)).values('job_type__name',
                                                                                                      'capacity'))

    project = Project.objects.get(pk=business.cur_project_id)
    first_node_id = get_start_node(project.flow_id)
    node = FlowNode.objects.get(pk=first_node_id)
    startRoleAlloc = BusinessRoleAllocation.objects.filter(business=business, node=node, can_start=1,
                                                           can_take_in=1).first()
    for item in business_team_counts:
        if item['job_type__name'] == startRoleAlloc.role.name:
            item['capacity'] -= 1

    company_id = None
    target_user_counts = []
    if business.target_part is not None:
        target_user_counts = list(Tuser.objects.filter(
            Q(tposition__parts_id=business.target_part.id, is_review=1) & ~Q(
                id=business.created_by_id if business.jumper_id is None else business.jumper_id)).values(
            'tposition__name').annotate(counts=Count('id')))
        company_id = business.target_part.company_id
    elif business.target_company is not None:
        target_user_counts = list(Tuser.objects.filter(
            Q(tcompany=business.target_company, is_review=1) & ~Q(
                id=business.created_by_id if business.jumper_id is None else business.jumper_id)).values(
            'tposition__name').annotate(counts=Count('id')))
        company_id = business.target_company.id

    if seted_users_fromInnerPermission:
        for xs in seted_users_fromInnerPermission:
            xIndex = next(
                (index for (index, xt) in enumerate(target_user_counts) if xt['tposition__name'] == xs['role_name']),
                None)
            if xIndex is None:
                target_user_counts.append({'tposition__name': xs['role_name'], 'counts': 1})
            else:
                target_user_counts[xIndex]['counts'] += 1

    moreResult = []
    for item in business_team_counts:
        matched_item = next((x for x in target_user_counts if x['tposition__name'] == item['job_type__name']), None)
        if matched_item is None:
            moreResult.append({'role_name': item['job_type__name'], 'moreCount': item['capacity']})
        elif matched_item['counts'] < item['capacity']:
            moreResult.append(
                {'role_name': item['job_type__name'], 'moreCount': item['capacity'] - matched_item['counts']})

    if moreResult:
        if len(seted_users_fromInnerPermission) != 0:
            return
        if TInnerPermission.objects.get(id=1).ownPositions.filter(parts__company_id=company_id).count() != 0 \
                and TInnerPermission.objects.get(id=1).ownPositions.filter(parts__company_id=company_id).filter(
                    tuser__isnull=False).values_list('tuser', flat=True).count() != 0:
            newNotification = TNotifications.objects.create(
                type='businessMoreTeammate_' + str(business_id),
                content=str(moreResult),
                link='business-moreTeammates',
                role_id=5,
                mode=0
            )
            newNotification.save()
            for userID in TInnerPermission.objects.get(id=1).ownPositions.filter(parts__company_id=company_id).filter(
                    tuser__isnull=False).values_list('tuser', flat=True):
                newNotification.targets.add(Tuser.objects.get(id=userID))
        elif TInnerPermission.objects.get(id=1).ownPositions.filter(parts__company_id=company_id).count() == 0 \
                or TInnerPermission.objects.get(id=1).ownPositions.filter(parts__company_id=company_id).filter(
                    tuser__isnull=False).values_list('tuser', flat=True).count() == 0:
            newNotification = TNotifications.objects.create(
                type='businessMoreTeammate_' + str(business_id),
                content=str(moreResult),
                link='manager-moreTeammates',
                role_id=3,
                mode=0
            )
            newNotification.save()
            for userItem in TCompany.objects.get(id=company_id).tcompanymanagers_set.all():
                newNotification.targets.add(userItem.tuser)
        return

    # configuration team

    targetUnitUsers = []
    teammateList = list(
        BusinessRoleAllocation.objects.filter(business_id=business_id, project_id=business.cur_project_id).values(
            'role_id', 'role__job_type__name',
            'no').distinct())
    teammateList.pop(next((index for (index, x) in enumerate(teammateList) if
                           x['role_id'] == startRoleAlloc.role_id and x['no'] == startRoleAlloc.no), None))
    if business.target_part is not None:
        targetUnitUsers = [{
            'id': item.id,
            'position': item.tposition.name,
        } for item in Tuser.objects.filter(Q(tposition__parts_id=business.target_part.id, is_review=1) & ~Q(
            id=business.created_by_id if business.jumper_id is None else business.jumper_id))]
    elif business.target_company is not None:
        targetUnitUsers = [{
            'id': item.id,
            'position': item.tposition.name if item.tposition else None,
        } for item in Tuser.objects.filter(Q(tcompany_id=company_id, is_review=1) & ~Q(
            id=business.created_by_id if business.jumper_id is None else business.jumper_id))]

    newTeammate = BusinessTeamMember.objects.create(
        business_id=business_id,
        user_id=business.created_by_id if business.jumper_id is None else business.jumper_id,
        business_role_id=startRoleAlloc.role_id,
        no=startRoleAlloc.no,
        del_flag=0,
        project_id=business.cur_project_id
    )
    newTeammate.save()

    if seted_users_fromInnerPermission:
        for item in seted_users_fromInnerPermission:
            targetUnitUsers.append({'position': item['role_name'], 'id': item['userId']})

    for teamItem in teammateList:
        if teamItem['role__job_type__name'] is None:
            newTeammate = BusinessTeamMember.objects.create(
                business_id=business_id,
                user_id=None,
                business_role_id=teamItem['role_id'],
                no=teamItem['no'],
                del_flag=0,
                project_id=business.cur_project_id
            )
            newTeammate.save()
        else:
            print(targetUnitUsers)
            print("------")
            print(teamItem)
            selectedUser = random.choice(
                [a for a in targetUnitUsers if a['position'] == teamItem['role__job_type__name']])
            targetUnitUsers.pop(next((index for (index, x) in enumerate(targetUnitUsers) if x == selectedUser), None))
            newTeammate = BusinessTeamMember.objects.create(
                business_id=business_id,
                user_id=selectedUser['id'],
                business_role_id=teamItem['role_id'],
                no=teamItem['no'],
                del_flag=0,
                project_id=business.cur_project_id
            )
            newTeammate.save()

    Business.objects.filter(id=business_id).update(status=2)

    return 'team_configured'


def api_business_list(request):
    resp = auth_check(request, "GET")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    try:
        search = request.GET.get("search", None)  # 关键字
        page = int(request.GET.get("page", 1))  # 页码
        size = int(request.GET.get("size", const.ROW_SIZE))  # 页面条数
        status = int(request.GET.get("status", 1))  # 实验状态

        user = request.user
        bussinessIDsInTeam = BusinessTeamMember.objects.filter(user=user, del_flag=0).values_list('business_id',
                                                                                                  flat=True).distinct()
        qs = Business.objects.filter(
            Q(del_flag=0, pk__in=bussinessIDsInTeam) | Q(del_flag=0, created_by=request.user))

        if search:
            qs = qs.filter(Q(name__icontains=search) | Q(pk__icontains=search))
            qs = qs.filter(del_flag=0)
        if status == 3:
            qs = qs.filter(Q(status=1) | Q(status=2))
        else:
            qs = qs.filter(status=status)
        paginator = Paginator(qs, size)

        try:
            businesses = paginator.page(page)
        except EmptyPage:
            businesses = paginator.page(1)

        results = []

        for item in businesses:
            team_dict = [model_to_dict(member) for member in
                         BusinessTeamMember.objects.filter(business_id=item.id, project_id=item.cur_project_id)]

            project = Project.objects.filter(pk=item.project_id).first()
            if project:
                project_dict = {
                    'id': project.id, 'name': project.name,
                    'office_item': project.officeItem.name if project.officeItem else None
                }
            else:
                project_dict = None

            node = FlowNode.objects.filter(pk=item.node_id).first() if item.node else None
            if node:
                cur_node = {
                    'id': node.id, 'name': node.name, 'condition': node.condition,
                    'process_type': node.process.type if node.process else None,
                }
            else:
                cur_node = None

            business = {
                'id': item.id, 'name': item.name, 'show_nickname': item.show_nickname,
                'start_time': item.start_time.strftime('%Y-%m-%d') if item.start_time else None,
                'end_time': item.end_time.strftime('%Y-%m-%d') if item.end_time else None,
                'create_time': item.create_time.strftime('%Y-%m-%d %H:%M:%S') if item.create_time else None,
                'team': team_dict, 'status': item.status, 'created_by': user_simple_info(item.created_by_id),
                'node_id': item.node_id,
                'project': project_dict,
                'huanxin_id': item.huanxin_id,
                'node': cur_node, 'flow_id': project.flow_id if project else None,
                'officeItem': model_to_dict(item.officeItem) if item.officeItem else None,
                'jumper_id': item.jumper_id
            }
            results.append(business)

        paging = {
            'count': paginator.count,
            'has_previous': businesses.has_previous(),
            'has_next': businesses.has_next(),
            'num_pages': paginator.num_pages,
            'cur_page': businesses.number,
            'page_size': size
        }

        resp = code.get_msg(code.SUCCESS)
        resp['d'] = {'results': results, 'paging': paging}
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    except Exception as e:
        logger.exception('api_business_list Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


def api_business_detail(request):
    resp = auth_check(request, "GET")
    observable = False
    if resp != {}:
        observable = True

    try:
        business_id = request.GET.get("business_id")  # 实验id
        business = Business.objects.filter(pk=business_id, del_flag=0).first()
        if business is None:
            resp = code.get_msg(code.BUSINESS_NOT_EXIST)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        if not observable and BusinessTransPath.objects.filter(business_id=business_id,
                                                               project_id=business.cur_project_id).count() == 0:
            resp = api_business_start(request)
            if resp != 'success':
                if business.jumper_id and resp['c'] == code.TEAM_MEMBER_NOT_EXIST:
                    notification = TNotifications.objects.filter(type='businessMoreTeammate_' + business_id).first()
                    if notification:
                        resp = code.get_msg(code.SUCCESS)
                        resp['d'] = {
                            'jumper_member_not_exist': True,
                            'notification': eval(notification.content)
                        }
                return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        business = Business.objects.filter(pk=business_id, del_flag=0).first()
        data = get_business_detail(business)

        # 三期记录用户最后一次进入的实验id
        if (not observable):
            user = request.user
            user.last_business_id = business_id
            user.save()

        user_role_allocs = []
        if business.status == 1:
            control_status = 1
            path_id = None
            is_parallel = 0
        else:
            path = BusinessTransPath.objects.filter(business=business).last()
            control_status = path.control_status
            path_id = path.pk
            if (not observable):
                user_role_allocs = get_role_allocs_status_by_user(business, path, request.user)
                data['with_user_nodes'] = get_user_with_node_on_business(business, request.user)
            is_parallel = 1 if path.node.parallel_node_start == 1 else 0

        data['control_status'] = control_status
        data['path_id'] = path_id
        data['user_role_allocs'] = user_role_allocs
        data['is_parallel'] = is_parallel
        resp = code.get_msg(code.SUCCESS)
        resp['d'] = data

        # 三期 - 到达指定环节还有角色没有设置提示设置角色
        # node = business.node
        # if node:
        #     role_name_not_set = []
        #     role_allocs_node = BusinessRoleAllocation.objects.filter(business=business, node=node,
        #                                                              project_id=business.cur_project_id,
        #                                                              can_take_in=True)
        #     for role_alloc in role_allocs_node:
        #         btmExist = BusinessTeamMember.objects.filter(business=business, business_role=role_alloc.role,
        #                                                      no=role_alloc.no, project_id=business.cur_project_id,
        #                                                      del_flag=0).exists()
        #         if btmExist and role_alloc.role.type != const.ROLE_TYPE_OBSERVER:
        #             continue
        #         role_name_not_set.append(role_alloc.role.name)
        #     if len(role_name_not_set) > 0:
        #         logger.info('当前实验环节，以下角色还没有设置: ' + ','.join(role_name_not_set))
        #         # resp['c'] = code.get_msg(code.EXPERIMENT_ROLE_NOT_SET)
        #         resp['m'] = '当前实验环节，以下角色还没有设置: ' + ','.join(role_name_not_set)
        #         data['role_not_set'] = resp['m']
        #         return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    except Exception as e:
        logger.exception('api_business_detail Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


def api_business_start(request):
    business_id = request.GET.get("business_id")  # 实验id
    business = Business.objects.filter(pk=business_id).first()  # get Business
    logger.info('api_business_start:business_id=%s' % business_id)
    if business is None:
        resp = code.get_msg(code.BUSINESS_NOT_EXIST)
        return resp

    # get Project that this business is based on
    project = Project.objects.get(pk=business.cur_project_id)
    # 验证项目中是否有未配置的跳转项目 todo
    if not check_jump_project(project):
        resp = code.get_msg(code.BUSINESS_JUMP_PROJECT_SETUP_ERROR)
        return resp
    # get First Node ID and Node by project Flow ID
    first_node_id = get_start_node(project.flow_id)
    node = FlowNode.objects.get(pk=first_node_id)
    # get All Business Roles to check if all users are allocated to business Role Alloc
    businessRoles = BusinessRole.objects.filter(business=business,
                                                project_id=business.cur_project_id)  # get all Business Roles
    for role in businessRoles:
        for no in range(1, role.capacity + 1):
            teamMembers = BusinessTeamMember.objects.filter(business=business, business_role=role, no=no,
                                                            project_id=business.cur_project_id,
                                                            del_flag=0)  # get all team members with same business, role, no to check if user is allocated to this allocation
            if teamMembers.count() == 0:
                resp = code.get_msg(code.TEAM_MEMBER_NOT_EXIST)
                return resp

    # teamMembers = BusinessTeamMember.objects.filter(business=business, del_flag=0)
    # ids = list(teamMembers.values_list('user_id'))

    # 注册所有的群组用户到环信群组
    # easemob_success, easemob_result = easemob.create_groups(str(business.pk), str(request.user.pk), ids)
    # logger.info(u'easemob create_groups:{}{}'.format(easemob_success, easemob_result))

    # if easemob_success is False:
    #     resp = code.get_msg(code.BUSINESS_START_FAILED)
    #     return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    # get Start Role Allocation
    startRoleAlloc = BusinessRoleAllocation.objects.filter(business=business, node=node, can_start=1,
                                                           can_take_in=1, project_id=business.cur_project_id).first()
    # check if this user is Start User
    isStartUser = BusinessTeamMember.objects.filter(business=business, user=request.user,
                                                    business_role=startRoleAlloc.role,
                                                    project_id=business.cur_project_id,
                                                    no=startRoleAlloc.no).exists()
    if not isStartUser:
        resp = code.get_msg(code.BUSINESS_NO_ACCESS_TO_START)
        return resp
    # get all allocations take part in this first node of this business
    allocations = BusinessRoleAllocation.objects.filter(business=business, node=node, can_take_in=1,
                                                        project_id=business.cur_project_id)
    with transaction.atomic():
        # Create Business TransPath
        step = BusinessTransPath.objects.filter(business_id=business.id).count() + 1
        path = BusinessTransPath.objects.create(business=business, node=node,
                                                project_id=business.cur_project_id, task_id=node.task_id, step=step)
        for item in allocations:
            if item.can_brought:
                come_status = 1
            else:
                come_status = 9
            # 三期 - 不能直接创建， 在service中结束并走向下一环节的时候会创建角色状态，这里再创建一次就重复了
            brses = BusinessRoleAllocationStatus.objects.filter(
                business=business,
                business_role_allocation=item,
                # path=path
            )
            if brses.count() > 0:  # 存在则更新
                brs = brses.first()
                brs.come_status = come_status
                brs.save()
            else:  # 不存在则创建
                BusinessRoleAllocationStatus.objects.update_or_create(business=business,
                                                                      business_role_allocation=item,
                                                                      # path=path,
                                                                      come_status=come_status)

        # 环信id
        # huanxin_id = easemob_result['data']['groupid']
        # exp.huanxin_id = huanxin_id
        # 设置实验环节为开始环节,改变实验状态
        business.node = node
        business.path_id = path.id
        business.save()

        if node.parallel_node_start == 1:
            business.parallel_count = 1
            business.save()
            for item in FlowTrans.objects.filter(incoming=node.task_id):
                fn = FlowNode.objects.get(task_id=item.outgoing, flow_id=project.flow_id)
                business.parallel_nodes.create(
                    node=fn
                )
                fnallocs = BusinessRoleAllocation.objects.filter(business=business, node=fn, can_take_in=1,
                                                                 project_id=business.cur_project_id)
                for fnusers in fnallocs:
                    if fnusers.can_brought:
                        come_status = 1
                    else:
                        come_status = 9
                    # 三期 - 不能直接创建， 在service中结束并走向下一环节的时候会创建角色状态，这里再创建一次就重复了
                    brses = BusinessRoleAllocationStatus.objects.filter(
                        business=business,
                        business_role_allocation=fnusers,
                        # path=path
                    )
                    if brses.count() > 0:  # 存在则更新
                        brs = brses.first()
                        brs.come_status = come_status
                        brs.save()
                    else:  # 不存在则创建
                        BusinessRoleAllocationStatus.objects.update_or_create(business=business,
                                                                              business_role_allocation=fnusers,
                                                                              # path=path,
                                                                              come_status=come_status)

    return 'success'


# 实验环节详情
def api_business_node_detail(request):
    resp = auth_check(request, "GET")
    observable = False
    if resp != {}:
        observable = True

    try:
        business_id = request.GET.get('business_id', None)  # 实验id
        node_id = request.GET.get("node_id", None)  # 环节id
        roleAllocID = request.GET.get("roleAllocID", None)  # 角色id

        business = Business.objects.filter(pk=business_id, del_flag=0).first()
        if business is None:
            resp = code.get_msg(code.BUSINESS_NOT_EXIST)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        project = Project.objects.get(pk=business.cur_project_id)

        # 验证环节是否存在
        node = FlowNode.objects.filter(pk=node_id).first()
        if node is None:
            resp = code.get_msg(code.BUSINESS_NODE_ERROR)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        # 获取上个环节
        # pre_node = get_business_pre_node_path(business)
        pre_node_id = None
        # if pre_node:
        #     pre_node_id = pre_node.node_id

        # path = BusinessTransPath.objects.filter(business=business).last()

        # 当前环节所有角色状态
        role_alloc_status_list = get_all_simple_role_allocs_status(business, node)

        # 当前用户可选角色
        role_alloc_list = []
        if not observable and roleAllocID != 'observable':
            btmQs = BusinessTeamMember.objects.filter(business=business, project_id=business.cur_project_id,
                                                      user=request.user)
            for btm in btmQs:
                try:
                    roleAlloc = BusinessRoleAllocation.objects.filter(business=business, node=node,
                                                                      role=btm.business_role,
                                                                      no=btm.no, project_id=business.cur_project_id,
                                                                      can_take_in=True).first()
                    roleAllocStatus = BusinessRoleAllocationStatus.objects.filter(business=business,
                                                                                  business_role_allocation=roleAlloc).first()
                    role_alloc_list.append({
                        'alloc_id': roleAlloc.id, 'come_status': roleAllocStatus.come_status, 'no': roleAlloc.no,
                        'sitting_status': roleAllocStatus.sitting_status, 'stand_status': roleAllocStatus.stand_status,
                        'vote_status': roleAllocStatus.vote_status, 'show_status': roleAllocStatus.show_status,
                        'speak_times': 0,
                        'role': model_to_dict(roleAlloc.role), 'can_terminate': roleAlloc.can_terminate,
                        'can_brought': roleAlloc.can_brought
                    })
                except:
                    continue

            # 是否投票
            has_vote = BusinessRoleAllocationStatus.objects.filter(business=business,
                                                                   business_role_allocation_id=roleAllocID,
                                                                   business_role_allocation__can_take_in=1,
                                                                   business_role_allocation__node=node,
                                                                   # path=path,
                                                                   vote_status=0).exists()
            # if path.vote_status == 1:
        end_vote = False
        # else:
        #     end_vote = True

        # 场景动作列表
        process_action_list = []
        # 场景信息
        if node.process:
            pro = node.process
            process = {
                'id': pro.id, 'name': pro.name, 'type': pro.type, 'can_switch': pro.can_switch,
                'file': pro.file.url if pro.file else None,
                'image': pro.image.url if pro.image else None
            }
        else:
            process = None

        # 查询小组组长
        can_opt = True if business.created_by == request.user else False

        # 实验心得
        # experience = ExperimentExperience.objects.filter(experiment_id=exp.id, created_by=request.user.pk).first()
        # experience_data = {'status': 1, 'content': ''}
        # if experience:
        #     experience_data = {
        #         'id': experience.id, 'content': experience.content, 'status': experience.status,
        #         'created_by': user_simple_info(experience.created_by),
        #         'create_time': experience.create_time.strftime('%Y-%m-%d')
        #     }

        resp = code.get_msg(code.SUCCESS)
        resp['d'] = {
            'role_allocs': role_alloc_list, 'process': process, 'process_actions': process_action_list,
            'can_opt': can_opt,
            'role_alloc_status': role_alloc_status_list, 'id': business.id, 'name': business.name,
            # 'experience': experience_data,
            'node': {'id': node.id, 'name': node.name, 'condition': node.condition}, 'pre_node_id': pre_node_id,
            'huanxin_id': business.huanxin_id, 'control_status': 1,
            'entire_graph': project.entire_graph,
            # 'leader': team.leader if team else None,
            'flow_id': project.flow_id,
            'has_vote': False if not observable and roleAllocID != 'observable' and has_vote else True,
            'end_vote': end_vote
        }
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    except Exception as e:
        logger.exception('api_business_node_detail Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


def api_business_trans_path(request):
    resp = auth_check(request, "GET")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
    try:
        business_id = request.GET.get("business_id")  # 实验ID
        # 判断实验是否存在
        business = Business.objects.filter(pk=business_id, del_flag=0).first()
        if business:
            stop_node = []
            if business.status == const.EXPERIMENT_FINISHED:
                node = {'is_finished': True}
            else:
                node = []
                if business.node.parallel_node_start == 0:
                    next_node = FlowNode.objects.filter(pk=business.node_id).first()
                    node.append(next_node.task_id)
                else:
                    for pnItem in business.parallel_nodes.all():
                        if pnItem.node.is_parallel_merging == 1 and \
                                        business.parallel_passed_nodes.filter(
                                            node__task_id__in=FlowTrans.objects.filter(
                                                outgoing=pnItem.node.task_id).values_list('incoming',
                                                                                          flat=True)).count() != FlowTrans.objects.filter(
                                    outgoing=pnItem.node.task_id).count():
                            stop_node.append(pnItem.node.task_id)
                        else:
                            node.append(pnItem.node.task_id)
            project = Project.objects.get(pk=business.cur_project_id)
            flow = Flow.objects.get(pk=project.flow_id)
            paths = list(business.parallel_passed_nodes.all().values_list('node__task_id', flat=True))
            paths += list(BusinessTransPath.objects.filter(business_id=business.id,
                                                           project_id=business.cur_project_id).values_list('task_id',
                                                                                                           flat=True))
            xml = bpmn_color(flow.xml, paths, node, business.node.parallel_node_start, stop_node)

            resp = code.get_msg(code.SUCCESS)
            if business.node.parallel_node_start == 1:
                node.append("")
            resp['d'] = {'xml': xml, 'node': node}
        else:
            resp = code.get_msg(code.BUSINESS_NOT_EXIST)
    except Exception as e:
        logger.exception('api_business_trans_path Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)
    return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


# 实验环节聊天消息列表
def api_business_node_messages(request):
    resp = auth_check(request, "GET")
    observable = False
    if resp != {}:
        observable = True

    try:
        business_id = request.GET.get("business_id")  # 实验ID
        node_id = request.GET.get("node_id", None)  # 环节id
        is_paging = int(request.GET.get('is_paging', 1))  # 是否进行分页（1，分页；0，不分页）
        page = int(request.GET.get("page", 1))
        size = int(request.GET.get("size", const.ROW_SIZE))

        business = Business.objects.filter(pk=business_id).first()
        if business:
            path = BusinessTransPath.objects.filter(business_id=business_id).last()
            data = get_node_path_messages_on_business(business, node_id, path.pk, is_paging, page, size)

            for i in range(len(data['results'])):
                file_id = data['results'][i]['file_id']
                mid = data['results'][i]['id']
                m_ext = data['results'][i]['ext']
                opt_status = data['results'][i]['opt_status']

                # 更新扩展
                ext = json.loads(m_ext)
                ext['id'] = mid
                ext['business_role_alloc'] = model_to_dict(
                    BusinessRoleAllocation.objects.filter(pk=ext['role_alloc_id']).first())
                if opt_status == 1:
                    ext['opt_status'] = True
                else:
                    ext['opt_status'] = False

                data['results'][i]['type'] = 'groupchat'
                data['results'][i]['ext'] = ext
                data['results'][i]['to'] = business.huanxin_id

                # 音频文件
                if file_id:
                    audio = BusinessMessageFile.objects.filter(pk=file_id).first()
                    if audio:
                        # data['results'][i]['url'] = const.WEB_HOST + audio.file.url
                        data['results'][i]['url'] = const.WEB_HOST + audio.file.url
                        data['results'][i]['filename'] = audio.file.name
                        data['results'][i]['secret'] = ''
                        data['results'][i]['length'] = audio.length

            resp = code.get_msg(code.SUCCESS)
            resp['d'] = data
        else:
            resp = code.get_msg(code.BUSINESS_NOT_EXIST)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    except Exception as e:
        logger.exception('api_business_node_messages Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


# 实验功能按钮
def api_business_node_function(request):
    resp = auth_check(request, "GET")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    try:
        business_id = request.GET.get('business_id', None)  # 实验id
        node_id = request.GET.get("node_id", None)  # 环节id
        role_alloc_id = request.GET.get("role_alloc_id", None)  # 角色id

        user_id = request.user.id
        business = Business.objects.filter(pk=business_id).first()
        if business is None:
            resp = code.get_msg(code.BUSINESS_NOT_EXIST)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        project = Project.objects.filter(pk=business.cur_project_id).first()
        # 验证环节是否存在
        node = FlowNode.objects.filter(pk=node_id).first()
        if node is None:
            resp = code.get_msg(code.BUSINESS_NODE_ERROR)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        # 路径
        path = BusinessTransPath.objects.filter(business_id=business_id).last()
        # 判断该实验环节是否存在该角色
        if role_alloc_id is None:
            brases = BusinessRoleAllocationStatus.objects.filter(
                business_id=business_id,
                business_role_allocation__node_id=node_id,
                business_role_allocation__project_id=business.cur_project_id,
                # path_id=path.pk
            )
            role_alloc = None
            for bras in brases:
                bra = bras.business_role_allocation
                btm = BusinessTeamMember.objects.filter(business=business, user_id=user_id, business_role=bra.role,
                                                        no=bra.no, del_flag=0, project_id=business.cur_project_id)
                if btm.exists():
                    role_alloc = bra
                    break
            if role_alloc is None:
                resp = code.get_msg(code.BUSINESS_NODE_ROLE_NOT_EXIST)
                return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
            else:
                role_alloc_id = role_alloc.id

        role_alloc_status = BusinessRoleAllocationStatus.objects.filter(business_id=business_id,
                                                                        business_role_allocation_id=role_alloc_id).first()
        if role_alloc_status is None:
            resp = code.get_msg(code.BUSINESS_NODE_ROLE_NOT_EXIST)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        # 三期 - 根据上一步骤自动入席 判断是否入席
        bps = BusinessPositionStatus.objects.filter(business_id=business_id,
                                                    business_role_allocation__node_id=path.node_id, path_id=path.id,
                                                    business_role_allocation_id=role_alloc_id)
        if bps:
            business_position_status = bps.first()
            if business_position_status.sitting_status == 2:  # 已入席
                role_alloc_status.sitting_status = const.SITTING_DOWN_STATUS

        # 用户角色状态
        # 当前用户可选角色
        role_alloc_list = get_role_allocs_status_simple_by_user(business, node, path, user_id)

        # 功能动作列表根据环节和角色分配过滤
        role_alloc = BusinessRoleAllocation.objects.filter(pk=role_alloc_id).first()
        flow_action_ids = []
        process_action_ids = []
        if role_alloc:
            flowRoleAlloc = FlowRoleAllocation.objects.get(pk=role_alloc.flow_role_alloc_id)
            flow_actions = FlowRoleActionNew.objects.filter(flow_id=project.flow_id, node_id=node_id,
                                                            role_id=flowRoleAlloc.role_id, no=flowRoleAlloc.no,
                                                            del_flag=0).first()

            process_actions = ProcessRoleActionNew.objects.filter(flow_id=project.flow_id, node_id=node_id,
                                                                  role_id=flowRoleAlloc.role_id, no=flowRoleAlloc.no,
                                                                  del_flag=0).first()
            if process_actions and process_actions.actions:
                process_action_ids = json.loads(process_actions.actions)
            else:
                process_action_ids = [58, 60, 61]
            if flow_actions and flow_actions.actions:
                flow_action_ids = json.loads(flow_actions.actions)
            else:
                flow_action_ids = [17, 18, 19, 20, 21, 22, 23, 24, 25, 26, 27, 28, 29, 31, 32, 33, 34]

        # 当前角色动画
        process_action_list = get_role_alloc_process_actions(business, path, role_alloc_id, process_action_ids)

        # 功能按钮
        # 是否有结束环节的权限
        can_terminate = role_alloc.can_terminate

        function_action_list = []
        function_actions = FlowAction.objects.filter(id__in=flow_action_ids, del_flag=0)

        # 判断按钮是否可用
        for item in function_actions:
            if can_terminate:
                disable = False
            else:
                disable = False
                # 判断表达管理
                if path.control_status == 2:
                    if role_alloc_status:
                        if item.cmd == const.ACTION_DOC_SHOW:
                            if role_alloc_status.show_status != 1:
                                disable = True
                        if item.cmd == const.ACTION_DOC_SUBMIT:
                            if role_alloc_status.submit_status != 1:
                                disable = True
                    else:
                        disable = True
                else:
                    # 申请发言状态
                    if item.cmd == const.ACTION_ROLE_APPLY_SPEAK:
                        disable = True
                    if item.cmd == const.ACTION_DOC_APPLY_SUBMIT:
                        disable = True
                    if item.cmd == const.ACTION_DOC_APPLY_SHOW:
                        disable = True

            # 入席、退席互斥
            if role_alloc_status.sitting_status == const.SITTING_UP_STATUS:
                if item.cmd == const.ACTION_ROLE_LETOUT or item.cmd == const.ACTION_ROLE_LETIN \
                        or item.cmd == const.ACTION_ROLE_REQUEST_SIGN or item.cmd == const.ACTION_ROLE_SCHEDULE_REPORT \
                        or item.cmd == const.ACTION_ROLE_HIDE:
                    disable = True
                    # 约见
                    # if item.cmd == const.ACTION_ROLE_MEET and not can_brought:
                    #     disable = True
            else:
                if item.cmd == const.ACTION_ROLE_SHOW:
                    disable = True

                if item.cmd == const.ACTION_ROLE_HIDE:
                    report_exists = BusinessReportStatus.objects.filter(business_id=business_id,
                                                                        business_role_allocation__node_id=node_id,
                                                                        path_id=path.id,
                                                                        business_role_allocation_id=role_alloc_id,
                                                                        schedule_status=const.SCHEDULE_UP_STATUS).exists()
                    if report_exists:
                        disable = True

                # 起立坐下互斥
                if item.cmd == const.ACTION_ROLE_STAND:
                    if role_alloc_status.stand_status == 1:
                        disable = True

                if item.cmd == const.ACTION_ROLE_SITDOWN:
                    if role_alloc_status.stand_status == 2:
                        disable = True
                # 约见
                if item.cmd == const.ACTION_ROLE_MEET:
                    disable = True

            # 判断报告按钮状态
            if item.cmd == const.ACTION_ROLE_TOWARD_REPORT:
                disable = True
                report_exists = BusinessReportStatus.objects.filter(business_id=business_id,
                                                                    business_role_allocation__node_id=node_id,
                                                                    path_id=path.id,
                                                                    business_role_allocation_id=role_alloc_id,
                                                                    schedule_status=const.SCHEDULE_OK_STATUS).exists()
                if report_exists:
                    disable = False

            if item.cmd == const.ACTION_ROLE_EDN_REPORT:
                disable = True
                report_exists = BusinessReportStatus.objects.filter(business_id=business_id,
                                                                    business_role_allocation__node_id=node_id,
                                                                    path_id=path.id,
                                                                    business_role_allocation_id=role_alloc_id,
                                                                    schedule_status=const.SCHEDULE_UP_STATUS).exists()
                if report_exists:
                    disable = False

            btn = {
                'id': item.id, 'name': item.name, 'cmd': item.cmd, 'disable': disable
            }
            function_action_list.append(btn)

        resp = code.get_msg(code.SUCCESS)

        # 三期，如果进来的是老师观察者角色， 没有任何功能按钮， tmd
        if role_alloc.role.type == const.ROLE_TYPE_OBSERVER:
            function_action_list = []
            process_action_list = []

        resp['d'] = {'function_action_list': function_action_list,
                     'process_action_list': process_action_list,
                     'user_role_allocs': role_alloc_list}
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
    except Exception as e:
        logger.exception('api_experiment_node_function Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


def api_business_node_role_docs(request):
    resp = auth_check(request, "GET")
    observable = False
    if resp != {}:
        observable = True

    try:
        business_id = request.GET.get('business_id', None)  # 实验id
        node_id = request.GET.get("node_id", None)  # 环节id
        role_alloc_id = request.GET.get("role_alloc_id", None)  # 角色id

        business = Business.objects.filter(pk=business_id, del_flag=0).first()
        if business is None:
            resp = code.get_msg(code.BUSINESS_NOT_EXIST)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        project = Project.objects.get(pk=business.cur_project_id)

        # 验证环节是否存在
        node = FlowNode.objects.filter(pk=node_id).first()
        if node is None:
            resp = code.get_msg(code.BUSINESS_NODE_ERROR)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        # 路径
        path = BusinessTransPath.objects.filter(business_id=business_id).last()
        if not observable and role_alloc_id != 'observable':
            # 判断该实验环节是否存在该角色
            if role_alloc_id is None:
                user = request.user
                brases = BusinessRoleAllocationStatus.objects.filter(
                    business_id=business_id,
                    business_role_allocation__node_id=node_id,
                    business_role_allocation__project_id=business.cur_project_id,
                    # path_id=path.pk
                )
                role_alloc = None
                for bras in brases:
                    bra = bras.business_role_allocation
                    btm = BusinessTeamMember.objects.filter(business=business, user_id=user.id, business_role=bra.role,
                                                            no=bra.no, del_flag=0, project_id=business.cur_project_id)
                    if btm.exists():
                        role_alloc = bra
                        break
                if role_alloc is None:
                    resp = code.get_msg(code.BUSINESS_NODE_ROLE_NOT_EXIST)
                    return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
                else:
                    role_alloc_id = role_alloc.id
            # 前面所有环节素材
            pre_doc_list = get_pre_node_role_alloc_docs(business, node_id, project.pk, role_alloc_id)

        # 获取该环节角色项目所有素材
        docs = get_node_role_alloc_docs(business, node_id, project.pk, project.flow_id, role_alloc_id)

        resp = code.get_msg(code.SUCCESS)
        resp['d'] = {
            'operation_guides': docs['operation_guides'],
            'project_tips_list': docs['project_tips_list'],
            'cur_doc_list': docs['cur_doc_list'],
            'pre_doc_list': pre_doc_list if not observable and role_alloc_id != 'observable' else [],
            'id': business.id, 'name': business.name,
            'flow_id': project.flow_id
        }
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    except Exception as e:
        logger.exception('api_business_node_role_docs Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


def api_business_node_observable_list(request):
    try:
        size = request.GET.get("size", None)
        businesses = Business.objects.filter(Q(del_flag=0, status=2)).exclude(Q(node__isnull=True)).order_by(
            '-create_time')
        observable_business_node_list = []
        for business in businesses:
            parallel_node_ids = list(business.parallel_nodes.all().distinct().values_list('node_id', flat=True))
            if len(parallel_node_ids) == 0:
                nodes = FlowNode.objects.filter(pk=business.node_id, look_on=True, del_flag=0)
            else:
                nodes = FlowNode.objects.filter(id__in=parallel_node_ids, look_on=True, del_flag=0)
            for node in nodes:
                observable_business_node_list.append({
                    'business': {
                        'id': business.id, 'name': business.name,
                        'start_time': business.start_time.strftime('%Y-%m-%d') if business.start_time else None,
                        'end_time': business.end_time.strftime('%Y-%m-%d') if business.end_time else None,
                        'create_time': business.create_time.strftime(
                            '%Y-%m-%d %H:%M:%S') if business.create_time else None,
                        'status': business.status, 'created_by': user_simple_info(business.created_by_id),
                        'officeItem': model_to_dict(business.officeItem) if business.officeItem else None,
                        'jumper_id': business.jumper_id
                    },
                    'node': {
                        'id': node.id, 'flow_id': node.flow_id, 'name': node.name, 'process_type': node.process.type
                    }
                })
                if size and len(observable_business_node_list) == 5:
                    break;

            if size and len(observable_business_node_list) == 5:
                break;

        resp = code.get_msg(code.SUCCESS)
        resp['d'] = {
            'nodes': observable_business_node_list
        }
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    except Exception as e:
        logger.exception('api_business_node_observable_list Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


def api_business_messages(request):
    resp = auth_check(request, "GET")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    try:
        business_id = request.GET.get("business_id")  # 实验ID
        business = Business.objects.filter(pk=business_id).first()
        if business:
            paths = BusinessTransPath.objects.filter(business_id=business_id)
            node_list = []
            for item in paths:
                messages = BusinessMessage.objects.filter(business_id=business_id,
                                                          business_role_allocation__node_id=item.node_id,
                                                          path_id=item.id).order_by(
                    'timestamp')
                message_list = []
                for m in messages:
                    ext = json.loads(m.ext)
                    ext['id'] = m.id
                    ext['opt_status'] = m.opt_status
                    message = {
                        'id': m.id, 'from': m.user_id, 'to': business.huanxin_id, 'msg_type': m.msg_type,
                        'data': m.msg, 'type': 'groupchat', 'ext': ext
                    }
                    if m.file_id:
                        audio = BusinessMessageFile.objects.filter(pk=m.file_id).first()
                        if audio:
                            message['url'] = const.WEB_HOST + audio.file.url
                            message['filename'] = audio.file.name
                            message['secret'] = ''
                            message['length'] = audio.length

                    message_list.append(message)
                node = FlowNode.objects.filter(pk=item.node_id, del_flag=0).first()
                node_list.append({
                    'id': node.id, 'name': node.name, 'process_type': node.process.type,
                    'messages': message_list, 'count': messages.count()
                })

            resp = code.get_msg(code.SUCCESS)
            resp['d'] = {'results': node_list}
        else:
            resp = code.get_msg(code.BUSINESS_NOT_EXIST)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    except Exception as e:
        logger.exception('api_business_messages Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


# 实验所属项目素材查询
# modified by ser -- edit_module param is added
def api_business_templates(request):
    resp = auth_check(request, "GET")
    observable = False
    if resp != {}:
        observable = True

    try:
        business_id = request.GET.get('business_id', None)  # 实验id
        node_id = request.GET.get('node_id', None)
        role_alloc_id = request.GET.get('role_alloc_id', None)
        usage = request.GET.get("usage", None)  # 用途
        edit_module = request.GET.get('edit_module', None)

        # added by ser start
        # prevent initial loading on node
        if edit_module is not None and role_alloc_id is None:
            resp = code.get_msg(code.SUCCESS)
            resp['d'] = {};
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
        # added by ser end

        if None in (business_id, node_id):
            resp = code.get_msg(code.PARAMETER_ERROR)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        business = Business.objects.filter(pk=business_id, del_flag=0).first()
        bra = BusinessRoleAllocation.objects.filter(pk=role_alloc_id).first() if not observable and role_alloc_id != 'observable' else None
        pra = ProjectRoleAllocation.objects.filter(pk=bra.project_role_alloc_id).first() if bra else None
        if business:
            if usage and usage == '3' and role_alloc_id != 'observable' and not observable:
                if role_alloc_id is None:
                    resp = code.get_msg(code.PARAMETER_ERROR)
                    return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

                # 复制编辑模板
                if edit_module is None:
                    doc_ids = ProjectDocRole.objects.filter(project_id=business.cur_project_id, node_id=node_id,
                                                    role_id=pra.role_id, no=pra.no).values_list('doc_id', flat=True)
                else:
                    doc_ids = ProjectDocRole.objects.filter(project_id=business.cur_project_id, node_id=node_id,
                                                        no=pra.no).values_list('doc_id', flat=True)

                project_docs = ProjectDoc.objects.filter(pk__in=doc_ids, usage=3)
                for doc in project_docs:
                    if edit_module is None:
                        is_exists = BusinessDocContent.objects.filter(business_id=business_id, node_id=node_id,
                                                                      doc_id=doc.pk,
                                                                      business_role_allocation_id=role_alloc_id).exists()
                    else:
                        is_exists = BusinessDocContent.objects.filter(business_id=business_id, node_id=node_id,
                                                                      doc_id=doc.pk).exists()
                    if not is_exists:
                        path = business_template_save(business.pk, node_id, doc.name, doc.content)
                        BusinessDocContent.objects.create(business_id=business.pk, node_id=node_id, doc_id=doc.pk,
                                                          business_role_allocation_id=role_alloc_id, name=doc.name,
                                                          content=doc.content,
                                                          created_by=request.user, file_type=1, file=path)

            # modified by ser -- edit_module added for param
            doc_list = get_business_templates(business, node_id, role_alloc_id, usage, pra, edit_module)
            resp = code.get_msg(code.SUCCESS)
            resp['d'] = doc_list
        else:
            resp = code.get_msg(code.BUSINESS_NOT_EXIST)

        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    except Exception as e:
        logger.exception('api_business_templates Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


# 个人笔记列表
def api_business_note_list(request):
    resp = auth_check(request, "GET")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    try:
        business_id = request.GET.get('business_id', None)  # 实验id

        business = Business.objects.filter(pk=business_id).first()
        if business:
            project = Project.objects.get(pk=business.cur_project_id)
            nodes = FlowNode.objects.filter(flow_id=project.flow_id)

            note_list = []
            for item in nodes:
                note = BusinessNotes.objects.filter(business_id=business_id, node_id=item.id,
                                                    created_by_id=request.user.id, del_flag=0).first()
                can_edit = True if business.node_id == item.id else False
                if note:
                    note_dict = {'id': note.id, 'content': note.content}
                else:
                    note_dict = None

                note_list.append({
                    'node_id': item.id, 'node_name': item.name, 'note': note_dict,
                    'can_edit': can_edit
                })

            resp = code.get_msg(code.SUCCESS)
            resp['d'] = note_list
        else:
            resp = code.get_msg(code.BUSINESS_NOT_EXIST)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    except Exception as e:
        logger.exception('api_business_note_list Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


# 创建实验笔记
def api_business_note_create(request):
    resp = auth_check(request, "POST")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    try:
        business_id = request.POST.get("business_id")  # 实验id
        node_id = request.POST.get("node_id")  # 环节id
        content = request.POST.get("content")  # 内容

        business = Business.objects.filter(pk=business_id).first()
        if business:
            # 验证实验环节是否在该环节
            if business.node_id == int(node_id):
                note, created = BusinessNotes.objects.update_or_create(business_id=business_id,
                                                                       node_id=node_id,
                                                                       created_by_id=request.user.id,
                                                                       defaults={'content': content})

                resp = code.get_msg(code.SUCCESS)
                resp['d'] = {
                    'id': note.id, 'content': note.content, 'node_id': note.node_id,
                    'business_id': note.business.id,
                    'created_by': user_simple_info(note.created_by.id)
                }
            else:
                resp = code.get_msg(code.BUSINESS_NODE_ERROR)
        else:
            resp = code.get_msg(code.BUSINESS_NOT_EXIST)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    except Exception as e:
        logger.exception('api_business_note_create Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


# Get No-Deleted Business
def api_business_list_nodel(request):
    if request.session['login_type'] in [2, 6, 3, 7]:
        resp = auth_check(request, "GET")
        if resp != {}:
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        try:
            userIDInfo = request.user.id
            projectAvailableList = []
            if request.session['login_type'] == 2:
                projectAvailableList = getProjectIDByGroupManager(userIDInfo)
            elif request.session['login_type'] == 3:
                projectAvailableList = getProjectIDByCompanyManager(userIDInfo)
            elif request.session['login_type'] == 6:
                projectAvailableList = getProjectIDByGroupManagerAssistant(userIDInfo)
            elif request.session['login_type'] == 7:
                projectAvailableList = getProjectIDByCompanyManagerAssistant(userIDInfo)
            search = request.GET.get("search", None)  # 关键字
            page = int(request.GET.get("page", 1))  # 页码
            size = int(request.GET.get("size", const.ROW_SIZE))  # 页面条数

            qs = Business.objects.filter(Q(project_id__in=projectAvailableList) & Q(del_flag=0))

            if search:
                if search == '已完成':
                    qs = qs.filter(status=9)
                else:
                    qs = qs.filter(Q(name__icontains=search) | Q(pk__icontains=search))

            paginator = Paginator(qs, size)

            try:
                business = paginator.page(page)
            except EmptyPage:
                business = paginator.page(1)

            results = []

            for item in business:
                teamMembers = list(
                    BusinessTeamMember.objects.filter(business_id=item.id, project_id=item.cur_project_id).values_list(
                        'user__name', flat=True))
                project = Project.objects.get(pk=item.project_id)
                project_name = project.name
                workflow_name = Flow.objects.get(pk=project.flow_id).name

                user_roles = []
                bus = {
                    'id': item.id, 'name': item.name, 'project_name': project_name,
                    'workflow_name': workflow_name, 'officeItem': item.officeItem.name,
                    'start_time': item.create_time.strftime('%Y-%m-%d %H:%M:%S') if item.create_time else None,
                    'end_time': item.finish_time.strftime('%Y-%m-%d %H:%M:%S') if item.finish_time else None,
                    'members': teamMembers, 'created_by': item.created_by.name,
                    'status': '已完成' if item.status == 9 else '未完成'
                }
                results.append(bus)
            # 分页信息
            paging = {
                'count': paginator.count,
                'has_previous': business.has_previous(),
                'has_next': business.has_next(),
                'num_pages': paginator.num_pages,
                'cur_page': business.number,
                'page_size': size
            }

            resp = code.get_msg(code.SUCCESS)
            resp['d'] = {'results': results, 'paging': paging}
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        except Exception as e:
            logger.exception('api_experiment_list Exception:{0}'.format(str(e)))
            resp = code.get_msg(code.SYSTEM_ERROR)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
    else:
        resp = code.get_msg(code.PERMISSION_DENIED)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


# Get Deleted Business
def api_business_list_del(request):
    if request.session['login_type'] in [2, 6, 3, 7]:
        resp = auth_check(request, "GET")
        if resp != {}:
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        try:
            userIDInfo = request.user.id
            projectAvailableList = []
            if request.session['login_type'] == 2:
                projectAvailableList = getProjectIDByGroupManager(userIDInfo)
            elif request.session['login_type'] == 3:
                projectAvailableList = getProjectIDByCompanyManager(userIDInfo)
            elif request.session['login_type'] == 6:
                projectAvailableList = getProjectIDByGroupManagerAssistant(userIDInfo)
            elif request.session['login_type'] == 7:
                projectAvailableList = getProjectIDByCompanyManagerAssistant(userIDInfo)
            search = request.GET.get("search", None)  # 关键字
            page = int(request.GET.get("page", 1))  # 页码
            size = int(request.GET.get("size", const.ROW_SIZE))  # 页面条数

            qs = Business.objects.filter(Q(project_id__in=projectAvailableList) & Q(del_flag=1))

            if search:
                qs = qs.filter(Q(name__icontains=search) | Q(pk__icontains=search))

            paginator = Paginator(qs, size)

            try:
                business = paginator.page(page)
            except EmptyPage:
                business = paginator.page(1)

            results = []

            for item in business:
                teamMembers = list(
                    BusinessTeamMember.objects.filter(business_id=item.id, project_id=item.cur_project_id).values_list(
                        'user__name', flat=True))
                project = Project.objects.get(pk=item.project_id)
                project_name = project.name
                workflow_name = Flow.objects.get(pk=project.flow_id).name

                user_roles = []
                bus = {
                    'id': item.id, 'name': item.name, 'project_name': project_name,
                    'workflow_name': workflow_name, 'officeItem': item.officeItem.name,
                    'start_time': item.create_time.strftime('%Y-%m-%d %H:%M:%S') if item.create_time else None,
                    'end_time': item.finish_time.strftime('%Y-%m-%d %H:%M:%S') if item.finish_time else None,
                    'members': teamMembers, 'created_by': item.created_by.name,
                    'status': '已完成' if item.status == 9 else '未完成'
                }
                results.append(bus)
            # 分页信息
            paging = {
                'count': paginator.count,
                'has_previous': business.has_previous(),
                'has_next': business.has_next(),
                'num_pages': paginator.num_pages,
                'cur_page': business.number,
                'page_size': size
            }

            resp = code.get_msg(code.SUCCESS)
            resp['d'] = {'results': results, 'paging': paging}
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        except Exception as e:
            logger.exception('api_experiment_list Exception:{0}'.format(str(e)))
            resp = code.get_msg(code.SYSTEM_ERROR)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
    else:
        resp = code.get_msg(code.PERMISSION_DENIED)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


# Delete Business
def api_business_delete(request):
    if request.session['login_type'] in [2, 6, 3, 7]:
        resp = auth_check(request, "POST")
        if resp != {}:
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        try:
            data = request.POST.get("data")  # 实验id数组

            ids = json.loads(data)
            # 排除已经开始的实验
            # Experiment.objects.exclude(status=2).filter(id__in=ids).update(del_flag=1)
            Business.objects.filter(id__in=ids).update(del_flag=1)

            resp = code.get_msg(code.SUCCESS)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        except Exception as e:
            logger.exception('api_experiment_delete Exception:{0}'.format(str(e)))
            resp = code.get_msg(code.SYSTEM_ERROR)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
    else:
        resp = code.get_msg(code.PERMISSION_DENIED)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


# Recovery Business
def api_business_recovery(request):
    if request.session['login_type'] in [2, 6, 3, 7]:
        resp = auth_check(request, "POST")
        if resp != {}:
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        try:
            data = request.POST.get("data")  # 实验id数组

            ids = json.loads(data)
            # 排除已经开始的实验
            # Experiment.objects.exclude(status=2).filter(id__in=ids).update(del_flag=1)
            Business.objects.filter(id__in=ids).update(del_flag=0)

            resp = code.get_msg(code.SUCCESS)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        except Exception as e:
            logger.exception('api_experiment_delete Exception:{0}'.format(str(e)))
            resp = code.get_msg(code.SYSTEM_ERROR)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
    else:
        resp = code.get_msg(code.PERMISSION_DENIED)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


# get unit user list for more business teammates configuration
def get_unit_userList(request):
    if request.session['login_type'] in [3, 5]:
        resp = auth_check(request, "POST")
        if resp != {}:
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        try:
            business_id = request.POST.get("business_id")
            business = Business.objects.get(id=business_id)

            targetUnitUsers = []
            if business.target_part is not None:
                targetUnitUsers = [{
                    'id': item.id,
                    'text': item.username,
                } for item in Tuser.objects.filter(tposition__parts_id=business.target_part.id, is_review=1)]
            elif business.target_company is not None:
                targetUnitUsers = [{
                    'id': item.id,
                    'text': item.username,
                } for item in Tuser.objects.filter(tcompany_id=business.target_company.id, is_review=1)]

            resp = code.get_msg(code.SUCCESS)
            resp['d'] = {'results': targetUnitUsers}
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        except Exception as e:
            logger.exception('api_experiment_delete Exception:{0}'.format(str(e)))
            resp = code.get_msg(code.SYSTEM_ERROR)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
    else:
        resp = code.get_msg(code.PERMISSION_DENIED)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


# get unit user list for more business teammates configuration
def add_more_teammates(request):
    if request.session['login_type'] in [3, 5]:
        resp = auth_check(request, "POST")
        if resp != {}:
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        try:

            business_id = request.POST.get("business_id")
            data = eval(request.POST.get("data"))
            for item in data:
                if item['role_name'] == '':
                    item['role_name'] = None
                else:
                    item['role_name'] = unicode(item['role_name'], "utf8")

            if len(TNotifications.objects.filter(type='businessMoreTeammate_' + business_id)) != 0:
                if teammates_configuration(business_id, data) == 'team_configured':
                    TNotifications.objects.filter(type='businessMoreTeammate_' + business_id).delete()

            resp = code.get_msg(code.SUCCESS)
            resp['d'] = {'results': 'success'}
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        except Exception as e:
            logger.exception('api_experiment_delete Exception:{0}'.format(str(e)))
            resp = code.get_msg(code.SYSTEM_ERROR)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
    else:
        resp = code.get_msg(code.PERMISSION_DENIED)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


# 实验环节用户角色状态查询
def api_business_role_status(request):
    resp = auth_check(request, 'GET')
    observable = False
    if resp != {}:
        observable = True
    try:
        business_id = request.GET.get("business_id")  # 实验任务id
        node_id = request.GET.get("node_id")  # 环节id

        bus = Business.objects.filter(pk=business_id, del_flag=0).first()
        if bus is None:
            resp = code.get_msg(code.BUSINESS_NOT_EXIST)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        project = Project.objects.get(pk=bus.cur_project_id)
        path = BusinessTransPath.objects.filter(business_id=business_id).last()
        node = FlowNode.objects.filter(pk=node_id, del_flag=0).first()

        resp = code.get_msg(code.SUCCESS)
        resp['d'] = get_all_roles_status(bus, project, node, path)
    except Exception as e:
        logger.exception('experiment_role_status Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)
    return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


def api_business_message_push(request):
    """
        实验发送消息
    """
    resp = auth_check(request, "POST")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
    try:
        business_id = request.POST.get("business_id", None)
        node_id = request.POST.get("node_id", None)
        role_alloc_id = request.POST.get("role_alloc_id", None)
        type = request.POST.get("type", None)
        msg = request.POST.get("msg", '')
        cmd = request.POST.get("cmd", None)
        param = request.POST.get("param", None)
        file_id = request.POST.get("file_id", None)
        data = request.POST.get('data', None)
        force_txt_mode = request.POST.get('force_txt_mode',
                                          None)  # added by ser -- force txt mode working though there is no image
        logger.info('business_id:%s,node_id:%s,role_id:%s,type:%s,cmd:%s,param:%s,file_id:%s,'
                    'data:%s' % (business_id, node_id, role_alloc_id, type, cmd, param, file_id, data))

        user = request.user
        user_id = user.id
        if not all(v is not None for v in [user_id, business_id, node_id, role_alloc_id]):
            resp = code.get_msg(code.PARAMETER_ERROR)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        bus = Business.objects.filter(pk=business_id, del_flag=0).first()
        if bus is None:
            resp = code.get_msg(code.BUSINESS_NOT_EXIST)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        bra = BusinessRoleAllocation.objects.filter(pk=role_alloc_id, business_id=business_id,
                                                    project_id=bus.cur_project_id).first()
        if bra is None:
            resp = code.get_msg(code.BUSINESS_ROLE_ALLOCATE_ERROR)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
        role = bra.role
        # 三期 组长没有权限也可以执行一些操作
        # 当前环节不存在该角色 除了组长
        if role is None:
            resp = code.get_msg(code.BUSINESS_NODE_ROLE_NOT_EXIST)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        # todo 组长没有当前环节权限操作返回上一步提前结束等， 这里的判断是不是有问题
        # if exp.node_id != int(node_id):
        #     resp = code.get_msg(code.BUSINESS_NODE_ERROR)
        #     return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        path = BusinessTransPath.objects.filter(business_id=business_id).last()
        if path is None:
            resp = code.get_msg(code.BUSINESS_NODE_ERROR)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        # 如果启动了表达管理，验证当前角色发言次数，每申请一次最多发言三次，
        # 如果是结束权限者则可发言
        role_status = BusinessRoleAllocationStatus.objects.filter(
            business_id=business_id,
            business_role_allocation_id=role_alloc_id,
            # path_id=path.pk
        ).first()
        logger.info('cmd:%s,control_status:%s,param:%s,type:%s' % (cmd, path.control_status, param, type))

        # 是否有结束环节的权限
        can_terminate = bra.can_terminate
        cur_project_id = bus.cur_project_id
        # if role_status is None:
        #     resp = code.get_msg(code.BUSINESS_NODE_ERROR)
        #     return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        if path.control_status == 2 and can_terminate is False:
            if type == const.MSG_TYPE_TXT or type == const.MSG_TYPE_AUDIO:
                if role_status.speak_times == 0:
                    resp = code.get_msg(code.MESSAGE_SPEAKER_CONTROL)
                    return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

            if cmd == const.ACTION_DOC_SHOW:
                if role_status.show_status != 1:
                    resp = code.get_msg(code.MESSAGE_SPEAKER_CONTROL)
                    return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

            if cmd == const.ACTION_DOC_SUBMIT:
                if role_status.submit_status != 1:
                    resp = code.get_msg(code.MESSAGE_SPEAKER_CONTROL)
                    return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        # 三期 - 根据上一步骤自动入席 判断是否入席
        bps = BusinessPositionStatus.objects.filter(business_id=bus.id, path_id=path.id,
                                                    business_role_allocation_id=role_alloc_id)
        print path.id
        print bps
        if bps:
            business_position_status = bps.first()
            if business_position_status.sitting_status:  # 已入席
                role_status.sitting_status = const.SITTING_DOWN_STATUS

        name = request.user.name

        project = Project.objects.get(pk=cur_project_id)
        node = FlowNode.objects.filter(pk=bus.node_id, del_flag=0).first()

        # 角色形象
        image = get_role_image(bra.flow_role_alloc_id)
        if image is None and type != const.MSG_TYPE_CMD and force_txt_mode is None:  # modified by ser, if force_txt_mode then allow push
            resp = code.get_msg(code.BUSINESS_ROLE_IMAGE_NOT_EXIST)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        # 角色占位
        pos = get_role_position(bus, project, node, role, role_alloc_id)
        print(model_to_dict(role_status))
        print(pos)
        time = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        opt = None
        if type == const.MSG_TYPE_TXT:
            if node.process.type == 1:
                if pos is None:
                    resp = code.get_msg(code.BUSINESS_ROLE_POSITION_NOT_EXIST)
                    return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
                # 文本， 角色未入席不能说话
                if role_status.sitting_status == const.SITTING_UP_STATUS:
                    resp = code.get_msg(code.MESSAGE_SITTING_UP_CANNOT_SPEAKER)
                    return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

            msg = msg.strip()
            if msg == '' or len(msg) > 30000:
                resp = code.get_msg(code.PARAMETER_ERROR)
                return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
            msg = tools.filter_invalid_str(msg)
            msg_obj = {'type': const.MSG_TYPE_TXT, 'msg': msg}
            ext = {'business_id': business_id, 'username': name,
                   'alloc_id': role_alloc_id, 'role_name': role.name, 'avatar': image['avatar'] if image else '',
                   'cmd': const.ACTION_TXT_SPEAK, 'param': '', 'time': time, 'can_terminate': can_terminate,
                   'code_position': pos['code_position'] if pos else ''}

        elif type == const.MSG_TYPE_AUDIO:
            if pos is None:
                resp = code.get_msg(code.BUSINESS_ROLE_POSITION_NOT_EXIST)
                return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
            # 音频
            if role_status.sitting_status == const.SITTING_UP_STATUS:
                resp = code.get_msg(code.MESSAGE_SITTING_UP_CANNOT_SPEAKER)
                return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

            if file_id is None:
                resp = code.get_msg(code.PARAMETER_ERROR)
                return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

            audio = BusinessMessageFile.objects.filter(pk=file_id).first()
            if audio is None:
                resp = code.get_msg(code.PARAMETER_ERROR)
                return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

            msg_obj = {'type': const.MSG_TYPE_AUDIO, 'url': (const.WEB_HOST + audio.file.url) if audio.file else '',
                       'filename': audio.file.name, 'length': audio.length, 'secret': ''}

            ext = {'business_id': business_id, 'username': name,
                   'alloc_id': role_alloc_id, 'role_name': role.name, 'avatar': image['avatar'],
                   'cmd': '', 'param': '', 'time': time, 'can_terminate': can_terminate,
                   'code_position': pos['code_position'] if pos else ''}

        elif type == const.MSG_TYPE_CMD:
            # 命令
            if cmd is None:
                resp = code.get_msg(code.PARAMETER_ERROR)
                return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

            result, opt = False, {}
            # 判断
            if cmd == const.ACTION_ROLE_BANNED:
                # 表达管理 data = {'control_status': 1}
                data = json.loads(data)
                result, opt = action_role_banned(bus, node_id, path.pk, data['control_status'])
                # clear_cache(bus.pk)
            elif cmd == const.ACTION_ROLE_MEET:
                # 约见
                result, opt = action_role_meet(bus, path.pk, role, role_alloc_id)
                # clear_cache(bus.pk)
            elif cmd == const.ACTION_ROLE_APPLY_SPEAK:
                # 申请发言
                result, opt = True, {'role_id': role_alloc_id, 'role_name': role.name}
            elif cmd == const.ACTION_ROLE_APPLY_SPEAK_OPT:
                # 申请发言操作结果 data = {'msg_id':1,'role_id': 1, 'result': 1}
                data = json.loads(data)
                result, opt = action_role_speak_opt(bus, path.pk, data)
                # clear_cache(bus.pk)
            elif cmd == const.ACTION_DOC_APPLY_SHOW:
                # 申请展示
                result, opt = True, {'role_id': role_alloc_id, 'role_name': role.name}
            elif cmd == const.ACTION_DOC_REFRESH:
                # 刷新文件列表
                result, opt = True, {'role_id': role_alloc_id, 'role_name': role.name}
            elif cmd == const.ACTION_DOC_APPLY_SHOW_OPT:
                # 申请展示操作结果 data = {'msg_id':1,'doc_id': 1, 'result': 1}
                data = json.loads(data)
                result, opt = action_doc_apply_show_opt(bus, node_id, path.pk, data)
                # clear_cache(bus.pk)
            elif cmd == const.ACTION_DOC_SHOW:
                # 展示
                data = json.loads(data)
                result, opt = action_doc_show(data['doc_id'])
            elif cmd == const.ACTION_ROLE_LETOUT:
                # 送出 data [1, 2, 3, ...]
                role_alloc_ids = json.loads(data)
                result, lst = action_role_letout(bus, node, path.pk, role_alloc_ids)
                opt = {'data': lst}
                # clear_cache(bus.pk)
            elif cmd == const.ACTION_ROLE_LETIN:
                # 请入 data [1, 2, 3, ...]
                role_alloc_ids = json.loads(data)
                result, lst = action_role_letin(bus, node_id, path.pk, role_alloc_ids)
                opt = {'data': lst}
                # clear_cache(bus.pk)
            elif cmd == const.ACTION_ROLE_SITDOWN:
                if pos is None:
                    resp = code.get_msg(code.BUSINESS_ROLE_POSITION_NOT_EXIST)
                    return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
                # 坐下
                if role_status.sitting_status == const.SITTING_UP_STATUS:
                    resp = code.get_msg(code.MESSAGE_SITTING_UP_CANNOT_SPEAKER)
                    return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

                result, opt = action_role_sitdown(bus, path.pk, role, pos, role_alloc_id)
                # clear_cache(bus.pk)
            elif cmd == const.ACTION_ROLE_STAND:
                if pos is None:
                    resp = code.get_msg(code.BUSINESS_ROLE_POSITION_NOT_EXIST)
                    return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
                # 起立
                if role_status.sitting_status == const.SITTING_UP_STATUS:
                    resp = code.get_msg(code.MESSAGE_SITTING_UP_CANNOT_SPEAKER)
                    return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

                result, opt = action_role_stand(bus, path.pk, role, pos, role_alloc_id)
                # clear_cache(bus.pk)
            elif cmd == const.ACTION_ROLE_HIDE:
                if pos is None:
                    resp = code.get_msg(code.BUSINESS_ROLE_POSITION_NOT_EXIST)
                    return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
                # 退席
                if role_status.sitting_status == const.SITTING_UP_STATUS:
                    resp = code.get_msg(code.MESSAGE_SITTING_UP_CANNOT_SPEAKER)
                    return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

                result, opt = action_role_hide(bus, path.pk, role, pos, role_alloc_id)
                # clear_cache(bus.pk)
            elif cmd == const.ACTION_ROLE_SHOW:
                if pos is None:
                    resp = code.get_msg(code.BUSINESS_ROLE_POSITION_NOT_EXIST)
                    return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
                print (role_status.sitting_status, model_to_dict(role_status))
                # 入席
                if role_status.sitting_status == const.SITTING_DOWN_STATUS:
                    resp = code.get_msg(code.BUSINESS_ROLE_HAS_IN_POSITION)
                    return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

                # 占位状态
                position_status = BusinessPositionStatus.objects.filter(
                    business_id=bus.id,
                    business_role_allocation_id=role_alloc_id,
                    path_id=path.pk,
                    position_id=pos['position_id'],
                    sitting_status=const.SITTING_DOWN_STATUS
                ).exists()
                if position_status:
                    resp = code.get_msg(code.BUSINESS_POSITION_HAS_USE)
                    return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

                result, opt = action_role_show(bus, node_id, path.pk, role, pos, role_alloc_id)
                # clear_cache(bus.pk)
            elif cmd == const.ACTION_DOC_APPLY_SUBMIT:
                # 申请提交
                result, opt = True, {'role_id': role_alloc_id, 'role_name': role.name}
            elif cmd == const.ACTION_DOC_APPLY_SUBMIT_OPT:
                # 申请提交操作结果 data = {'msg_id':1,'role_id': 1, 'result': 1}
                data = json.loads(data)
                result, opt = action_doc_apply_submit_opt(bus, node_id, path.pk, data)
                # clear_cache(bus.pk)
            elif cmd == const.ACTION_DOC_SUBMIT:
                # 提交 实验文件id data = [1, 2, 3, ...]
                doc_ids = json.loads(data)
                result, lst = action_doc_submit(doc_ids)
                opt = {'data': lst}
            elif cmd == const.ACTION_EXP_RESTART:
                # 重新开始实验
                # result, opt = action_exp_restart(bus, request.user.pk)
                if not result:
                    return HttpResponse(json.dumps(opt, ensure_ascii=False), content_type="application/json")
                    # clear_cache(bus.pk)
            elif cmd == const.ACTION_EXP_BACK:
                # 返回上一步
                # result, opt = action_exp_back(bus)
                if not result:
                    return HttpResponse(json.dumps(opt, ensure_ascii=False), content_type="application/json")
                    # clear_cache(bus.pk)
            elif cmd == const.ACTION_EXP_NODE_END:
                # 结束环节 opt = {'next_node_id': 1, 'status': 1, 'process_type': 1},
                # data={'tran_id': 1, 'project_id': 0}
                data = json.loads(data)
                result, opt = action_exp_node_end(bus, role_alloc_id, data)
                if not result:
                    return HttpResponse(json.dumps(opt, ensure_ascii=False), content_type="application/json")
                    # clear_cache(bus.pk)
            elif cmd == const.ACTION_EXP_FINISH:
                result, opt = action_exp_finish(bus, request.user.id)
                if not result:
                    return HttpResponse(json.dumps(opt, ensure_ascii=False), content_type="application/json")
                    # clear_cache(bus.pk)
            elif cmd == const.ACTION_SUBMIT_EXPERIENCE:
                # 提交实验心得 data = {"content": ""}
                data = json.loads(data)
                result, opt = action_submit_experience(bus, data['content'], user_id)
                if not result:
                    return HttpResponse(json.dumps(opt, ensure_ascii=False), content_type="application/json")
            elif cmd == const.ACTION_ROLE_VOTE:
                # 提交实验投票 data = {"status": 1}
                data = json.loads(data)
                result, opt = action_role_vote(bus, node_id, path, role_alloc_id, data['status'])
                if not result:
                    return HttpResponse(json.dumps(opt, ensure_ascii=False), content_type="application/json")
                    # clear_cache(bus.pk)
            elif cmd == const.ACTION_ROLE_VOTE_END:
                # 提交实验投票结束 data = {}
                result, opt = action_role_vote_end(bus, node_id, path)
                if not result:
                    return HttpResponse(json.dumps(opt, ensure_ascii=False), content_type="application/json")
                    # clear_cache(bus.pk)
            elif cmd == const.ACTION_ROLE_REQUEST_SIGN:
                # 要求签字 data = {"doc_id": 1, "doc_name": "xxx", "role_id": 1, "role_name": "xx"}
                data = json.loads(data)
                result, opt = action_role_request_sign(bus, node_id, data)
                if not result:
                    return HttpResponse(json.dumps(opt, ensure_ascii=False), content_type="application/json")
            elif cmd == const.ACTION_ROLE_SIGN:
                # 签字 data = {'msg_id':1,'result': 1,"doc_id": 1, "doc_name": 'xxx'}
                data = json.loads(data)
                sign = '{0}({1})'.format(request.user.name, role.name)
                result, opt = action_role_sign(bus, sign, node_id, role_alloc_id, data)
                if not result:
                    return HttpResponse(json.dumps(opt, ensure_ascii=False), content_type="application/json")
                    # clear_cache(bus.pk)
            elif cmd == const.ACTION_ROLE_SCHEDULE_REPORT:
                # 安排报告 data = {"role_id": 1, "role_name": "xx"}
                data = json.loads(data)
                result, opt = action_role_schedule_report(bus, node_id, path.pk, data)
                if not result:
                    return HttpResponse(json.dumps(opt, ensure_ascii=False), content_type="application/json")

            elif cmd == const.ACTION_ROLE_TOWARD_REPORT:
                # 走向发言席 data = {}
                if pos is None:
                    resp = code.get_msg(code.BUSINESS_ROLE_POSITION_NOT_EXIST)
                    return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

                result, opt = action_role_toward_report(bus, node_id, path.pk, bra, pos)
                if not result:
                    return HttpResponse(json.dumps(opt, ensure_ascii=False), content_type="application/json")
                    # clear_cache(bus.pk)
            elif cmd == const.ACTION_ROLE_EDN_REPORT:
                # 走下发言席 data = {}
                if pos is None:
                    resp = code.get_msg(code.BUSINESS_ROLE_POSITION_NOT_EXIST)
                    return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

                result, opt = action_role_end_report(bus, node_id, path.pk, bra, pos)
                if not result:
                    return HttpResponse(json.dumps(opt, ensure_ascii=False), content_type="application/json")
                    # clear_cache(bus.pk)
            elif cmd == const.ACTION_ROLES_EXIT:
                # 退出实验 data = {}
                result, lst = action_roles_exit(bus, node, path.pk, user_id)
                opt = {'data': lst}
                # clear_cache(bus.pk)
            elif cmd == const.ACTION_TRANS:
                result, opt = True, {}
            else:
                logger.info('action cmd %s' % cmd)
                resp = code.get_msg(code.MESSAGE_ACTION_ERROR)
                return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

            # if not result:
            #     raise Exception(opt)

            msg_obj = {'type': const.MSG_TYPE_TXT, 'msg': msg}
            ext = {'business_id': business_id, 'node_id': node_id, 'username': name,
                   'role_alloc_id': role_alloc_id, 'role_name': role.name, 'avatar': image['avatar'] if image else '',
                   'cmd': cmd, 'param': param, 'time': time, 'opt': opt, 'can_terminate': can_terminate,
                   'code_position': pos['code_position'] if pos else ''}
        else:
            resp = code.get_msg(code.PARAMETER_ERROR)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        # 保存消息，得到消息id
        user = request.user
        message = BusinessMessage()
        if role_alloc_id:
            bra = BusinessRoleAllocation.objects.filter(pk=role_alloc_id,
                                                        business_id=business_id,
                                                        project_id=cur_project_id
                                                        ).first()
            if bra:
                message = BusinessMessage.objects.create(business_id=business_id, user_id=user.pk,
                                                         business_role_allocation_id=role_alloc_id,
                                                         file_id=file_id, msg=msg, msg_type=type,
                                                         path_id=path.id, user_name=user.name, role_name=bra.role.name,
                                                         ext=json.dumps(ext))
        ext['id'] = message.pk
        ext['opt_status'] = False

        msgDict = model_to_dict(message) if message else {}
        msgDict['ext'] = ext
        msgDict['from'] = message.user_id if message else None
        msgDict['type'] = 'groupchat'
        msgDict['to'] = None
        with SocketIO(u'localhost', 4000, LoggingNamespace) as socketIO:
            socketIO.emit('message', msgDict)
            socketIO.wait_for_callbacks(seconds=1)
        resp = code.get_msg(code.SUCCESS)
        if opt:
            resp['d'] = opt

        if can_terminate is False:
            # 角色发言次数减1
            if path.control_status == 2 and type != const.MSG_TYPE_CMD:
                brat = BusinessRoleAllocationStatus.objects.filter(
                    business_role_allocation_id=role_alloc_id,
                    business_id=business_id,
                    # path_id=path.pk
                ).first()
                brat.speak_times -= 1
                brat.save(update_fields=['speak_times'])
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
    except Exception as e:
        logger.exception('api_business_message_push Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


# 待请入角色列表
def api_business_role_in_list(request):
    resp = auth_check(request, "GET")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    try:
        business_id = request.GET.get('business_id', None)  # 实验id
        node_id = request.GET.get("node_id", None)  # 环节id
        role_alloc_id = request.GET.get("role_alloc_id", 0)  # 角色id

        bus = Business.objects.filter(pk=business_id).first()

        # 判断实验是否存在以及实验当前环节是否是node_id
        if bus and bus.node_id == int(node_id):
            project = Project.objects.filter(pk=bus.cur_project_id).first()
            path = BusinessTransPath.objects.filter(business_id=business_id).last()

            role_alloc_ids = BusinessRoleAllocationStatus.objects.filter(
                business_id=business_id,
                business_role_allocation__node_id=bus.node_id,
                # path_id=path.pk,
                sitting_status=1
            ).exclude(come_status=9).values_list('business_role_allocation_id', flat=True)
            role_alloc_list = []
            for id in role_alloc_ids:
                if id == int(role_alloc_id):
                    continue

                bra = BusinessRoleAllocation.objects.get(pk=id)
                br = BusinessRole.objects.get(pk=bra.role_id)
                role_position = FlowRolePosition.objects.filter(flow_id=project.flow_id, node_id=node_id, no=bra.no,
                                                                role_id=br.flow_role_id, del_flag=0).first()
                if role_position:
                    pos_status = BusinessPositionStatus.objects.filter(
                        business_id=business_id,
                        business_role_allocation_id=id,
                        path_id=path.pk,
                        position_id=role_position.position_id
                    ).first()

                    if pos_status and pos_status.sitting_status == const.SITTING_DOWN_STATUS:
                        continue

                    pos = FlowPosition.objects.filter(pk=role_position.position_id).first()
                    if pos:
                        code_position = pos.code_position
                    else:
                        continue
                else:
                    continue
                role_alloc_list.append({
                    'id': bra.id, 'name': br.name, 'code_position': code_position
                })
            resp = code.get_msg(code.SUCCESS)
            resp['d'] = role_alloc_list
        elif bus is None:
            resp = code.get_msg(code.EXPERIMENT_NOT_EXIST)
        else:
            resp = code.get_msg(code.EXPERIMENT_NODE_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    except Exception as e:
        logger.exception('api_experiment_role_in_list Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


# 待请出角色列表
def api_business_role_out_list(request):
    resp = auth_check(request, "GET")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    try:
        business_id = request.GET.get('business_id', None)  # 实验id
        node_id = request.GET.get("node_id", None)  # 环节id
        role_alloc_id = request.GET.get("role_alloc_id", 0)  # 角色id

        bus = Business.objects.filter(pk=business_id).first()
        if bus and bus.node_id == int(node_id):
            project = Project.objects.filter(pk=bus.cur_project_id).first()
            path = BusinessTransPath.objects.filter(business_id=business_id).last()

            role_alloc_ids = BusinessRoleAllocationStatus.objects.filter(
                business_id=business_id,
                business_role_allocation__node_id=bus.node_id,
                # path_id=path.pk,
                sitting_status=2
            ).exclude(come_status=9).values_list('business_role_allocation_id', flat=True)

            role_alloc_list = []
            for id in role_alloc_ids:
                if id == int(role_alloc_id):
                    continue

                bra = BusinessRoleAllocation.objects.get(pk=id)
                br = BusinessRole.objects.get(pk=bra.role_id)
                role_position = FlowRolePosition.objects.filter(flow_id=project.flow_id, node_id=node_id, no=bra.no,
                                                                role_id=br.flow_role_id, del_flag=0).first()

                if role_position:
                    pos = FlowPosition.objects.filter(pk=role_position.position_id).first()
                    if pos:
                        code_position = pos.code_position
                    else:
                        continue
                else:
                    continue
                role_alloc_list.append({
                    'id': bra.id, 'name': br.name, 'code_position': code_position
                })
            resp = code.get_msg(code.SUCCESS)
            resp['d'] = role_alloc_list
        elif bus is None:
            resp = code.get_msg(code.EXPERIMENT_NOT_EXIST)
        else:
            resp = code.get_msg(code.EXPERIMENT_NODE_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    except Exception as e:
        logger.exception('api_experiment_role_out_list Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


# 实验中上传文件
def api_business_docs_create(request):
    resp = auth_check(request, "POST")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    try:
        upload_file = request.FILES["file"]  # 文件
        business_id = request.POST.get("business_id")  # 实验
        node_id = request.POST.get("node_id")  # 环节
        role_alloc_id = request.POST.get("role_alloc_id", None)  # 角色id
        cmd = request.POST.get('cmd', None)
        logger.info('business_id:%s,node_id:%s,role_id:%s,cmd:%s' % (business_id, node_id, role_alloc_id, cmd))
        path = BusinessTransPath.objects.filter(business_id=business_id).last()
        if path is None:
            resp = code.get_msg(code.EXPERIMENT_NODE_ERROR)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        if len(upload_file.name) > 60:
            resp = code.get_msg(code.UPLOAD_FILE_NAME_TOOLONG_ERROR)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        file_type = tools.check_file_type(upload_file.name)
        doc = BusinessDoc.objects.create(
            filename=upload_file.name,
            file=upload_file,
            business_id=business_id,
            node_id=node_id,
            business_role_allocation_id=role_alloc_id,
            path_id=path.id,
            file_type=file_type,
            created_by=request.user
        )

        resp = code.get_msg(code.SUCCESS)
        resp['d'] = {
            'id': doc.id, 'filename': doc.filename, 'file': doc.file.url if doc.file else None,
            'business_id': doc.business_id, 'node_id': doc.node_id, 'file_type': file_type,
            'created_by': user_simple_info(doc.created_by_id)
        }
        # clear_cache(business_id)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
    except Exception as e:
        logger.exception('api_experiment_docs_create Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


# 实验文件展示列表
def api_business_file_display_list(request):
    resp = auth_check(request, "GET")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    try:
        business_id = request.GET.get('business_id', None)  # 实验id
        node_id = request.GET.get('node_id', None)
        path_id = request.GET.get("path_id", None)  # 环节id

        bus = Business.objects.filter(pk=business_id).first()
        if bus:
            doc_list = get_business_display_files(bus, node_id, path_id)
            # 分页信息
            paging = {
                'count': len(doc_list),
                'has_previous': False,
                'has_next': False,
                'num_pages': 1,
                'cur_page': 1,
            }

            resp = code.get_msg(code.SUCCESS)
            resp['d'] = {'results': doc_list, 'paging': paging}

        else:
            resp = code.get_msg(code.EXPERIMENT_NOT_EXIST)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    except Exception as e:
        logger.exception('api_experiment_file_display_list Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


# 待安排报告角色列表
def api_business_role_schedule_report_list(request):
    resp = auth_check(request, "GET")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    try:
        business_id = request.GET.get('business_id', None)  # 实验id
        node_id = request.GET.get("node_id")  # 环节id
        bus = Business.objects.filter(pk=business_id).first()

        # 判断实验是否存在以及实验当前环节是否是node_id
        if bus and bus.node_id == int(node_id):
            path = BusinessTransPath.objects.filter(business_id=business_id).last()

            role_alloc_ids = BusinessReportStatus.objects.filter(
                business_id=business_id,
                business_role_allocation__node_id=bus.node_id,
                path_id=path.pk
            ).exclude(schedule_status=0).values_list('business_role_allocation_id', flat=True)

            role_alloc_status = BusinessRoleAllocationStatus.objects.filter(
                business_id=business_id,
                business_role_allocation__node_id=node_id,
                # path_id=path.pk
            )
            role_list = []
            for item in role_alloc_status:
                if item.business_role_allocation_id not in role_alloc_ids:
                    role = BusinessRoleAllocation.objects.filter(pk=item.business_role_allocation_id).first().role
                    if role.name != const.ROLE_TYPE_OBSERVER:
                        role_list.append({'id': item.business_role_allocation_id, 'name': role.name})

            resp = code.get_msg(code.SUCCESS)
            resp['d'] = role_list

        elif bus is None:
            resp = code.get_msg(code.EXPERIMENT_NOT_EXIST)
        else:
            resp = code.get_msg(code.EXPERIMENT_NODE_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    except Exception as e:
        logger.exception('api_experiment_wait_report_list Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


# 实验环节用户可签字角色查询
def api_business_request_sign_roles(request):
    resp = auth_check(request, "GET")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    try:
        business_id = request.GET.get('business_id', None)  # 实验id
        node_id = request.GET.get("node_id", None)  # 环节id
        role_alloc_id = request.GET.get("role_alloc_id", 0)  # 角色id
        bus = Business.objects.filter(pk=business_id).first()

        # 判断实验是否存在以及实验当前环节是否是node_id
        if bus and bus.node_id == int(node_id):
            path = BusinessTransPath.objects.filter(business_id=business_id).last()

            role_alloc_status_list = BusinessRoleAllocationStatus.objects.filter(
                business_id=business_id,
                business_role_allocation__node_id=bus.node_id,
                # path_id=path.pk,
                sitting_status=2)
            role_list = []
            for item in role_alloc_status_list:
                if item.business_role_allocation_id == int(role_alloc_id):
                    continue

                role = BusinessRoleAllocation.objects.get(pk=item.business_role_allocation_id).role
                # 三期 老师以实验指导登录进来，老师只观察只给一个观察者的角色;
                # 老师以实验者登录进来，要去掉老师的观察者角色
                if role.name != const.ROLE_TYPE_OBSERVER:
                    role_list.append({'id': item.business_role_allocation_id, 'name': role.name})
            resp = code.get_msg(code.SUCCESS)
            resp['d'] = role_list
        elif bus is None:
            resp = code.get_msg(code.EXPERIMENT_NOT_EXIST)
        else:
            logger.info('=====================2422======================')
            resp = code.get_msg(code.EXPERIMENT_NODE_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    except Exception as e:
        logger.exception('api_experiment_request_sign_roles Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


def api_business_report_generate(request):
    resp = auth_check(request, "GET")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    try:
        business_id = request.GET.get("business_id")  # 实验ID
        user_id = request.GET.get("user_id", None)  # 用户
        user_id = user_id if user_id else request.user.id
        busi = Business.objects.filter(pk=business_id).first()
        if busi:
            project = Project.objects.filter(pk=busi.project_id).first()
            flow = Flow.objects.filter(pk=project.flow_id).first()
            members = BusinessTeamMember.objects.filter(business_id=business_id, del_flag=0,
                                                        project_id=busi.cur_project_id).values_list('user_id',
                                                                                                    flat=True)

            # 小组成员
            member_list = []

            for uid in members:
                if uid:
                    user = Tuser.objects.get(pk=int(uid))
                    member_list.append(user.name)

            # 各环节提交文件信息和聊天信息
            paths = BusinessTransPath.objects.filter(business_id=busi.id)
            node_list = []

            for item in paths:
                node = FlowNode.objects.filter(pk=item.node_id, del_flag=0).first()
                if node.process:
                    if node.process.type == const.PROCESS_NEST_TYPE:
                        continue
                    doc_list = []
                    vote_status = []
                    if node.process.type == 2:
                        # 如果是编辑
                        # 应用模板
                        contents = BusinessDocContent.objects.filter(business_id=business_id, node_id=item.node_id,
                                                                     has_edited=True)
                        for d in contents:
                            doc_list.append({
                                'id': d.doc_id, 'filename': d.name, 'content': d.content, 'file_type': d.file_type,
                                'signs': [{'sign_status': d.sign_status, 'sign': d.sign}],
                                'url': d.file.url if d.file else None
                            })
                        # 提交的文件
                        docs = BusinessDoc.objects.filter(business_id=business_id, node_id=item.node_id,
                                                          path_id=item.pk)
                        for d in docs:
                            sign_list = BusinessDocSign.objects.filter(doc_id=d.pk).values('sign', 'sign_status')
                            doc_list.append({
                                'id': d.id, 'filename': d.filename, 'content': d.content, 'file_type': d.file_type,
                                'signs': list(sign_list), 'url': d.file.url if d.file else None
                            })
                    elif node.process.type == 3:
                        project_docs = BusinessDoc.objects.filter(business_id=business_id, node_id=item.node_id,
                                                                  path_id=item.pk)
                        for d in project_docs:
                            doc_list.append({
                                'id': d.id, 'filename': d.filename, 'signs': [],
                                'url': d.file.url if d.file else None, 'content': d.content, 'file_type': d.file_type,
                            })
                    elif node.process.type == 5:
                        # 如果是投票   三期 - 增加投票结果数量汇总
                        vote_status_0_temp = BusinessRoleAllocationStatus.objects.filter(
                            business_id=business_id,
                            business_role_allocation__node_id=item.node_id,
                            # path_id=item.id,
                            vote_status=0)
                        vote_status_0 = []
                        # 去掉老师观察者角色的数据
                        for item0 in vote_status_0_temp:
                            role_alloc_temp = item0.business_role_allocation
                            if role_alloc_temp.role.name != const.ROLE_TYPE_OBSERVER:
                                vote_status_0.append(item0)

                        vote_status_1_temp = BusinessRoleAllocationStatus.objects.filter(
                            business_id=business_id,
                            business_role_allocation__node_id=item.node_id,
                            # path_id=item.id,
                            vote_status=1)
                        vote_status_1 = []
                        # 去掉老师观察者角色的数据
                        for item1 in vote_status_1_temp:
                            role_alloc_temp = item1.business_role_allocation
                            if role_alloc_temp.name != const.ROLE_TYPE_OBSERVER:
                                vote_status_1.append(item1)

                        vote_status_2_temp = BusinessRoleAllocationStatus.objects.filter(
                            business_id=business_id,
                            business_role_allocation__node_id=item.node_id,
                            # path_id=item.id,
                            vote_status=2)
                        vote_status_2 = []
                        # 去掉老师观察者角色的数据
                        for item2 in vote_status_2_temp:
                            role_alloc_temp = item2.business_role_allocation
                            if role_alloc_temp.name != const.ROLE_TYPE_OBSERVER:
                                vote_status_2.append(item2)

                        vote_status_9_temp = BusinessRoleAllocationStatus.objects.filter(
                            business_id=business_id,
                            business_role_allocation__node_id=item.node_id,
                            # path_id=item.id,
                            vote_status=9)
                        vote_status_9 = []
                        # 去掉老师观察者角色的数据
                        for item9 in vote_status_9_temp:
                            role_alloc_temp = item9.business_role_allocation
                            if role_alloc_temp.name != const.ROLE_TYPE_OBSERVER:
                                vote_status_9.append(item9)
                        vote_status = [{'status': '同意', 'num': len(vote_status_1)},
                                       {'status': '不同意', 'num': len(vote_status_2)},
                                       {'status': '弃权', 'num': len(vote_status_9)},
                                       {'status': '未投票', 'num': len(vote_status_0)}]
                        pass
                    else:
                        # 提交的文件
                        docs = BusinessDoc.objects.filter(business_id=business_id, node_id=item.node_id,
                                                          path_id=item.id)
                        for d in docs:
                            sign_list = BusinessDocSign.objects.filter(doc_id=d.pk).values('sign', 'sign_status')
                            doc_list.append({
                                'id': d.id, 'filename': d.filename, 'content': d.content,
                                'signs': list(sign_list), 'url': d.file.url if d.file else None, 'file_type': d.file_type
                            })

                    # 消息
                    messages = BusinessMessage.objects.filter(business_id=business_id,
                                                              business_role_allocation__node_id=item.node_id,
                                                              path_id=item.id).order_by('timestamp')
                    message_list = []
                    for m in messages:
                        message = {
                            'user_name': m.user_name, 'role_name': m.role_name,
                            'msg': m.msg, 'msg_type': m.msg_type, 'ext': json.loads(m.ext),
                            'timestamp': m.timestamp.strftime('%Y-%m-%d %H:%M:%S')
                        }
                        message_list.append(message)

                    # 个人笔记
                    note = BusinessNotes.objects.filter(business_id=business_id,
                                                        node_id=item.node_id, created_by_id=user_id).first()
                    node_list.append({
                        'docs': doc_list, 'messages': message_list, 'id': node.id, 'node_name': node.name,
                        'note': note.content if note else None, 'type': node.process.type if node.process else 0,
                        'vote_status': vote_status
                    })

            experience = BusinessExperience.objects.filter(business_id=busi.id, created_by_id=user_id).first()
            experience_data = {'status': 1, 'content': ''}
            if experience:
                experience_data = {
                    'id': experience.id, 'content': experience.content, 'status': experience.status,
                    'created_by': user_simple_info(user_id),
                    'create_time': experience.create_time.strftime('%Y-%m-%d')
                }

            resp = code.get_msg(code.SUCCESS)
            resp['d'] = {
                'name': u'{0} {1}'.format(busi.id, busi.name), 'project_name': project.name,
                'flow_name': flow.name, 'members': member_list,
                'finish_time': busi.finish_time.strftime('%Y-%m-%d') if busi.finish_time else None,
                'start_time': busi.start_time.strftime('%Y-%m-%d') if busi.start_time else None,
                'end_time': busi.end_time.strftime('%Y-%m-%d') if busi.end_time else None,
                'create_time': busi.create_time.strftime('%Y-%m-%d'),
                'nodes': node_list, 'experience': experience_data,
            }
        else:
            resp = code.get_msg(code.BUSINESS_NOT_EXIST)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    except Exception as e:
        logger.exception('api_business_report_generate Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


# 设置样式
def set_style(height, bold=False):
    style = xlwt.XFStyle()  # 初始化样式
    pattern = xlwt.Pattern()
    pattern.pattern = xlwt.Pattern.SOLID_PATTERN
    pattern.pattern_fore_colour = xlwt.Style.colour_map['aqua']
    font = xlwt.Font()  # 为样式创建字体
    font.bold = bold
    font.color_index = 4
    font.height = height
    style.font = font
    style.pattern = pattern
    return style


def api_business_report_export(request):
    resp = auth_check(request, "GET")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    try:
        business_id = request.GET.get("business_id")  # 实验ID
        user_id = request.GET.get("user_id", None)  # 用户
        user_id = user_id if user_id else request.user.id
        busi = Business.objects.filter(pk=business_id).first()

        docTitle = [u'文件名', u'文件类型', u'签字', u'url']
        messageTitle = [u'user_name', u'role_name', u'time', u'msg']
        voteTitle = [u'同意', u'不同意', u'弃权', u'未投票']
        noteTitle = u'Note'
        experienceTitle = [u'user_name', u'time', u'content']

        if busi:
            project = Project.objects.filter(pk=busi.project_id).first()
            members = BusinessTeamMember.objects.filter(business_id=business_id, del_flag=0, project_id=busi.cur_project_id).values_list('user_id', flat=True)

            # 小组成员
            member_list = []

            for uid in members:
                if uid is not None:
                    user = Tuser.objects.get(pk=int(uid))
                    member_list.append(user.name)

            # 各环节提交文件信息和聊天信息
            paths = BusinessTransPath.objects.filter(business_id=busi.id)
            report = xlwt.Workbook(encoding='utf8')
            for item in paths:
                node = FlowNode.objects.filter(pk=item.node_id, del_flag=0).first()
                if node.process.type == const.PROCESS_NEST_TYPE:
                    continue
                doc_list = []
                vote_status = []
                sheet = report.add_sheet(node.name)

                if node.process.type == 2:
                    # 如果是编辑
                    # 应用模板
                    contents = BusinessDocContent.objects.filter(business_id=business_id, node_id=item.node_id,
                                                                 has_edited=True)
                    for d in contents:
                        doc_list.append({
                            'id': d.doc_id, 'filename': d.name, 'content': d.content, 'file_type': d.file_type,
                            'signs': [{'sign_status': d.sign_status, 'sign': d.sign}],
                            'url': d.file.url if d.file else None
                        })
                    # 提交的文件
                    docs = BusinessDoc.objects.filter(business_id=business_id, node_id=item.node_id,
                                                      path_id=item.pk)
                    for d in docs:
                        sign_list = BusinessDocSign.objects.filter(doc_id=d.pk).values('sign', 'sign_status')
                        doc_list.append({
                            'id': d.id, 'filename': d.filename, 'content': d.content, 'file_type': d.file_type,
                            'signs': list(sign_list), 'url': d.file.url if d.file else None
                        })
                elif node.process.type == 3:
                    project_docs = BusinessDoc.objects.filter(business_id=business_id, node_id=item.node_id,
                                                              path_id=item.pk)
                    for d in project_docs:
                        doc_list.append({
                            'id': d.id, 'filename': d.filename, 'signs': [],
                            'url': d.file.url if d.file else None, 'content': d.content, 'file_type': d.file_type,
                        })
                elif node.process.type == 5:
                    # 如果是投票   三期 - 增加投票结果数量汇总
                    vote_status_0_temp = BusinessRoleAllocationStatus.objects.filter(
                        business_id=business_id,
                        business_role_allocation__node_id=item.node_id,
                        # path_id=item.id,
                        vote_status=0)
                    vote_status_0 = []
                    # 去掉老师观察者角色的数据
                    for item0 in vote_status_0_temp:
                        role_alloc_temp = item0.business_role_allocation
                        if role_alloc_temp.role.name != const.ROLE_TYPE_OBSERVER:
                            vote_status_0.append(item0)

                    vote_status_1_temp = BusinessRoleAllocationStatus.objects.filter(
                        business_id=business_id,
                        business_role_allocation__node_id=item.node_id,
                        # path_id=item.id,
                        vote_status=1)
                    vote_status_1 = []
                    # 去掉老师观察者角色的数据
                    for item1 in vote_status_1_temp:
                        role_alloc_temp = item1.business_role_allocation
                        if role_alloc_temp.name != const.ROLE_TYPE_OBSERVER:
                            vote_status_1.append(item1)

                    vote_status_2_temp = BusinessRoleAllocationStatus.objects.filter(
                        business_id=business_id,
                        business_role_allocation__node_id=item.node_id,
                        # path_id=item.id,
                        vote_status=2)
                    vote_status_2 = []
                    # 去掉老师观察者角色的数据
                    for item2 in vote_status_2_temp:
                        role_alloc_temp = item2.business_role_allocation
                        if role_alloc_temp.name != const.ROLE_TYPE_OBSERVER:
                            vote_status_2.append(item2)

                    vote_status_9_temp = BusinessRoleAllocationStatus.objects.filter(
                        business_id=business_id,
                        business_role_allocation__node_id=item.node_id,
                        # path_id=item.id,
                        vote_status=9)
                    vote_status_9 = []
                    # 去掉老师观察者角色的数据
                    for item9 in vote_status_9_temp:
                        role_alloc_temp = item9.business_role_allocation
                        if role_alloc_temp.name != const.ROLE_TYPE_OBSERVER:
                            vote_status_9.append(item9)
                    vote_status = [{'status': '同意', 'num': len(vote_status_1)},
                                   {'status': '不同意', 'num': len(vote_status_2)},
                                   {'status': '弃权', 'num': len(vote_status_9)},
                                   {'status': '未投票', 'num': len(vote_status_0)}]
                    pass
                else:
                    # 提交的文件
                    docs = BusinessDoc.objects.filter(business_id=business_id, node_id=item.node_id,
                                                      path_id=item.id)
                    for d in docs:
                        sign_list = BusinessDocSign.objects.filter(doc_id=d.pk).values('sign', 'sign_status')
                        doc_list.append({
                            'id': d.id, 'filename': d.filename, 'content': d.content,
                            'signs': list(sign_list), 'url': d.file.url if d.file else None, 'file_type': d.file_type
                        })

                # 消息
                messages = BusinessMessage.objects.filter(business_id=business_id,
                                                          business_role_allocation__node_id=item.node_id,
                                                          path_id=item.id).order_by('timestamp')
                message_list = []
                for m in messages:
                    message = {
                        'user_name': m.user_name, 'role_name': m.role_name,
                        'msg': m.msg, 'msg_type': m.msg_type, 'ext': json.loads(m.ext),
                        'timestamp': m.timestamp.strftime('%Y-%m-%d %H:%M:%S')
                    }
                    message_list.append(message)

                # 个人笔记
                note = BusinessNotes.objects.filter(business_id=business_id,
                                                    node_id=item.node_id, created_by_id=user_id).first()

                # 设置样式
                for i in range(0, len(docTitle)):
                    sheet.write(0, i, docTitle[i], set_style(220, True))
                row = 1
                for d in doc_list:
                    sheet.write(row, 0, d['filename'])
                    sheet.write(row, 1, d['file_type'])
                    topRow = row
                    for sign in d['signs']:
                        if int(sign['sign_status']) == 1:
                            sheet.write(topRow, 2, sign['sign'] + u'--已签字')
                            topRow += 1
                        elif int(sign['sign_status']) == 2:
                            sheet.write(topRow, 2, sign['sign'] + u'--已拒绝签字')
                            topRow += 1
                    sheet.write(row, 3, d['url'])
                    row = topRow
                    row += 1

                row += 2
                for i in range(0, len(messageTitle)):
                    sheet.write(row, i, messageTitle[i], set_style(220, True))
                row += 1
                for m in message_list:
                    sheet.write(row, 0, m['user_name'])
                    sheet.write(row, 1, m['role_name'])
                    sheet.write(row, 2, m['timestamp'])
                    if m['ext']['cmd'] == 'action_submit_experience':
                        sheet.write(row, 3, m['ext']['opt']['content'])
                    else:
                        sheet.write(row, 3, m['msg'])
                    row += 1

                if node.process.type == 5:
                    row += 2
                    for i in range(0, len(voteTitle)):
                        sheet.write(row, i, voteTitle[i], set_style(220, True))
                    row += 1
                    for i in range(0, len(vote_status)):
                        sheet.write(row, i, vote_status[i].num)
                    row += 1
                if note:
                    row += 2
                    sheet.write(row, 0, noteTitle, set_style(220, True))
                    row += 1
                    sheet.write(row, 0, note.content, set_style(220, True))

            experiences = BusinessExperience.objects.filter(business_id=busi.id)
            sheet = report.add_sheet(u'Experience')  # 设置样式
            for i in range(0, len(experienceTitle)):
                sheet.write(0, i, experienceTitle[i], set_style(220, True))
            row = 1
            for e in experiences:
                sheet.write(row, 0, e.created_by.name)
                sheet.write(row, 1, e.create_time.strftime('%Y-%m-%d'))
                sheet.write(row, 2, e.content)
                row += 1

            response = HttpResponse(content_type='application/vnd.ms-excel')
            filename = urlquote(u'心得')
            response['Content-Disposition'] = u'attachment;filename=%s.xls' % filename
            report.save(response)
            return response
        else:
            resp = code.get_msg(code.BUSINESS_NOT_EXIST)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    except Exception as e:
        logger.exception('api_business_report_export Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


def api_business_experience_list(request):
    resp = auth_check(request, "GET")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    try:
        business_id = request.GET.get("business_id")  # 实验ID
        # 判断实验是否存在
        business = Business.objects.filter(pk=business_id, del_flag=0)
        if business:
            lst = BusinessExperience.objects.filter(business_id=business_id)
            data = []
            for item in lst:
                data.append({
                    'id': item.id, 'content': item.content, 'created_by': user_simple_info(item.created_by_id),
                    'create_time': item.create_time.strftime('%Y-%m-%d'), 'status': item.status
                })
            resp = code.get_msg(code.SUCCESS)
            resp['d'] = data
        else:
            resp = code.get_msg(code.BUSINESS_NOT_EXIST)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    except Exception as e:
        logger.exception('api_business_experience_list Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


#
def api_business_save_experience(request):
    resp = auth_check(request, "POST")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
    try:
        business_id = request.POST.get('business_id', None)  # 实验id
        content = request.POST.get('content', '')
        if business_id is None:
            resp = code.get_msg(code.PARAMETER_ERROR)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        busi = Business.objects.filter(pk=business_id, del_flag=0).first()
        if busi is None:
            resp = code.get_msg(code.BUSINESS_NOT_EXIST)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        if busi.status == 2:
            if content is None or len(content) > 30000:
                resp = code.get_msg(code.PARAMETER_ERROR)
                return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

            instance, flag = BusinessExperience.objects.update_or_create(business_id=business_id,
                                                                         created_by_id=request.user.id,
                                                                         defaults={'content': content,
                                                                                   'created_by_id': request.user.id})
            data = {
                'id': instance.pk, 'content': instance.content, 'status': instance.status,
                'created_by': user_simple_info(instance.created_by_id),
                'create_time': instance.create_time.strftime('%Y-%m-%d')
            }
            resp = code.get_msg(code.SUCCESS)
            resp['d'] = data
        elif busi.status == 1:
            resp = code.get_msg(code.BUSINESS_HAS_NOT_STARTED)
        else:
            resp = code.get_msg(code.BUSINESS_HAS_FINISHED)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    except Exception as e:
        logger.exception('api_business_save_experience Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


def api_business_post(request):
    resp = auth_check(request, "GET")
    observable = False
    if resp != {}:
        observable = True
    try:
        business_id = request.GET.get('business_id', None)  # 实验id
        node_id = request.GET.get('node_id', None)
        if not all(v is not None for v in [business_id, node_id]):
            resp = code.get_msg(code.PARAMETER_ERROR)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        if (observable and not is_look_on_node(node_id)):
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
        business = Business.objects.filter(pk=business_id, del_flag=0).first()
        if business is None:
            resp = code.get_msg(code.BUSINESS_NOT_EXIST)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        if int(business.node_id) != int(node_id):
            resp = code.get_msg(code.PARAMETER_ERROR)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        if business.status == 2:
            businessPost = BusinessPost.objects.filter(business_id=business_id, node_id=node_id).first()
            resp = code.get_msg(code.SUCCESS)
            resp['d'] = model_to_dict(businessPost) if businessPost else {}
        elif business.status == 1:
            resp = code.get_msg(code.BUSINESS_HAS_NOT_STARTED)
        else:
            resp = code.get_msg(code.BUSINESS_HAS_FINISHED)

        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    except Exception as e:
        logger.exception('api_business_post_info Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


def api_business_post_create(request):
    resp = auth_check(request, "POST")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
    try:
        business_id = request.POST.get('business_id', None)  # 实验id
        node_id = request.POST.get('node_id', None)
        post_name = request.POST.get("post_name", None)  # 名称
        post_content = request.POST.get("post_content", None)  # 名称
        if not all(v is not None for v in [business_id, node_id, post_name, post_content]):
            resp = code.get_msg(code.PARAMETER_ERROR)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        business = Business.objects.filter(pk=business_id, del_flag=0).first()
        if business is None:
            resp = code.get_msg(code.BUSINESS_NOT_EXIST)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        if int(business.node_id) != int(node_id):
            resp = code.get_msg(code.PARAMETER_ERROR)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        if business.status == 2:
            post_name = post_name.strip()
            post_content = post_content.strip()
            if len(post_name) == 0 or len(post_name) > 32 or len(post_content) == 0:
                resp = code.get_msg(code.PARAMETER_ERROR)
                return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
            saved_data_id = str(business.id) + '-' + str(node_id)
            html_file = saved_data_id + '.html'
            docx_file = saved_data_id + '.docx'
            html_filename = 'media/files/business/' + html_file
            docx_file_name = 'media/files/business/' + docx_file

            f = codecs.open(html_filename, "a+", "utf-8")
            f.write(post_content)
            f.close()
            pypandoc.convert_file(html_filename, 'docx', outputfile=docx_file_name)
            docxObj = UploadFile.objects.create(filename=docx_file, file='files/business/' + docx_file,
                                                created_by=request.user.id)
            htmlObj = UploadFile.objects.create(filename=html_file, file='files/business/' + html_file,
                                                created_by=request.user.id)
            bp = BusinessPost.objects.filter(business_id=business_id, node_id=node_id)
            if bp.first():
                bp.update(name=post_name, content=post_content, docx_id=docxObj.id, html_id=htmlObj.id,
                          created_by=request.user)
            else:
                BusinessPost.objects.create(business_id=business_id, node_id=node_id, name=post_name,
                                            content=post_content, docx_id=docxObj.id, html_id=htmlObj.id,
                                            created_by=request.user)

            resp = code.get_msg(code.SUCCESS)
        elif business.status == 1:
            resp = code.get_msg(code.BUSINESS_HAS_NOT_STARTED)
        else:
            resp = code.get_msg(code.BUSINESS_HAS_FINISHED)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    except Exception as e:
        logger.exception('api_business_display_application Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


def api_business_post_info(request):
    resp = auth_check(request, "GET")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
    try:
        post_id = request.GET.get('id', None)  # 实验id
        if post_id is None:
            resp = code.get_msg(code.PARAMETER_ERROR)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        businessPost = BusinessPost.objects.filter(pk=post_id).first()
        if businessPost is None:
            resp = code.get_msg(code.PARAMETER_ERROR)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        resp = code.get_msg(code.SUCCESS)
        resp['d'] = model_to_dict(businessPost)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    except Exception as e:
        logger.exception('api_business_post_info Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


#
def api_vote_get_init_data(request):
    resp = auth_check(request, "POST")
    observable = False
    if resp != {}:
        observable = True
    try:
        business_id = int(request.POST.get('business_id', None))
        node_id = int(request.POST.get('node_id', None))
        role = int(request.POST.get('role', None))
        if business_id is None or node_id is None:
            resp = code.get_msg(code.PARAMETER_ERROR)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
        if observable and (not is_look_on_node(node_id) or role != 0):
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
        bus = Business.objects.filter(pk=business_id, del_flag=0).first()
        if bus is None:
            resp = code.get_msg(code.BUSINESS_NOT_EXIST)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        if bus.status == 2:
            vote = Vote.objects.filter(
                business_id=business_id,
                node_id=node_id
            ).first()

            if role == 0:
                if vote is None:
                    resp = code.get_msg(code.SUCCESS)
                    resp['d'] = {'status': 2, 'data': "还没进行表决设置，请等待"}
                    return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
                else:
                    return am_i_vote_member(request, None)

            if vote is None:
                data = {
                    'node_members': [{
                        'value': BusinessTeamMember.objects.filter(business_role_id=item.role_id,
                                                                   no=item.no).first().user_id,
                        'text': BusinessTeamMember.objects.filter(business_role_id=item.role_id,
                                                                  no=item.no).first().user.name
                        if BusinessTeamMember.objects.filter(business_role_id=item.role_id,
                                                             no=item.no).first().user else ''
                    } for item in BusinessRoleAllocation.objects.filter(business_id=business_id, node_id=node_id)],
                }
                resp = code.get_msg(code.SUCCESS)
                resp['d'] = {'status': 1, 'data': data}
            elif vote.end_time > timezone.now():
                statusMsg = "参与人员在输入表决选项，请等待" if vote.mode == 4 else "参与人员在投票，请等待"
                return am_i_vote_member(request, statusMsg)
            elif vote.mode == 4:
                data = {
                    'title': vote.title,
                    'description': vote.description,
                    'mode': 4,
                    'node_members': [{
                        'value': BusinessTeamMember.objects.filter(business_role_id=n_member.role_id,
                                                                   no=n_member.no).first().user_id,
                        'text': BusinessTeamMember.objects.filter(business_role_id=n_member.role_id,
                                                                  no=n_member.no).first().user.name
                    } for n_member in BusinessRoleAllocation.objects.filter(business_id=business_id, node_id=node_id)],
                    'items': [{
                        'id': item.pk,
                        'text': item.content
                    } for item in vote.items.all()]
                }
                resp = code.get_msg(code.SUCCESS)
                resp['d'] = {'status': 3, 'data': data}
            elif vote.mode == 3:
                data = {
                    'title': vote.title,
                    'description': vote.description,
                    'mode': 3,
                    'method': vote.method,
                    'max_vote': vote.max_vote,
                    'lost_vote': vote.lost_vote,
                    'node_members': [{
                        'value': BusinessTeamMember.objects.filter(business_role_id=n_member.role_id,
                                                                   no=n_member.no).first().user_id,
                        'text': BusinessTeamMember.objects.filter(business_role_id=n_member.role_id,
                                                                  no=n_member.no).first().user.name
                    } for n_member in BusinessRoleAllocation.objects.filter(business_id=business_id, node_id=node_id)],
                    'members': [{
                        'id': member.pk,
                        'username': member.user.name,
                        'voted': member.voted,
                    } for member in vote.members.all()],
                    'items': [{
                        'id': item.pk,
                        'text': item.content,
                        'voted_count': item.voted_count,
                        'voted_users': [user.name for user in item.voted_users.all()]
                    } for item in vote.items.all()]
                }
                resp = code.get_msg(code.SUCCESS)
                resp['d'] = {'status': 4, 'data': data}
            else:
                data = {
                    'title': vote.title,
                    'description': vote.description,
                    'mode': vote.mode,
                    'method': vote.method,
                    'members': [{
                        'id': member.pk,
                        'username': member.user.name,
                        'voted': member.voted,
                    } for member in vote.members.all()],
                    'items': [{
                        'id': item.pk,
                        'text': item.content,
                        'voted_count': item.voted_count,
                        'voted_users': [user.name for user in item.voted_users.all()]
                    } for item in vote.items.all()]
                }
                resp = code.get_msg(code.SUCCESS)
                resp['d'] = {'status': 5, 'data': data}
        elif bus.status == 1:
            resp = code.get_msg(code.BUSINESS_HAS_NOT_STARTED)
        else:
            resp = code.get_msg(code.BUSINESS_HAS_FINISHED)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    except Exception as e:
        logger.exception('api_business_display_application Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


#
def api_vote_save_vote_data(request):
    resp = auth_check(request, "POST")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
    try:
        voteMode = request.POST.get('voteMode', None)
        voteData = eval(request.POST.get('voteData', None))
        voteSetting = eval(request.POST.get('voteSetting', None))

        business_id = request.POST.get('business_id', None)
        node_id = request.POST.get('node_id', None)

        bus = Business.objects.filter(pk=business_id, del_flag=0).first()
        if bus is None:
            resp = code.get_msg(code.BUSINESS_NOT_EXIST)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        if bus.status == 2:
            vote = Vote.objects.filter(
                business_id=business_id,
                node_id=node_id
            ).first()

            if vote is not None:
                vote.items.all().delete()
                vote.members.all().delete()
                vote.delete()
            newVote = Vote(
                business_id=business_id,
                node_id=node_id,
                mode=voteMode,
                title=voteData['voteTitle'],
                description=voteData['voteDescription'],
                method=voteSetting['voteMethod'] if int(voteMode) != 4 else None,
                end_time=voteSetting['voteEndTime'],
                max_vote=voteSetting['voteMaxVote'] if int(voteMode) == 3 else None,
                lost_vote=voteSetting['voteLostVote'] if int(voteMode) == 3 else None
            )
            newVote.save()

            if int(voteMode) != 4:
                for item in voteData['voteItems']:
                    newVoteItem = VoteItem(content=item['text'])
                    newVoteItem.save()
                    newVote.items.add(newVoteItem)
            for vm in voteSetting['members']:
                newMember = VoteMember(user_id=vm)
                newMember.save()
                newVote.members.add(newMember)

            statusMsg = "参与人员在输入表决选项，请等待" if int(voteMode) == 4 else "参与人员在投票，请等待"
            return am_i_vote_member(request, statusMsg)
        elif bus.status == 1:
            resp = code.get_msg(code.BUSINESS_HAS_NOT_STARTED)
        else:
            resp = code.get_msg(code.BUSINESS_HAS_FINISHED)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    except Exception as e:
        logger.exception('api_business_display_application Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


#
def api_vote_finish_mode_3(request):
    resp = auth_check(request, "POST")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
    try:
        business_id = request.POST.get('business_id', None)
        node_id = request.POST.get('node_id', None)

        bus = Business.objects.filter(pk=business_id, del_flag=0).first()
        if bus is None:
            resp = code.get_msg(code.BUSINESS_NOT_EXIST)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        if bus.status == 2:
            Vote.objects.filter(
                business_id=business_id,
                node_id=node_id,
                mode=3
            ).update(mode=0)

            vote = Vote.objects.filter(
                business_id=business_id,
                node_id=node_id,
                mode=0
            ).first()

            data = {
                'title': vote.title,
                'description': vote.description,
                'mode': vote.mode,
                'method': vote.method,
                'members': [{
                    'id': member.pk,
                    'username': member.user.name,
                    'voted': member.voted,
                } for member in vote.members.all()],
                'items': [{
                    'id': item.pk,
                    'text': item.content,
                    'voted_count': item.voted_count,
                    'voted_users': [user.name for user in item.voted_users.all()]
                } for item in vote.items.all()]
            }
            resp = code.get_msg(code.SUCCESS)
            resp['d'] = {'status': 5, 'data': data}
        elif bus.status == 1:
            resp = code.get_msg(code.BUSINESS_HAS_NOT_STARTED)
        else:
            resp = code.get_msg(code.BUSINESS_HAS_FINISHED)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    except Exception as e:
        logger.exception('api_business_display_application Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


def api_business_jump_start(request):
    resp = auth_check(request, "POST")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    try:
        business_id = request.POST.get("business_id", None)
        project_id = request.POST.get("project_id", None)  # 项目ID
        use_to = request.POST.get("use_to", None)
        role_alloc_id = request.POST.get("role_alloc_id", None)
        process_type = request.POST.get("process_type", None)
        tran_id = request.POST.get("tran_id", None)

        if not all(v is not None for v in [business_id, project_id, role_alloc_id, process_type, tran_id]):
            resp = code.get_msg(code.PARAMETER_ERROR)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
        business = Business.objects.filter(pk=business_id).first()
        if business is None:
            resp = code.get_msg(code.BUSINESS_NOT_EXIST)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        if not BusinessRoleAllocation.objects.filter(pk=role_alloc_id, project_id=business.cur_project_id,
                                                     can_terminate=True).exists():
            resp = code.get_msg(code.PERMISSION_DENIED)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
        project = Project.objects.filter(pk=project_id).first()
        if project is None:
            resp = code.get_msg(code.PROJECT_NOT_EXIST)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        roles = ProjectRole.objects.filter(project_id=project_id)
        if roles.exists() is False:
            resp = code.get_msg(code.PROJECT_ROLE_NOT_EXIST)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        with transaction.atomic():
            if project.created_role_id in [3, 7]:
                company_id = project.created_by.tcompanymanagers_set.get().tcompany.id if project.created_role_id == 3 else project.created_by.t_company_set_assistants.get().id if project.created_role_id == 7 else None
            if business.path_id is None:
                resp = code.get_msg(code.PERMISSION_DENIED)
                return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
            if int(process_type) == 9:
                BusinessProjectTrack.objects.create(business_id=business_id, project_id=business.cur_project_id,
                                                    process_type=process_type, flow_trans_id=tran_id)
            business.path_id = None
            business.node_id = None
            business.cur_project_id = project_id
            business.jumper_id = request.user.id
            business.target_company_id = use_to if project.created_role_id in [2,
                                                                               6] else company_id if project.created_role_id in [
                3, 7] and project.use_to_id is None else None
            business.target_part_id = project.use_to_id if project.created_role_id in [3,
                                                                                       7] and project.use_to_id is not None else None
            business.save()

            business_roles = []
            for item in roles:
                business_roles.append(BusinessRole(business=business, name=item.name,
                                                   type=item.type, flow_role_id=item.flow_role_id,
                                                   project_role_id=item.id, project_id=project_id,
                                                   category=item.category, capacity=item.capacity,
                                                   job_type=item.job_type))
            BusinessRole.objects.bulk_create(business_roles)

            # 复制流程角色分配设置
            business_allocations = []
            allocations = ProjectRoleAllocation.objects.filter(project_id=project_id)
            for item in allocations:
                # 将角色分配中的role_id设置为ProjectRole id
                role = BusinessRole.objects.filter(business=business, project_role_id=item.role_id).first()
                projectRole = ProjectRole.objects.filter(pk=item.role_id).first()
                if projectRole is None:
                    continue
                flow_role_alloc = FlowRoleAllocation.objects.filter(flow_id=project.flow_id, node_id=item.node_id,
                                                                    role_id=projectRole.flow_role_id,
                                                                    no=item.no).first()
                if flow_role_alloc is None:
                    continue
                if role:
                    business_allocations.append(
                        BusinessRoleAllocation(business=business, node=FlowNode.objects.get(pk=item.node_id),
                                               project_id=project_id,
                                               project_role_alloc_id=item.id,
                                               flow_role_alloc_id=flow_role_alloc.id,
                                               role=role,
                                               can_start=item.can_start,
                                               can_terminate=item.can_terminate,
                                               can_brought=item.can_brought,
                                               can_take_in=item.can_take_in,
                                               no=item.no))
            BusinessRoleAllocation.objects.bulk_create(business_allocations)

            res = teammates_configuration(business.id, [])
            if (res == 'team_configured'):
                return api_business_detail(request)
            resp = code.get_msg(code.SUCCESS)
            resp['d'] = 'jump_team_not_configured'
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    except Exception as e:
        logger.exception('api_business_create Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


#
def am_i_vote_member(request, statusMsg):
    resp = auth_check(request, "POST")
    observable = False
    if resp != {}:
        observable = True
    try:
        business_id = request.POST.get('business_id', None)
        node_id = request.POST.get('node_id', None)

        bus = Business.objects.filter(pk=business_id, del_flag=0).first()
        if bus is None:
            resp = code.get_msg(code.BUSINESS_NOT_EXIST)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        if bus.status == 2:
            vote = Vote.objects.filter(
                business_id=business_id,
                node_id=node_id
            ).first()
            if not vote.members.filter(user_id=request.user.pk).exists():
                resp = code.get_msg(code.SUCCESS)
                if statusMsg is None:
                    resp['d'] = {'status': 2, 'data': "您不能参与表决"}
                elif vote.mode != 4 and not vote.members.filter(voted=0).exists():
                    data = {
                        'title': vote.title,
                        'description': vote.description,
                        'mode': vote.mode,
                        'method': vote.method,
                        'members': [{
                            'id': member.pk,
                            'username': member.user.name,
                            'voted': member.voted,
                        } for member in vote.members.all()],
                        'items': [{
                            'id': item.pk,
                            'text': item.content,
                            'voted_count': item.voted_count,
                            'voted_users': [user.name for user in item.voted_users.all()]
                        } for item in vote.items.all()]
                    }
                    resp['d'] = {'status': 5, 'data': data}
                else:
                    resp['d'] = {'status': 2, 'data': statusMsg}
                return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
            elif vote.members.get(user_id=request.user.pk).voted == 1:
                resp = code.get_msg(code.SUCCESS)
                if vote.mode == 4:
                    resp['d'] = {'status': 2, 'data': "您已经输入了表决选项"}
                elif vote.end_time <= timezone.now() or not vote.members.filter(voted=0).exists():
                    data = {
                        'title': vote.title,
                        'description': vote.description,
                        'mode': vote.mode,
                        'method': vote.method,
                        'members': [{
                            'id': member.pk,
                            'username': member.user.name,
                            'voted': member.voted,
                        } for member in vote.members.all()],
                        'items': [{
                            'id': item.pk,
                            'text': item.content,
                            'voted_count': item.voted_count,
                            'voted_users': [user.name for user in item.voted_users.all()]
                        } for item in vote.items.all()]
                    }
                    resp['d'] = {'status': 5, 'data': data}
                else:
                    resp['d'] = {'status': 2, 'data': "您已经进行表决"}
                return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
            elif vote.end_time <= timezone.now():
                resp = code.get_msg(code.SUCCESS)
                if vote.mode == 4:
                    resp['d'] = {'status': 2, 'data': "输入表决选项的时间已经过了"}
                else:
                    data = {
                        'title': vote.title,
                        'description': vote.description,
                        'mode': vote.mode,
                        'method': vote.method,
                        'members': [{
                            'id': member.pk,
                            'username': member.user.name,
                            'voted': member.voted,
                        } for member in vote.members.all()],
                        'items': [{
                            'id': item.pk,
                            'text': item.content,
                            'voted_count': item.voted_count,
                            'voted_users': [user.name for user in item.voted_users.all()]
                        } for item in vote.items.all()]
                    }
                    resp['d'] = {'status': 5, 'data': data}
                return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
            elif vote.mode == 4:
                resp = code.get_msg(code.SUCCESS)
                data = {
                    'title': vote.title,
                    'description': vote.description,
                }
                resp['d'] = {'status': 6, 'data': data}
                return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
            else:
                resp = code.get_msg(code.SUCCESS)
                data = {
                    'title': vote.title,
                    'description': vote.description,
                    'mode': vote.mode,
                    'method': vote.method,
                    'max_vote': vote.max_vote,
                    'items': [{
                        'value': item.pk,
                        'text': item.content,
                    } for item in vote.items.all()]
                }
                resp['d'] = {'status': 7, 'data': data}
                return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        elif bus.status == 1:
            resp = code.get_msg(code.BUSINESS_HAS_NOT_STARTED)
        else:
            resp = code.get_msg(code.BUSINESS_HAS_FINISHED)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    except Exception as e:
        logger.exception('am_i_vote_member Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


#
def api_user_vote_save(request):
    resp = auth_check(request, "POST")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
    try:
        business_id = request.POST.get('business_id', None)
        node_id = request.POST.get('node_id', None)
        items = eval(request.POST.get('items', None))

        bus = Business.objects.filter(pk=business_id, del_flag=0).first()
        if bus is None:
            resp = code.get_msg(code.BUSINESS_NOT_EXIST)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        if bus.status == 2:
            vote = Vote.objects.filter(
                business_id=business_id,
                node_id=node_id
            ).first()

            for item in items:
                if vote.method == 0:
                    VoteItem.objects.get(pk=item).voted_users.add(request.user.pk)
                elif vote.method == 1:
                    VoteItem.objects.filter(pk=item).update(voted_count=F('voted_count') + 1)

            vote.members.filter(user_id=request.user.pk).update(voted=1)

            resp = code.get_msg(code.SUCCESS)
            if vote.end_time <= timezone.now() or not vote.members.filter(voted=0).exists():
                data = {
                    'title': vote.title,
                    'description': vote.description,
                    'mode': vote.mode,
                    'method': vote.method,
                    'members': [{
                        'id': member.pk,
                        'username': member.user.name,
                        'voted': member.voted,
                    } for member in vote.members.all()],
                    'items': [{
                        'id': item.pk,
                        'text': item.content,
                        'voted_count': item.voted_count,
                        'voted_users': [user.name for user in item.voted_users.all()]
                    } for item in vote.items.all()]
                }
                resp['d'] = {'status': 5, 'data': data}
            else:
                resp['d'] = {'status': 2, 'data': "您已经进行表决"}
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        elif bus.status == 1:
            resp = code.get_msg(code.BUSINESS_HAS_NOT_STARTED)
        else:
            resp = code.get_msg(code.BUSINESS_HAS_FINISHED)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    except Exception as e:
        logger.exception('am_i_vote_member Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


#
def api_user_vote_item_save(request):
    resp = auth_check(request, "POST")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
    try:
        business_id = request.POST.get('business_id', None)
        node_id = request.POST.get('node_id', None)
        new_item = request.POST.get('new_item', None)

        bus = Business.objects.filter(pk=business_id, del_flag=0).first()
        if bus is None:
            resp = code.get_msg(code.BUSINESS_NOT_EXIST)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        if bus.status == 2:
            vote = Vote.objects.filter(
                business_id=business_id,
                node_id=node_id
            ).first()

            newVoteItem = VoteItem(content=new_item)
            newVoteItem.save()
            vote.items.add(newVoteItem)
            vote.members.filter(user_id=request.user.pk).update(voted=1)

            resp = code.get_msg(code.SUCCESS)
            resp['d'] = {'status': 2, 'data': "您已经输入了表决选项"}
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        elif bus.status == 1:
            resp = code.get_msg(code.BUSINESS_HAS_NOT_STARTED)
        else:
            resp = code.get_msg(code.BUSINESS_HAS_FINISHED)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    except Exception as e:
        logger.exception('am_i_vote_member Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


#
def api_get_poll_init_data(request):
    resp = auth_check(request, "POST")
    observable = False
    if resp != {}:
        observable = True
    try:
        business_id = request.POST.get('business_id', None)
        node_id = request.POST.get('node_id', None)
        role = int(request.POST.get('role', None))
        is_observable = request.POST.get('observable', None)
        if observable and is_observable is None and role == 1:
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        if business_id is None or node_id is None:
            resp = code.get_msg(code.PARAMETER_ERROR)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        bus = Business.objects.filter(pk=business_id, del_flag=0).first()
        if bus is None:
            resp = code.get_msg(code.BUSINESS_NOT_EXIST)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        if bus.status == 2:
            poll = Poll.objects.filter(
                business_id=business_id,
                node_id=node_id
            ).first()

            if role == 1:
                if poll is None:
                    data = {
                        'node_members': [],
                    }
                    for item in BusinessRoleAllocation.objects.filter(business_id=business_id, node_id=node_id):
                        btm = BusinessTeamMember.objects.filter(business_role_id=item.role_id, no=item.no).first()
                        if not btm.user_id:
                            continue
                        data['node_members'].append({
                            'value': btm.user_id,
                            'text': btm.user.name
                        })
                    resp = code.get_msg(code.SUCCESS)
                    resp['d'] = {'status': 1, 'data': data}
                elif poll.end_time > timezone.now():
                    if not poll.members.filter(poll_status=0).exists():
                        data = {
                            'title': poll.title,
                            'method': poll.method,
                            'share': poll.share,
                            'items': {
                                '1': [x.user.name for x in poll.members.filter(poll_status=1)],
                                '2': [x.user.name for x in poll.members.filter(poll_status=2)],
                                '3': [x.user.name for x in poll.members.filter(Q(poll_status=3) | Q(poll_status=0))],
                            } if poll.share == 0 else {
                                '1': poll.members.filter(poll_status=1).count(),
                                '2': poll.members.filter(poll_status=2).count(),
                                '3': poll.members.filter(Q(poll_status=3) | Q(poll_status=0)).count(),
                            },
                        }
                        resp = code.get_msg(code.SUCCESS)
                        resp['d'] = {'status': 4, 'data': data}
                    elif not poll.members.filter(user_id=request.user.pk).exists():
                        resp = code.get_msg(code.SUCCESS)
                        resp['d'] = {'status': 2, 'data': "请等到投票结束"}
                    else:
                        if poll.members.get(user_id=request.user.pk).poll_status == 0:
                            data = {
                                'title': poll.title,
                                'method': poll.method,
                                'share': poll.share,
                            }
                            resp = code.get_msg(code.SUCCESS)
                            resp['d'] = {'status': 3, 'data': data}
                        elif not poll.members.filter(poll_status=0).exists():
                            data = {
                                'title': poll.title,
                                'method': poll.method,
                                'share': poll.share,
                                'items': {
                                    '1': [x.user.name for x in poll.members.filter(poll_status=1)],
                                    '2': [x.user.name for x in poll.members.filter(poll_status=2)],
                                    '3': [x.user.name for x in
                                          poll.members.filter(Q(poll_status=3) | Q(poll_status=0))],
                                } if poll.share == 0 else {
                                    '1': poll.members.filter(poll_status=1).count(),
                                    '2': poll.members.filter(poll_status=2).count(),
                                    '3': poll.members.filter(Q(poll_status=3) | Q(poll_status=0)).count(),
                                },
                            }
                            resp = code.get_msg(code.SUCCESS)
                            resp['d'] = {'status': 4, 'data': data}
                        else:
                            resp = code.get_msg(code.SUCCESS)
                            resp['d'] = {'status': 2, 'data': "已经投票了"}
                else:
                    data = {
                        'title': poll.title,
                        'method': poll.method,
                        'share': poll.share,
                        'items': {
                            '1': [x.user.name for x in poll.members.filter(poll_status=1)],
                            '2': [x.user.name for x in poll.members.filter(poll_status=2)],
                            '3': [x.user.name for x in poll.members.filter(Q(poll_status=3) | Q(poll_status=0))],
                        } if poll.share == 0 else {
                            '1': poll.members.filter(poll_status=1).count(),
                            '2': poll.members.filter(poll_status=2).count(),
                            '3': poll.members.filter(Q(poll_status=3) | Q(poll_status=0)).count(),
                        },
                    }
                    resp = code.get_msg(code.SUCCESS)
                    resp['d'] = {'status': 4, 'data': data}
            else:
                if poll is None:
                    resp = code.get_msg(code.SUCCESS)
                    resp['d'] = {'status': 2, 'data': "还没有进行投票设置"}
                elif is_observable != '1' and not poll.members.filter(user_id=request.user.pk).exists():
                    resp = code.get_msg(code.SUCCESS)
                    resp['d'] = {'status': 2, 'data': "不能参与投票"}
                elif poll.end_time <= timezone.now():
                    if poll.share == 0 or is_observable == 1:
                        data = {
                            'title': poll.title,
                            'method': poll.method,
                            'share': poll.share,
                            'items': {
                                '1': [x.user.name for x in poll.members.filter(poll_status=1)],
                                '2': [x.user.name for x in poll.members.filter(poll_status=2)],
                                '3': [x.user.name for x in poll.members.filter(Q(poll_status=3) | Q(poll_status=0))],
                            } if poll.share == 0 else {
                                '1': poll.members.filter(poll_status=1).count(),
                                '2': poll.members.filter(poll_status=2).count(),
                                '3': poll.members.filter(Q(poll_status=3) | Q(poll_status=0)).count(),
                            },
                        }
                        resp = code.get_msg(code.SUCCESS)
                        resp['d'] = {'status': 4, 'data': data}
                    else:
                        resp = code.get_msg(code.SUCCESS)
                        resp['d'] = {'status': 2, 'data': "已经过了投票时间"}
                else:
                    if is_observable != '1' and poll.members.get(user_id=request.user.pk).poll_status == 0:
                        data = {
                            'title': poll.title,
                            'method': poll.method,
                            'share': poll.share,
                        }
                        resp = code.get_msg(code.SUCCESS)
                        resp['d'] = {'status': 3, 'data': data}
                    elif (is_observable == 1 or poll.share == 0) and not poll.members.filter(poll_status=0).exists():
                        data = {
                            'title': poll.title,
                            'method': poll.method,
                            'share': poll.share,
                            'items': {
                                '1': [x.user.name for x in poll.members.filter(poll_status=1)],
                                '2': [x.user.name for x in poll.members.filter(poll_status=2)],
                                '3': [x.user.name for x in poll.members.filter(Q(poll_status=3) | Q(poll_status=0))],
                            } if poll.share == 0 else {
                                '1': poll.members.filter(poll_status=1).count(),
                                '2': poll.members.filter(poll_status=2).count(),
                                '3': poll.members.filter(Q(poll_status=3) | Q(poll_status=0)).count(),
                            },
                        }
                        resp = code.get_msg(code.SUCCESS)
                        resp['d'] = {'status': 4, 'data': data}
                    else:
                        resp = code.get_msg(code.SUCCESS)
                        resp['d'] = {'status': 2, 'data': "已经投票了"}
        elif bus.status == 1:
            resp = code.get_msg(code.BUSINESS_HAS_NOT_STARTED)
        else:
            resp = code.get_msg(code.BUSINESS_HAS_FINISHED)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    except Exception as e:
        logger.exception('api_business_display_application Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


#
def api_save_poll_data(request):
    resp = auth_check(request, "POST")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
    try:
        business_id = request.POST.get('business_id', None)
        node_id = request.POST.get('node_id', None)
        data = eval(request.POST.get('data', None))

        bus = Business.objects.filter(pk=business_id, del_flag=0).first()
        if bus is None:
            resp = code.get_msg(code.BUSINESS_NOT_EXIST)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        if bus.status == 2:
            poll = Poll.objects.filter(
                business_id=business_id,
                node_id=node_id
            ).first()

            if poll is not None:
                poll.members.all().delete()
                poll.delete()
            newPoll = Poll(
                business_id=business_id,
                node_id=node_id,
                title=data['pollTitle'],
                method=data['pollMethod'],
                end_time=data['pollEndTime'],
                share=data['pollShare'],
            )
            newPoll.save()

            for pm in data['members']:
                newMember = PollMember(user_id=pm)
                newMember.save()
                newPoll.members.add(newMember)

            poll = newPoll

            if not poll.members.filter(user_id=request.user.pk).exists():
                resp = code.get_msg(code.SUCCESS)
                resp['d'] = {'status': 2, 'data': "请等到投票结束"}
            else:
                if poll.members.get(user_id=request.user.pk).poll_status == 0:
                    data = {
                        'title': poll.title,
                        'method': poll.method,
                        'share': poll.share,
                    }
                    resp = code.get_msg(code.SUCCESS)
                    resp['d'] = {'status': 3, 'data': data}
                elif not poll.members.filter(poll_status=0).exists():
                    data = {
                        'title': poll.title,
                        'method': poll.method,
                        'share': poll.share,
                        'items': {
                            '1': [x.user.name for x in poll.members.filter(poll_status=1)],
                            '2': [x.user.name for x in poll.members.filter(poll_status=2)],
                            '3': [x.user.name for x in poll.members.filter(Q(poll_status=3) | Q(poll_status=0))],
                        } if poll.share == 0 else {
                            '1': poll.members.filter(poll_status=1).count(),
                            '2': poll.members.filter(poll_status=2).count(),
                            '3': poll.members.filter(Q(poll_status=3) | Q(poll_status=0)).count(),
                        },
                    }
                    resp = code.get_msg(code.SUCCESS)
                    resp['d'] = {'status': 4, 'data': data}
                else:
                    resp = code.get_msg(code.SUCCESS)
                    resp['d'] = {'status': 2, 'data': "已经投票了"}
        elif bus.status == 1:
            resp = code.get_msg(code.BUSINESS_HAS_NOT_STARTED)
        else:
            resp = code.get_msg(code.BUSINESS_HAS_FINISHED)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    except Exception as e:
        logger.exception('api_save_poll_data Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


#
def api_user_poll_save(request):
    resp = auth_check(request, "POST")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
    try:
        business_id = request.POST.get('business_id', None)
        node_id = request.POST.get('node_id', None)
        set_poll = request.POST.get('poll', None)

        bus = Business.objects.filter(pk=business_id, del_flag=0).first()
        if bus is None:
            resp = code.get_msg(code.BUSINESS_NOT_EXIST)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        if bus.status == 2:
            poll = Poll.objects.filter(
                business_id=business_id,
                node_id=node_id
            ).first()

            poll.members.filter(user_id=request.user.pk).update(poll_status=set_poll)

            resp = code.get_msg(code.SUCCESS)
            if poll.share == 0 and not poll.members.filter(poll_status=0).exists():
                data = {
                    'title': poll.title,
                    'method': poll.method,
                    'share': poll.share,
                    'items': {
                        '1': [x.user.name for x in poll.members.filter(poll_status=1)],
                        '2': [x.user.name for x in poll.members.filter(poll_status=2)],
                        '3': [x.user.name for x in poll.members.filter(Q(poll_status=3) | Q(poll_status=0))],
                    } if poll.share == 0 else {
                        '1': poll.members.filter(poll_status=1).count(),
                        '2': poll.members.filter(poll_status=2).count(),
                        '3': poll.members.filter(Q(poll_status=3) | Q(poll_status=0)).count(),
                    },
                }
                resp['d'] = {'status': 4, 'data': data}
            else:
                resp['d'] = {'status': 2, 'data': "已经投票了"}
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        elif bus.status == 1:
            resp = code.get_msg(code.BUSINESS_HAS_NOT_STARTED)
        else:
            resp = code.get_msg(code.BUSINESS_HAS_FINISHED)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    except Exception as e:
        logger.exception('am_i_vote_member Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


def api_business_result(request):
    resp = auth_check(request, "GET")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    try:
        business_id = request.GET.get("business_id")  # 实验ID
        user_id = request.user.id
        business = Business.objects.filter(pk=business_id).first()
        if business:
            project = Project.objects.filter(pk=business.project_id).first()
            flow = Flow.objects.filter(pk=project.flow_id).first()
            members = BusinessTeamMember.objects.filter(business_id=business_id, del_flag=0).values_list('user_id',
                                                                                                         flat=True)

            # 小组成员
            member_list = []
            for uid in members:
                if uid is None:
                    continue
                user = Tuser.objects.get(pk=uid)
                member_list.append(user.name)

            # 各环节提交文件信息和聊天信息
            paths = BusinessTransPath.objects.filter(business_id=business_id)
            node_list = []
            for item in paths:
                node = FlowNode.objects.filter(pk=item.node_id, del_flag=0).first()
                if node.process.type == const.PROCESS_NEST_TYPE:
                    continue
                # 个人笔记
                note = BusinessNotes.objects.filter(business_id=business_id,
                                                    node_id=item.node_id, created_by=user_id).first()

                # 角色项目素材
                project_doc_list = []
                operation_guide_list = []
                project_tips_list = []

                doc_ids = FlowNodeDocs.objects.filter(flow_id=flow.pk,
                                                      node_id=item.node_id).values_list('doc_id', flat=True)
                if doc_ids:
                    operation_docs = FlowDocs.objects.filter(id__in=doc_ids, usage__in=(1, 2, 3))
                    for d in operation_docs:
                        url = ''
                        if d.file:
                            url = d.file.url
                        if d.usage == 1:
                            operation_guide_list.append({
                                'id': d.id, 'name': d.name, 'type': d.type, 'usage': d.usage,
                                'content': d.content, 'url': url, 'file_type': d.file_type
                            })
                        else:
                            project_doc_list.append({
                                'id': d.id, 'name': d.name, 'type': d.type, 'usage': d.usage,
                                'content': d.content, 'url': url, 'file_type': d.file_type
                            })

                # 获取该环节角色分配项目素材id
                doc_ids = ProjectDocRole.objects.filter(project_id=item.project_id,
                                                        node_id=item.node_id).values_list('doc_id', flat=True)

                if doc_ids:
                    # logger.info(doc_ids)
                    project_docs = ProjectDoc.objects.filter(id__in=doc_ids)
                    for d in project_docs:
                        if d.usage in [3, 4, 5, 7]:
                            is_exist = False
                            if d.usage == 3:
                                for t in project_doc_list:
                                    if d.name == t['name']:
                                        is_exist = True
                                        break
                            if not is_exist:
                                project_doc_list.append({
                                    'id': d.id, 'name': d.name, 'type': d.type, 'usage': d.usage,
                                    'content': d.content, 'url': d.file.url, 'file_type': d.file_type
                                })

                doc_list = []
                vote_status = []
                if node.process.type == 2:
                    # 如果是编辑
                    # 应用模板
                    contents = BusinessDocContent.objects.filter(business_id=business_id, node_id=item.node_id,
                                                                 has_edited=True)
                    for d in contents:
                        doc_list.append({
                            'id': d.doc_id, 'filename': d.name, 'content': d.content, 'file_type': d.file_type,
                            'signs': [{'sign_status': d.sign_status, 'sign': d.sign}],
                            'url': d.file.url if d.file else None
                        })
                    # 提交的文件
                    docs = BusinessDoc.objects.filter(business_id=business_id, node_id=item.node_id,
                                                      path_id=item.pk)
                    for d in docs:
                        sign_list = BusinessDocSign.objects.filter(doc_id=d.pk).values('sign', 'sign_status')
                        doc_list.append({
                            'id': d.id, 'filename': d.filename, 'content': d.content, 'file_type': d.file_type,
                            'signs': list(sign_list), 'url': d.file.url if d.file else None
                        })
                elif node.process.type == 3:
                    project_docs = BusinessDoc.objects.filter(business_id=business_id, node_id=item.node_id,
                                                              path_id=item.pk)
                    for d in project_docs:
                        doc_list.append({
                            'id': d.id, 'filename': d.filename, 'signs': [],
                            'url': d.file.url if d.file else None, 'content': d.content, 'file_type': d.file_type,
                        })
                elif node.process.type == 5:
                    # 如果是投票   三期 - 增加投票结果数量汇总  todo 去掉老师观察者的数量 WTF
                    vote_status_0_temp = BusinessRoleAllocationStatus.objects.filter(
                        business_id=business_id,
                        business_role_allocation__node_id=item.node_id,
                        # path_id=item.id,
                        vote_status=0)
                    vote_status_0 = []
                    # 去掉老师观察者角色的数据
                    for item0 in vote_status_0_temp:
                        role_alloc_temp = item0.business_role_allocation
                        if role_alloc_temp.role.name != const.ROLE_TYPE_OBSERVER:
                            vote_status_0.append(item0)

                    vote_status_1_temp = BusinessRoleAllocationStatus.objects.filter(
                        business_id=business_id,
                        business_role_allocation__node_id=item.node_id,
                        # path_id=item.id,
                        vote_status=1)
                    vote_status_1 = []
                    # 去掉老师观察者角色的数据
                    for item1 in vote_status_1_temp:
                        role_alloc_temp = item1.business_role_allocation
                        if role_alloc_temp.name != const.ROLE_TYPE_OBSERVER:
                            vote_status_1.append(item1)

                    vote_status_2_temp = BusinessRoleAllocationStatus.objects.filter(
                        business_id=business_id,
                        business_role_allocation__node_id=item.node_id,
                        # path_id=item.id,
                        vote_status=2)
                    vote_status_2 = []
                    # 去掉老师观察者角色的数据
                    for item2 in vote_status_2_temp:
                        role_alloc_temp = item2.business_role_allocation
                        if role_alloc_temp.name != const.ROLE_TYPE_OBSERVER:
                            vote_status_2.append(item2)

                    vote_status_9_temp = BusinessRoleAllocationStatus.objects.filter(
                        business_id=business_id,
                        business_role_allocation__node_id=item.node_id,
                        # path_id=item.id,
                        vote_status=9)
                    vote_status_9 = []
                    # 去掉老师观察者角色的数据
                    for item9 in vote_status_9_temp:
                        role_alloc_temp = item9.business_role_allocation
                        if role_alloc_temp.name != const.ROLE_TYPE_OBSERVER:
                            vote_status_9.append(item9)
                    vote_status = [{'status': '未投票', 'num': len(vote_status_0)},
                                   {'status': '同意', 'num': len(vote_status_1)},
                                   {'status': '不同意', 'num': len(vote_status_2)},
                                   {'status': '弃权', 'num': len(vote_status_9)}]
                    pass
                else:
                    # 提交的文件
                    docs = BusinessDoc.objects.filter(business_id=business_id, node_id=item.node_id,
                                                      path_id=item.id)
                    for d in docs:
                        sign_list = BusinessDocSign.objects.filter(doc_id=d.pk).values('sign', 'sign_status')
                        doc_list.append({
                            'id': d.id, 'filename': d.filename, 'content': d.content, 'file_type': d.file_type,
                            'signs': list(sign_list), 'url': d.file.url if d.file else None
                        })
                # 消息
                messages = BusinessMessage.objects.filter(business_id=business_id,
                                                          business_role_allocation__node_id=item.node_id,
                                                          path_id=item.id).order_by('timestamp')
                message_list = []
                for m in messages:
                    message = {
                        'user_name': m.user_name, 'role_name': m.role_name,
                        'msg': m.msg, 'msg_type': m.msg_type, 'ext': json.loads(m.ext),
                        'timestamp': m.timestamp.strftime('%Y-%m-%d %H:%M:%S')
                    }
                    message_list.append(message)

                node_list.append({
                    'docs': doc_list, 'messages': message_list, 'id': node.id, 'node_name': node.name,
                    'project_docs': project_doc_list,
                    'operation_guides': operation_guide_list,
                    'project_tips_list': project_tips_list,
                    'note': note.content if note else None, 'type': node.process.type if node.process else 0,
                    'vote_status': vote_status
                })

            detail = {'name': u'{0} {1}'.format(business.id, business.name), 'project_name': project.name,
                      'flow_name': flow.name, 'members': member_list,
                      'finish_time': business.finish_time.strftime('%Y-%m-%d') if business.finish_time else None,
                      'start_time': business.start_time.strftime('%Y-%m-%d') if business.start_time else None,
                      'end_time': business.end_time.strftime('%Y-%m-%d') if business.end_time else None,
                      'create_time': business.create_time.strftime('%Y-%m-%d'),
                      'flow_xml': flow.xml}
            resp = code.get_msg(code.SUCCESS)
            resp['d'] = {'detail': detail, 'nodes': node_list}
        else:
            resp = code.get_msg(code.BUSINESS_NOT_EXIST)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    except Exception as e:
        logger.exception('api_business_result Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


# added by ser for edit module
def api_business_template_create(request):
    resp = auth_check(request, "POST")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    business_id = request.POST.get('business_id', None)  # 实验id
    node_id = request.POST.get('node_id', None)  # 环节id
    doc_id = request.POST.get('doc_id', None)  # 模板素材id
    content = request.POST.get('content', '')  # 内容

    if None in (business_id, node_id, doc_id):
        resp = code.get_msg(code.PARAMETER_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    try:
        bus = Business.objects.filter(pk=business_id, del_flag=0).first()
        if bus:
            doc = BusinessDocContent.objects.filter(pk=doc_id).first()
            path = business_template_save(business_id, node_id, doc.name, content)
            BusinessDocContent.objects.filter(pk=doc_id).update(content=content, created_by=request.user.id, file=path,
                                                                has_edited=True)

            clear_cache(bus.pk)
            resp = code.get_msg(code.SUCCESS)
        else:
            resp = code.get_msg(code.BUSINESS_NOT_EXIST)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
    except Exception as e:
        logger.exception('api_business_template_create Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


# added by ser for edit module
def api_business_template_new(request):
    resp = auth_check(request, "POST")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    try:
        business_id = request.POST.get('business_id', None)
        node_id = request.POST.get('node_id', None)
        name = request.POST.get('name', '')
        content = request.POST.get('content', '')
        role_alloc_id = request.POST.get('role_alloc_id', None)

        if None in (business_id, node_id, name):
            resp = code.get_msg(code.PARAMETER_ERROR)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        bus = Business.objects.filter(pk=business_id, del_flag=0).first()
        if bus:
            name = '%s.docx' % name
            path = business_template_save(business_id, node_id, name, content)
            bdc = BusinessDocContent.objects.create(business_id=business_id, node_id=node_id,
                                                    business_role_allocation_id=role_alloc_id,
                                                    content=content, name=name, created_by=request.user,
                                                    file_type=1, file=path, has_edited=True)

            clear_cache(bus.pk)
            resp = code.get_msg(code.SUCCESS)

            resp['d'] = {'id': bdc.id, 'name': bdc.name, 'type': '', 'usage': 3,
                         'content': bdc.content, 'file_type': bdc.file_type,
                         'has_edited': bdc.has_edited, 'from': 1,
                         'sign_status': bdc.sign_status, 'sign': bdc.sign,
                         'role_alloc_id': bdc.business_role_allocation_id, 'url': bdc.file.url}
        else:
            resp = code.get_msg(code.EXPERIMENT_NOT_EXIST)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
    except Exception as e:
        logger.exception('api_business_template_new Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


# added by ser for edit module
def api_business_template_sign(request):
    resp = auth_check(request, "POST")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    business_id = request.POST.get('business_id', None)  # 实验id
    node_id = request.POST.get('node_id', None)  # 环节id
    project_role_id = request.POST.get('project_role_id', None)
    doc_id = request.POST.get('doc_id', None)  # 模板素材id
    status = request.POST.get('status', None)
    content = request.POST.get('content', None)

    logger.info('experiment_id:%s,node_id:%s,role_id:%s,doc_id:%s,status:%s' % (business_id, node_id,
                                                                                project_role_id, doc_id, status))
    if None in (business_id, node_id, doc_id, project_role_id):
        resp = code.get_msg(code.PARAMETER_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    try:
        role = ProjectRole.objects.filter(pk=project_role_id).first()
        if role is None:
            resp = code.get_msg(code.BUSINESS_NODE_ROLE_NOT_EXIST)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        bus = Business.objects.filter(pk=business_id, del_flag=0).first()
        if bus:
            sign = content
            if status == '1' or status == '2':
                BusinessDocContent.objects.filter(pk=doc_id, business_id=business_id,
                                                  node_id=node_id).update(sign_status=status, sign=sign,
                                                                          has_edited=True)
            else:
                BusinessDocContent.objects.filter(pk=doc_id, business_id=business_id,
                                                  node_id=node_id).update(sign_status=0, sign='', has_edited=True)

            doc_sign = BusinessDocContent.objects.filter(pk=doc_id).first()
            if doc_sign is None:
                resp = code.get_msg(code.PARAMETER_ERROR)
                return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

            resp = code.get_msg(code.SUCCESS)
            resp['d'] = {'sign_status': doc_sign.sign_status, 'sign': doc_sign.sign}

            clear_cache(bus.pk)
        else:
            resp = code.get_msg(code.BUSINESS_NOT_EXIST)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
    except Exception as e:
        logger.exception('api_experiment_template_sign Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


# added by ser for display module
def api_business_docs_delete(request):
    resp = auth_check(request, "POST")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    try:
        business_id = request.POST.get("business_id")  # 实验
        node_id = request.POST.get("node_id")  # 环节
        doc_id = request.POST.get("doc_id")  # 文件
        logger.info('experiment_id:%s,node_id:%s' % (business_id, node_id))
        path = BusinessTransPath.objects.filter(business_id=business_id).last()
        if path is None:
            resp = code.get_msg(code.EXPERIMENT_NODE_ERROR)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        doc = BusinessDoc.objects.filter(id=doc_id)
        if doc:
            doc.delete()

        resp = code.get_msg(code.SUCCESS)
        clear_cache(business_id)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
    except Exception as e:
        logger.exception('api_business_docs_delete Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


# added by ser
def api_business_step_status(request):
    resp = auth_check(request, "GET")
    observable = False
    if resp != {}:
        observable = True

    try:
        business_id = request.GET.get("business_id", None)
        node_id = request.GET.get("node_id", None)

        if None in (business_id, node_id):
            resp = code.get_msg(code.SYSTEM_ERROR)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        bss, created = BusinessStepStatus.objects.get_or_create(business_id=business_id, node_id=node_id,
                                                                defaults={'business_id': business_id,
                                                                          'node_id': node_id});

        resp = code.get_msg(code.SUCCESS)
        resp['d'] = {'step': bss.step}
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
    except Exception as e:
        logger.exception('get_business_step_status Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


# added by ser
def api_business_step_status_update(request):
    resp = auth_check(request, "POST")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    try:
        business_id = request.POST.get("business_id", None)
        node_id = request.POST.get("node_id", None)
        step = request.POST.get("step", 0)

        if None in (business_id, node_id, step):
            resp = code.get_msg(code.SYSTEM_ERROR)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        bss = BusinessStepStatus.objects.filter(business_id=business_id, node_id=node_id).first();
        if bss is None:  # if not exist then create it
            resp = code.get_msg(code.SYSTEM_ERROR)  # have to change to something
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        BusinessStepStatus.objects.update_or_create(business_id=business_id, node_id=node_id, defaults={'step': step})
        resp = code.get_msg(code.SUCCESS)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
    except Exception as e:
        logger.exception('get_business_step_status Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


# added by ser
def api_business_doc_team_status(request):
    resp = auth_check(request, "GET")
    observable = False
    if resp != {}:
        observable = True

    try:
        bdts_list = []

        business_id = request.GET.get("business_id", None)
        node_id = request.GET.get("node_id", None)
        business_doc_id = request.GET.get("business_doc_id", None)
        user_id = request.GET.get("user_id", None)

        if None in (business_id, node_id):
            resp = code.get_msg(code.SYSTEM_ERROR)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        if user_id is None:
            # get all members
            team_members = BusinessTeamMember.objects.filter(business_id=business_id)
        else:
            team_members = BusinessTeamMember.objects.filter(business_id=business_id, user_id=user_id)

        if business_doc_id is None:
            docs = BusinessDoc.objects.filter(business_id=business_id, node_id=node_id)
        else:
            docs = BusinessDoc.objects.filter(pk=business_doc_id)

        r = tools.generate_code(6)

        for doc in docs:
            left_users = ''
            status = 0
            for member in team_members:
                b = BusinessDocTeamStatus.objects.filter(business_id=business_id, node_id=node_id,
                                                         business_doc_id=doc.pk,
                                                         business_team_member_id=member.pk).first();
                if not member.user_id:
                    continue;
                if b is not None:
                    if b.status == 0:
                        user = Tuser.objects.filter(pk=member.user_id).first().name
                        left_users = left_users + user + ','
                    else:
                        status = 1

            if user_id is None or b is not None:
                url = '{0}?{1}'.format(doc.file.url, r) if doc.file else None
                # url = doc.file.url if doc.file else None
                bdts_list.append({'doc_id': doc.pk, 'doc_name': doc.filename, 'doc_url': url, 'left_users': left_users,
                                  'status': status})

        resp = code.get_msg(code.SUCCESS)
        resp['d'] = bdts_list
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    except Exception as e:
        logger.exception('api_business_doc_team_status Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


def get_group_userList(request):
    resp = auth_check(request, "GET")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    try:
        if request.session['login_type'] != 5:
            resp = code.get_msg(code.PERMISSION_DENIED)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
        group_id = request.user.tcompany.group_id
        guserList = [{'value': item.pk, 'text': item.name} for item in
                     Tuser.objects.filter(roles=5, tcompany__group_id=group_id)]
        resp = code.get_msg(code.SUCCESS)
        resp['d'] = {'result': guserList}
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
    except Exception as e:
        logger.exception('api_business_result Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


# added by ser
def api_business_doc_team_status_create(request):
    resp = auth_check(request, "POST")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    try:
        business_id = request.POST.get("business_id", None)
        node_id = request.POST.get("node_id", None)
        business_doc_id = request.POST.get("business_doc_id", None)
        user_id = request.POST.get("user_id", None)

        if None in (business_id, node_id):
            resp = code.get_msg(code.SYSTEM_ERROR)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        if user_id is None:
            # get all members
            team_members = BusinessTeamMember.objects.filter(business_id=business_id)
        else:
            team_members = BusinessTeamMember.objects.filter(business_id=business_id, user_id=user_id)

        if business_doc_id is None:
            docs = BusinessDoc.objects.filter(business_id=business_id, node_id=node_id)
        else:
            docs = BusinessDoc.objects.filter(pk=business_doc_id)

        for doc in docs:
            for member in team_members:
                BusinessDocTeamStatus.objects.create(business_id=business_id, node_id=node_id,
                                                     business_doc_id=doc.pk,
                                                     business_team_member_id=member.pk);

        resp = code.get_msg(code.SUCCESS)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    except Exception as e:
        logger.exception('api_business_doc_team_status Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


def api_business_doc_team_staus_update(request):
    resp = auth_check(request, "POST")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
    try:
        business_id = request.POST.get("business_id", None)
        business_doc_id = request.POST.get("business_doc_id", None)
        node_id = request.POST.get("node_id", None)
        user_id = request.POST.get("user_id", None)
        status = request.POST.get("status", None)

        if None in (business_id, business_doc_id, node_id, user_id, status):
            resp = code.get_msg(code.SYSTEM_ERROR)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        b = BusinessTeamMember.objects.filter(business_id=business_id, user_id=user_id).first()

        if b is None:
            resp = code.get_msg(code.SYSTEM_ERROR)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        BusinessDocTeamStatus.objects.filter(business_id=business_id, node_id=node_id, business_team_member_id=b.pk,
                                             business_doc_id=business_doc_id).update(status=1);
        resp = code.get_msg(code.SUCCESS)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
    except Exception as e:
        logger.exception('api_business_doc_teaem_staus_update Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


def set_none_user(request):
    resp = auth_check(request, "POST")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    try:
        if request.session['login_type'] != 5:
            resp = code.get_msg(code.PERMISSION_DENIED)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
        role_alloc_id = request.POST.get("role_alloc_id", None)
        user_id = request.POST.get("user_id", None)
        role_id = BusinessRoleAllocation.objects.get(pk=role_alloc_id).role_id
        no = BusinessRoleAllocation.objects.get(pk=role_alloc_id).no
        BusinessTeamMember.objects.filter(business_role_id=role_id, no=no).update(user_id=user_id)
        resp = code.get_msg(code.SUCCESS)
        resp['d'] = {'result': 'success'}
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
    except Exception as e:
        logger.exception('api_business_result Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


def api_buisness_prev_doc_get(request):
    resp = auth_check(request, "GET")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    try:
        business_id = request.GET.get("business_id", None)
        node_id = request.GET.get("node_id", None)
        doc_list = []

        if business_id is None:
            resp = code.get_msg(code.SYSTEM_ERROR)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        bus = Business.objects.filter(pk=business_id).first();

        if bus is None:
            resp = code.get_msg(code.SYSTEM_ERROR)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        # get project_doc
        project_docs = ProjectDoc.objects.filter(project_id=bus.project_id, usage=3)
        for item in project_docs:
            doc_list.append({
                'id': item.id, 'name': item.name, 'url': item.file.url, 'type': 'project_doc'
            })

        # get business_doc
        # exclude current node docs
        business_docs = BusinessDoc.objects.filter(business_id=bus.pk).exclude(node_id=node_id)
        for item in business_docs:
            doc_list.append({
                'id': item.id, 'name': item.filename, 'url': item.file.url, 'type': 'business_doc'
            })
        # get business_doc_content
        business_doc_contents = BusinessDocContent.objects.filter(business_id=business_id).exclude(node_id=node_id)
        for item in business_doc_contents:
            doc_list.append({
                'id': item.id, 'name': item.name, 'url': item.file.url, 'type': 'business_doc_content'
            })

        resp = code.get_msg(code.SUCCESS)
        resp['d'] = doc_list
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    except Exception as e:
        logger.exception('api_buisness_prev_doc_get Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


def api_business_doc_create_from_prev(request):
    resp = auth_check(request, "POST")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    try:
        business_id = request.POST.get("business_id", None)
        node_id = request.POST.get("node_id", None)
        doc_from = request.POST.get("doc_from", None)
        doc_id = request.POST.get("doc_id", None)
        role_alloc_id = request.POST.get("role_alloc_id", None)
        path_id = request.POST.get("path_id", None)

        if None in (business_id, node_id, doc_from, doc_id):
            resp = code.get_msg(code.SYSTEM_ERROR)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        # get doc
        if doc_from == 'business_doc':
            doc = BusinessDoc.objects.filter(pk=doc_id).first()
            filename = doc.filename
        elif doc_from == 'business_doc_content':
            doc = BusinessDocContent.objects.filter(pk=doc_id).first()
            filename = doc.name
        elif doc_from == 'project_doc':
            doc = ProjectDoc.objects.filter(pk=doc_id).first()
            filename = doc.name

        bdoc = BusinessDoc.objects.create(
            filename=filename,
            file=doc.file,
            business_id=business_id,
            node_id=node_id,
            business_role_allocation_id=role_alloc_id,
            path_id=path_id,
            file_type=doc.file_type,
            created_by=request.user
        )

        resp = code.get_msg(code.SUCCESS)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    except Exception as e:
        logger.exception('api_business_doc_create_copy Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


def api_business_survey(request):
    resp = auth_check(request, "GET")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    try:
        business_id = request.GET.get("business_id", None)
        node_id = request.GET.get("node_id", None)

        if None in (business_id, node_id):
            resp = code.get_msg(code.PARAMETER_ERROR)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        business = Business.objects.filter(pk=business_id).first();
        if business is None:
            resp = code.get_msg(code.PARAMETER_ERROR)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        qs = BusinessSurvey.objects.filter(
            business_id=business_id,
            project_id=business.cur_project_id,
            node_id=node_id
        ).first()
        bsurvey = {}
        if qs:
            bsurvey = {
                'id': qs.id, 'business_id': qs.business_id, 'project_id': qs.project_id, 'node_id': qs.node_id,
                'title': qs.title,
                'description': qs.description, 'step': qs.step,
                'start_time': qs.start_time.strftime('%Y-%m-%d') if qs.start_time else '',
                'end_time': qs.end_time.strftime(
                    '%Y-%m-%d') if qs.end_time else '', 'end_quote': qs.end_quote, 'target': qs.target
            }
            selectQuestions = BusinessQuestion.objects.filter(
                survey_id=qs.pk, type=0
            )
            blankQuestions = BusinessQuestion.objects.filter(
                survey_id=qs.pk, type=1
            )
            normalQuestions = BusinessQuestion.objects.filter(
                survey_id=qs.pk, type=2
            )

            bsurvey['select_questions'] = []
            for item in selectQuestions:
                questionCases = item.businessquestioncase_set.all()
                selectQuestion = model_to_dict(item)
                selectQuestion['question_cases'] = [model_to_dict(qc) for qc in questionCases]
                bsurvey['select_questions'].append(selectQuestion)
            bsurvey['blank_questions'] = [model_to_dict(bq) for bq in blankQuestions]
            bsurvey['normal_questions'] = [model_to_dict(nq) for nq in normalQuestions]

        resp = code.get_msg(code.SUCCESS)
        resp['d'] = bsurvey
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    except Exception as e:
        logger.exception('api_business_survey Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


def api_business_survey_create_or_update(request):
    resp = auth_check(request, "POST")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    try:
        business_id = request.POST.get("business_id", None)
        node_id = request.POST.get("node_id", None)
        title = request.POST.get("title", None)
        description = request.POST.get("description", None)
        end_quote = request.POST.get("end_quote", None)

        if None in (business_id, node_id):
            resp = code.get_msg(code.PARAMETER_ERROR)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        business = Business.objects.filter(pk=business_id).first();
        if business is None:
            resp = code.get_msg(code.PARAMETER_ERROR)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        qs = BusinessSurvey.objects.filter(
            business_id=business_id,
            project_id=business.cur_project_id,
            node_id=node_id
        ).first()
        if qs:
            if end_quote:
                qs.end_quote = end_quote
                if qs.step is None or qs.step < 5:
                    qs.step = 5
            else:
                if title:
                    qs.title = title
                if description:
                    qs.description = description
                if qs.step is None or qs.step < 1:
                    qs.step = 1
            qs.save()
        elif end_quote is None:
            BusinessSurvey.objects.create(
                business_id=business_id,
                project_id=business.cur_project_id,
                node_id=node_id,
                title=title,
                description=description,
                step=1
            )
        resp = code.get_msg(code.SUCCESS)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    except Exception as e:
        logger.exception('api_business_survey_create_or_update Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


def api_business_survey_set_select_questions(request):
    resp = auth_check(request, "POST")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    try:
        business_id = request.POST.get("business_id", None)
        node_id = request.POST.get("node_id", None)
        survey_id = request.POST.get("survey_id", None)
        select_questions = request.POST.get("select_questions", None)

        if None in (business_id, node_id, survey_id, select_questions):
            resp = code.get_msg(code.PARAMETER_ERROR)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
        select_questions = json.loads(select_questions)
        business = Business.objects.filter(pk=business_id).first()
        if business is None:
            resp = code.get_msg(code.BUSINESS_NOT_EXIST)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
        businessSurvey = BusinessSurvey.objects.filter(pk=survey_id).first()
        if businessSurvey is None:
            resp = code.get_msg(code.PARAMETER_ERROR)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
        question_ids = [sq['id'] for sq in select_questions if sq['id']]

        BusinessQuestion.objects.filter(survey_id=survey_id, type=0).exclude(id__in=question_ids).delete()
        for sq in select_questions:
            if sq['id']:
                bq = BusinessQuestion.objects.filter(pk=sq['id']).first()
                if bq is None:
                    continue
                bq.select_option = sq['select_option']
                bq.title = sq['title']
                bq.save()
                question_case_ids = [qc['id'] for qc in sq['question_cases'] if qc['id']]
                print question_case_ids
                BusinessQuestionCase.objects.filter(question_id=sq['id']).exclude(id__in=question_case_ids).delete()
                for qc in sq['question_cases']:
                    if qc['id']:
                        continue
                    bqc = BusinessQuestionCase.objects.create(
                        question_id=sq['id'],
                        case=qc['case']
                    )
            else:
                bq = BusinessQuestion.objects.create(
                    survey_id=survey_id,
                    type=0,
                    select_option=sq['select_option'],
                    title=sq['title']
                )
                for qc in sq['question_cases']:
                    bqc = BusinessQuestionCase.objects.create(
                        question_id=bq.id,
                        case=qc['case']
                    )

        if businessSurvey.step is None or businessSurvey.step < 2:
            businessSurvey.step = 2
            businessSurvey.save()
        resp = code.get_msg(code.SUCCESS)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    except Exception as e:
        logger.exception('api_business_survey_set_select_questions Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


def api_business_survey_set_blank_questions(request):
    resp = auth_check(request, "POST")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    try:
        business_id = request.POST.get("business_id", None)
        node_id = request.POST.get("node_id", None)
        survey_id = request.POST.get("survey_id", None)
        blank_questions = request.POST.get("blank_questions", None)

        if None in (business_id, node_id, survey_id, blank_questions):
            resp = code.get_msg(code.PARAMETER_ERROR)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
        blank_questions = json.loads(blank_questions)
        business = Business.objects.filter(pk=business_id).first()
        if business is None:
            resp = code.get_msg(code.BUSINESS_NOT_EXIST)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
        businessSurvey = BusinessSurvey.objects.filter(pk=survey_id).first()
        if businessSurvey is None:
            resp = code.get_msg(code.PARAMETER_ERROR)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
        question_ids = [bq['id'] for bq in blank_questions if bq['id']]
        BusinessQuestion.objects.filter(survey_id=survey_id, type=1).exclude(id__in=question_ids).delete()

        for bq in blank_questions:
            if bq['id']:
                bqQs = BusinessQuestion.objects.filter(pk=bq['id']).first()
                if bqQs is None:
                    continue
                bqQs.title = bq['title']
                bqQs.type = 1
                bqQs.save()
            else:
                bqQs = BusinessQuestion.objects.create(
                    survey_id=survey_id,
                    type=1,
                    title=bq['title']
                )

        if businessSurvey.step is None or businessSurvey.step < 3:
            businessSurvey.step = 3
            businessSurvey.save()
        resp = code.get_msg(code.SUCCESS)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    except Exception as e:
        logger.exception('api_business_survey_set_blank_questions Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


def api_business_survey_set_normal_questions(request):
    resp = auth_check(request, "POST")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    try:
        business_id = request.POST.get("business_id", None)
        node_id = request.POST.get("node_id", None)
        survey_id = request.POST.get("survey_id", None)
        normal_questions = request.POST.get("normal_questions", None)

        if None in (business_id, node_id, survey_id, normal_questions):
            resp = code.get_msg(code.PARAMETER_ERROR)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
        normal_questions = json.loads(normal_questions)
        business = Business.objects.filter(pk=business_id).first()
        if business is None:
            resp = code.get_msg(code.BUSINESS_NOT_EXIST)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
        businessSurvey = BusinessSurvey.objects.filter(pk=survey_id).first()
        if businessSurvey is None:
            resp = code.get_msg(code.PARAMETER_ERROR)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
        question_ids = [nq['id'] for nq in normal_questions if nq['id']]
        BusinessQuestion.objects.filter(survey_id=survey_id, type=2).exclude(id__in=question_ids).delete()

        for nq in normal_questions:
            if nq['id']:
                bqQs = BusinessQuestion.objects.filter(pk=nq['id']).first()
                if bqQs is None:
                    continue
                bqQs.title = nq['title']
                bqQs.type = 2
                bqQs.save()
            else:
                bqQs = BusinessQuestion.objects.create(
                    survey_id=survey_id,
                    type=2,
                    title=nq['title']
                )

        if businessSurvey.step is None or businessSurvey.step < 4:
            businessSurvey.step = 4
            businessSurvey.save()
        resp = code.get_msg(code.SUCCESS)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    except Exception as e:
        logger.exception('api_business_survey_set_blank_questions Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


def api_business_survey_publish(request):
    resp = auth_check(request, "POST")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    try:
        business_id = request.POST.get("business_id", None)
        node_id = request.POST.get("node_id", None)
        start_date = request.POST.get("start_date", None)
        end_date = request.POST.get("end_date", None)
        target = request.POST.get("target", None)

        if None in (business_id, node_id, target):
            resp = code.get_msg(code.PARAMETER_ERROR)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        business = Business.objects.filter(pk=business_id).first();
        if business is None:
            resp = code.get_msg(code.PARAMETER_ERROR)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        businessSurvey = BusinessSurvey.objects.filter(
            business_id=business_id,
            project_id=business.cur_project_id,
            node_id=node_id
        ).first()
        if businessSurvey is None:
            resp = code.get_msg(code.PARAMETER_ERROR)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
        businessSurvey.target = target
        if start_date and end_date:
            businessSurvey.start_time = datetime.strptime(start_date, '%Y-%m-%d')
            businessSurvey.end_time = datetime.strptime(end_date, '%Y-%m-%d')
        if businessSurvey.step is None or businessSurvey.step < 6:
            businessSurvey.step = 6
        businessSurvey.save()
        resp = code.get_msg(code.SUCCESS)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    except Exception as e:
        logger.exception('api_business_survey_publish Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


def api_business_survey_answer(request):
    try:
        survey_id = request.POST.get("survey_id", None)
        answers = request.POST.get("answers", None)
        user = request.user if request.user else None

        if None in (survey_id, answers):
            resp = code.get_msg(code.PARAMETER_ERROR)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        businessSurvey = BusinessSurvey.objects.filter(pk=survey_id).first();
        if businessSurvey is None:
            resp = code.get_msg(code.PARAMETER_ERROR)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
        answers = json.loads(answers)
        if user.id is not None:
            basu = BusinessSurveyAnsweredUser.objects.filter(survey_id=survey_id, user_id=user.id).first()
            if not basu:
                bsau = BusinessSurveyAnsweredUser.objects.create(survey_id=survey_id, user_id=user.id)
        else:
            bsau = BusinessSurveyAnsweredUser.objects.create(survey_id=survey_id)

        for answer in answers:
            bqQs = BusinessQuestion.objects.filter(pk=answer['id']).first()
            if not bqQs:
                continue
            if user.id is not None:
                ba = BusinessAnswer.objects.filter(survey_id=survey_id, question_id=bqQs.id, question_title=bqQs.title,
                                                   user=user).first()
                if not ba:
                    ba = BusinessAnswer.objects.create(survey_id=survey_id, question_id=bqQs.id,
                                                       question_title=bqQs.title,
                                                       user=user)
            else:
                ba = BusinessAnswer.objects.create(survey_id=survey_id, question_id=bqQs.id, question_title=bqQs.title,
                                                   answeredUser=bsau)
            if bqQs.type == 0:
                ba.question_cases.clear()
                if bqQs.select_option == 0:
                    ba.question_cases.add(BusinessQuestionCase.objects.get(pk=int(answer['answer'])))
                else:
                    for answer_id in answer['answers']:
                        ba.question_cases.add(BusinessQuestionCase.objects.get(pk=int(answer_id)))
            else:
                ba.answer = answer['answer']
                ba.save()
        resp = code.get_msg(code.SUCCESS)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    except Exception as e:
        logger.exception('api_business_survey_publish Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


def api_business_survey_public_list(request):
    try:
        page = request.GET.get("page", 1)
        size = request.GET.get("size", 10)
        search = request.GET.get("search", "")

        bsQs = BusinessSurvey.objects.filter(target=0, title__contains=search).order_by('-create_time')
        paginator = Paginator(bsQs, size)
        try:
            bsurveyQs = paginator.page(page)
        except EmptyPage:
            bsurveyQs = paginator.page(1)

        bsurveys = []

        for qs in bsurveyQs:
            bsurvey = {
                'id': qs.id, 'business_id': qs.business_id, 'project_id': qs.project_id, 'node_id': qs.node_id,
                'title': qs.title,
                'description': qs.description, 'step': qs.step,
                'created_at': qs.create_time.strftime('%Y-%m-%d') if qs.create_time else '',
                'start_time': qs.start_time.strftime('%Y-%m-%d') if qs.start_time else '',
                'end_time': qs.end_time.strftime(
                    '%Y-%m-%d') if qs.end_time else '', 'end_quote': qs.end_quote, 'target': qs.target,
                'link': '/survey/' + str(qs.id),
                'is_ended': qs.business.node_id != qs.node_id
            }
            selectQuestions = BusinessQuestion.objects.filter(
                survey_id=qs.pk, type=0
            )
            blankQuestions = BusinessQuestion.objects.filter(
                survey_id=qs.pk, type=1
            )
            normalQuestions = BusinessQuestion.objects.filter(
                survey_id=qs.pk, type=2
            )

            bsurvey['select_questions'] = []
            for item in selectQuestions:
                questionCases = item.businessquestioncase_set.all()
                selectQuestion = model_to_dict(item)
                selectQuestion['question_cases'] = [model_to_dict(qc) for qc in questionCases]
                bsurvey['select_questions'].append(selectQuestion)
            bsurvey['blank_questions'] = [model_to_dict(bq) for bq in blankQuestions]
            bsurvey['normal_questions'] = [model_to_dict(nq) for nq in normalQuestions]
            bsurveys.append(bsurvey)

        paging = {
            'count': paginator.count,
            'has_previous': bsurveyQs.has_previous(),
            'has_next': bsurveyQs.has_next(),
            'num_pages': paginator.num_pages,
            'cur_page': bsurveyQs.number,
            'page_size': size
        }

        resp = code.get_msg(code.SUCCESS)
        resp['d'] = {'surveys': bsurveys, 'paging': paging}
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    except Exception as e:
        logger.exception('api_business_survey_public_list Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


def api_business_selectDecide_get_setting(request):
    resp = auth_check(request, "POST")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
    try:
        flowNode_id = request.POST.get('flowNode_id', None)
        alloc_id = request.POST.get('alloc_id', None)
        settings = FlowNodeSelectDecide.objects.filter(flowNode_id=flowNode_id).first()
        result = SelectDecideResult.objects.filter(business_role_allocation_id=alloc_id).first()
        resp = code.get_msg(code.SUCCESS)
        if settings is None:
            resp['d'] = {
                'status': 0
            }
        elif result is not None:
            resp['d'] = {
                'status': 1
            }
        else:
            resp['d'] = {
                'title': settings.title,
                'description': settings.description,
                'mode': settings.mode,
                'items': [{
                    'value': item.pk,
                    'text': item.itemTitle,
                    'description': item.itemDescription,
                } for item in settings.items.all()],
                'status': 2
            }
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
    except Exception as e:
        logger.exception('api_workflow_role_allocation_image_update Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


def api_business_survey_public_detail(request):
    try:
        survey_id = request.GET.get("survey_id", None)

        if survey_id is None:
            resp = code.get_msg(code.PARAMETER_ERROR)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        qs = BusinessSurvey.objects.filter(pk=survey_id).first()
        if not qs or qs.target != 0:
            resp = code.get_msg(code.PERMISSION_DENIED)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
        cur_node_id = qs.business.node_id
        bsurvey = {
            'id': qs.id, 'business_id': qs.business_id, 'project_id': qs.project_id, 'node_id': qs.node_id,
            'title': qs.title,
            'description': qs.description, 'step': qs.step,
            'created_at': qs.create_time.strftime('%Y-%m-%d') if qs.create_time else '',
            'start_time': qs.start_time.strftime('%Y-%m-%d') if qs.start_time else '',
            'end_time': qs.end_time.strftime(
                '%Y-%m-%d') if qs.end_time else '', 'end_quote': qs.end_quote, 'target': qs.target,
            'link': '/survey/' + str(qs.id),
            'is_ended': cur_node_id != qs.node_id
        }
        selectQuestions = BusinessQuestion.objects.filter(
            survey_id=qs.pk, type=0
        )
        blankQuestions = BusinessQuestion.objects.filter(
            survey_id=qs.pk, type=1
        )
        normalQuestions = BusinessQuestion.objects.filter(
            survey_id=qs.pk, type=2
        )

        bsurvey['select_questions'] = []
        for item in selectQuestions:
            questionCases = item.businessquestioncase_set.all()
            selectQuestion = model_to_dict(item)
            selectQuestion['question_cases'] = [model_to_dict(qc) for qc in questionCases]
            bsurvey['select_questions'].append(selectQuestion)
        bsurvey['blank_questions'] = [model_to_dict(bq) for bq in blankQuestions]
        bsurvey['normal_questions'] = [model_to_dict(nq) for nq in normalQuestions]

        resp = code.get_msg(code.SUCCESS)
        resp['d'] = bsurvey
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    except Exception as e:
        logger.exception('api_business_survey_public_list Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


def api_business_survey_report(request):
    try:
        survey_id = request.GET.get("survey_id", None)

        if survey_id is None:
            resp = code.get_msg(code.PARAMETER_ERROR)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        bsurvey = BusinessSurvey.objects.filter(pk=survey_id).first()
        if not bsurvey:
            resp = code.get_msg(code.PARAMETER_ERROR)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        selectQuestions = BusinessQuestion.objects.filter(
            survey_id=bsurvey.pk, type=0
        )
        blankQuestions = BusinessQuestion.objects.filter(
            survey_id=bsurvey.pk, type=1
        )
        normalQuestions = BusinessQuestion.objects.filter(
            survey_id=bsurvey.pk, type=2
        )

        report = {'select_questions': [], 'blank_questions': [], 'normal_questions': []}

        for item in selectQuestions:
            selectQuestion = model_to_dict(item)
            totalAnswers = BusinessAnswer.objects.filter(question_id=item.id)
            selectQuestion['total_answers'] = totalAnswers.count()
            questionCases = item.businessquestioncase_set.all()
            selectQuestion['question_cases'] = []
            for qc in questionCases:
                questionCase = model_to_dict(qc)
                answers = BusinessAnswer.objects.filter(question_id=item.id, question_cases__id=qc.id)
                questionCase['answers'] = answers.count()
                selectQuestion['question_cases'].append(questionCase)
            report['select_questions'].append(selectQuestion)
        for item in blankQuestions:
            blankQuestion = model_to_dict(item)
            answers = BusinessAnswer.objects.filter(question_id=item.id)
            blankQuestion['total_answers'] = answers.count()
            blankQuestion['answers'] = [answer.answer for answer in answers]
            report['blank_questions'].append(blankQuestion)
        for item in normalQuestions:
            normalQuestion = model_to_dict(item)
            answers = BusinessAnswer.objects.filter(question_id=item.id)
            normalQuestion['total_answers'] = answers.count()
            normalQuestion['answers'] = [answer.answer for answer in answers]
            report['normal_questions'].append(normalQuestion)

        resp = code.get_msg(code.SUCCESS)
        resp['d'] = report
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    except Exception as e:
        logger.exception('api_business_survey_report Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


def api_business_survey_report_export(request):
    try:
        survey_id = request.GET.get("survey_id", None)

        if survey_id is None:
            resp = code.get_msg(code.PARAMETER_ERROR)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        bsurvey = BusinessSurvey.objects.filter(pk=survey_id).first()
        if not bsurvey:
            resp = code.get_msg(code.PARAMETER_ERROR)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        selectQuestions = BusinessQuestion.objects.filter(
            survey_id=bsurvey.pk, type=0
        )
        blankQuestions = BusinessQuestion.objects.filter(
            survey_id=bsurvey.pk, type=1
        )
        normalQuestions = BusinessQuestion.objects.filter(
            survey_id=bsurvey.pk, type=2
        )
        report = {'select_questions': [], 'blank_questions': [], 'normal_questions': []}
        for item in selectQuestions:
            selectQuestion = model_to_dict(item)
            totalAnswers = BusinessAnswer.objects.filter(question_id=item.id)
            selectQuestion['total_answers'] = totalAnswers.count()
            questionCases = item.businessquestioncase_set.all()
            selectQuestion['question_cases'] = []
            for qc in questionCases:
                questionCase = model_to_dict(qc)
                answers = BusinessAnswer.objects.filter(question_id=item.id, question_cases__id=qc.id)
                questionCase['answers'] = answers.count()
                selectQuestion['question_cases'].append(questionCase)
            report['select_questions'].append(selectQuestion)
        for item in blankQuestions:
            blankQuestion = model_to_dict(item)
            answers = BusinessAnswer.objects.filter(question_id=item.id)
            blankQuestion['total_answers'] = answers.count()
            blankQuestion['answers'] = [answer.answer for answer in answers]
            report['blank_questions'].append(blankQuestion)
        for item in normalQuestions:
            normalQuestion = model_to_dict(item)
            answers = BusinessAnswer.objects.filter(question_id=item.id)
            normalQuestion['total_answers'] = answers.count()
            normalQuestion['answers'] = [answer.answer for answer in answers]
            report['normal_questions'].append(normalQuestion)

        workbook = xlwt.Workbook(encoding='utf8')
        sheet = workbook.add_sheet(u"选择题")
        title = [u'题目', u'选项', u'回答人数', u'比例']
        for i in range(0, len(title)):
            sheet.write(0, i, title[i], set_style(220, True))
        row = 1
        for item in report['select_questions']:
            for qc in item['question_cases']:
                sheet.write(row, 0, item['title'])
                sheet.write(row, 1, qc['case'])
                sheet.write(row, 2, qc['answers'])
                sheet.write(row, 3, str(qc['answers'] / float(item['total_answers']) * 100) + '%')
                row += 1
            row += 1

        sheet = workbook.add_sheet(u"填空题")
        title = [u'题目', u'Answer']
        for i in range(0, len(title)):
            sheet.write(0, i, title[i], set_style(220, True))
        row = 1
        for item in report['blank_questions']:
            for answer in item['answers']:
                sheet.write(row, 0, html2text.html2text(item['title'].replace('', '_')))
                sheet.write(row, 1, answer)
                row += 1
            row += 1

        sheet = workbook.add_sheet(u"问答题")
        title = [u'题目', u'Answer']
        for i in range(0, len(title)):
            sheet.write(0, i, title[i], set_style(220, True))
        row = 1
        for item in report['normal_questions']:
            for answer in item['answers']:
                sheet.write(row, 0, html2text.html2text(item['title']))
                sheet.write(row, 1, answer)
                row += 1
            row += 1

        response = HttpResponse(content_type='application/vnd.ms-excel')
        filename = urlquote(u'Survey Report')
        response['Content-Disposition'] = u'attachment;filename=%s.xls' % filename
        workbook.save(response)
        return response
    except Exception as e:
        logger.exception('api_business_survey_report_export Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


def api_business_selectDecide_save_result(request):
    resp = auth_check(request, "POST")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
    try:
        selectedItems = eval(request.POST.get('selectedItems', []))
        alloc_id = request.POST.get('alloc_id', None)
        resp = code.get_msg(code.SUCCESS)
        result = SelectDecideResult.objects.create(
            business_role_allocation_id=alloc_id
        )
        for item in selectedItems:
            result.selectedItems.add(SelectDecideItem.objects.get(pk=item))

        resp['d'] = {'results': 'success'}
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
    except Exception as e:
        logger.exception('api_workflow_role_allocation_image_update Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


def api_business_get_guider_list(request):
    resp = auth_check(request, "POST")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    try:
        login_type = request.session['login_type']
        id = request.POST.get('id', None)
        bid = request.POST.get('bid', None)
        if login_type not in [5]:
            resp = code.get_msg(code.PERMISSION_DENIED)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        businessGuide = BusinessGuide.objects.filter(business_id=bid, request=request.user).first()
        if businessGuide is None:
            GI = request.user.tcompany.group.groupInstructors.filter(instructorItems__id=id)
            GIA = request.user.tcompany.group.groupInstructorAssistants.filter(instructorItems__id=id)
            result = [{
                'value': {'id': instructor.id, 'role': 4},
                'html': instructor.username + " (指导者)"
            } for instructor in GI]

            result += [{
                'value': {'id': instructor.id, 'role': 8},
                'html': instructor.username + " (指导者助理)"
            } for instructor in GIA]

            resp = code.get_msg(code.SUCCESS)
            resp['d'] = {'status': 0, 'results': result}
        else:
            resp = code.get_msg(code.SUCCESS)
            resp['d'] = {'status': 1, 'results': model_to_dict(businessGuide, fields=['id', 'business_id', 'guider_id',
                                                                                      'guider__username', 'role_id',
                                                                                      'request_id'])}

        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    except Exception as e:
        logger.exception('get_own_group Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


def api_business_set_guider(request):
    resp = auth_check(request, "POST")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    try:
        login_type = request.session['login_type']
        guiderRole = request.POST.get('guiderRole', None)
        guiderId = request.POST.get('guiderId', None)
        bid = request.POST.get('bid', None)
        if login_type not in [5]:
            resp = code.get_msg(code.PERMISSION_DENIED)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        businessGuide = BusinessGuide.objects.create(
            guider_id=guiderId,
            role_id=guiderRole,
            business_id=bid,
            request_id=request.user.id,
        )

        resp = code.get_msg(code.SUCCESS)
        resp['d'] = {'status': 1, 'results': model_to_dict(businessGuide,
                                                           fields=['id', 'business_id', 'guider_id', 'guider__username',
                                                                   'role_id', 'request_id'])}

        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    except Exception as e:
        logger.exception('get_own_group Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


def api_business_get_guider_message(request):
    resp = auth_check(request, "GET")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
    try:
        guider_id = request.GET.get("guider_id", None)

        message = GuideChatLog.objects.filter(guide_id=guider_id)
        message_list = []
        for m in message:
            message_list.append({
                'sender': m.sender_id,
                'name': m.sender.username,
                'msg': m.msg,
                'create_time': m.create_time.strftime('%Y-%m-%d %H:%M:%S')
            })
        resp = code.get_msg(code.SUCCESS)
        resp['d'] = {
            'results': message_list
        }
    except Exception as e:
        logger.exception('api_student_msg_list Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)

    return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


def api_business_get_chatRoom_id(request):
    resp = auth_check(request, "GET")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
    try:
        officeItem = request.GET.get("officeItem", None)

        group_id = request.user.tcompany.group_id
        ba = BusinessAsk.objects.filter(office_id=officeItem, group_id=group_id).first()
        if ba is None:
            ba = BusinessAsk.objects.create(
                office_id=officeItem,
                group_id=group_id
            )

        resp = code.get_msg(code.SUCCESS)
        resp['d'] = {'room_id': ba.id}
    except Exception as e:
        logger.exception('api_business_get_chatRoom_id Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)

    return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


def api_business_send_guider_message(request):
    resp = auth_check(request, "POST")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
    try:
        guide_id = request.POST.get("guide_id", None)
        msg = request.POST.get("msg", None)

        if None in [guide_id, msg]:
            resp = code.get_msg(code.PARAMETER_ERROR)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        gcl = GuideChatLog.objects.create(guide_id=guide_id, msg=msg, sender=request.user)
        resp = code.get_msg(code.SUCCESS)
        msgDict = {'guide_id': gcl.guide_id, 'sender': gcl.sender_id, 'name': gcl.sender.username, 'msg': gcl.msg,
                   'create_time': gcl.create_time.strftime('%Y-%m-%d %H:%M:%S')}
        with SocketIO(u'localhost', 4000, LoggingNamespace) as socketIO:
            socketIO.emit('guider_message', msgDict)
            socketIO.wait_for_callbacks(seconds=1)
    except Exception as e:
        logger.exception('api_business_send_guider_message Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)

    return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


def api_business_get_business_guide_list(request):
    resp = auth_check(request, "GET")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    try:
        businessGuide = BusinessGuide.objects.filter(guider_id=request.user.id, role_id=request.session['login_type'])

        resp = code.get_msg(code.SUCCESS)
        if businessGuide is not None:
            resp['d'] = [{
                'id': bg.id,
                'business_id': bg.business_id,
                'business_name': bg.business.name,
                'create_time': bg.create_time.strftime('%Y-%m-%d %H:%M:%S'),
                'request': bg.request.username,
            } for bg in businessGuide]
        else:
            resp['d'] = []

        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    except Exception as e:
        logger.exception('get_own_group Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


def api_business_get_chatRooms(request):
    resp = auth_check(request, "GET")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    try:
        login_type = request.session['login_type']
        if login_type not in [4, 8]:
            resp = code.get_msg(code.PERMISSION_DENIED)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        group_id = None
        if login_type == 4:
            group_id = request.user.allgroups_set_instructors.get().id
        else:
            group_id = request.user.allgroups_set_instructor_assistants.get().id

        resp = code.get_msg(code.SUCCESS)
        resp['d'] = []
        for item in request.user.instructorItems.all():
            ba = BusinessAsk.objects.filter(office=item, group_id=group_id).first()
            if ba is None:
                ba = BusinessAsk.objects.create(
                    office=item,
                    group_id=group_id
                )
            resp['d'].append({
                'room_id': ba.id,
                'group_id': group_id,
                'officeItem': item.name
            })

        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    except Exception as e:
        logger.exception('get_own_group Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


def api_business_get_chatRoom_messages(request):
    resp = auth_check(request, "POST")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
    try:
        room_id = request.POST.get("room_id", None)

        messages = AskChatLog.objects.filter(ask_id=room_id)
        message_list = []
        for m in messages:
            message_list.append({
                'sender': m.sender_id,
                'name': m.sender.username,
                'role': m.sender_role_id,
                'msg': m.msg,
                'create_time': m.create_time.strftime('%Y-%m-%d %H:%M:%S')
            })
        resp = code.get_msg(code.SUCCESS)
        resp['d'] = {
            'results': message_list
        }
    except Exception as e:
        logger.exception('api_student_msg_list Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)

    return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


def api_business_send_ask_messages(request):
    resp = auth_check(request, "POST")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
    try:
        login_type = request.session['login_type']
        room_id = request.POST.get("room_id", None)
        msg = request.POST.get("msg", None)

        if None in [room_id, msg]:
            resp = code.get_msg(code.PARAMETER_ERROR)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

        acl = AskChatLog.objects.create(ask_id=room_id, msg=msg, sender=request.user, sender_role_id=login_type)
        resp = code.get_msg(code.SUCCESS)
        msgDict = {'room_id': acl.ask_id, 'sender': acl.sender_id, 'name': acl.sender.username,
                   'role': acl.sender_role_id, 'msg': acl.msg,
                   'create_time': acl.create_time.strftime('%Y-%m-%d %H:%M:%S')}
        with SocketIO(u'localhost', 4000, LoggingNamespace) as socketIO:
            socketIO.emit('ask_message', msgDict)
            socketIO.wait_for_callbacks(seconds=1)
    except Exception as e:
        logger.exception('api_business_send_guider_message Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)

    return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


def api_business_get_init_evaluation(request):
    resp = auth_check(request, "GET")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    try:
        login_type = request.session['login_type']
        if login_type not in [4, 8]:
            resp = code.get_msg(code.PERMISSION_DENIED)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
        page = request.GET.get("page", 1)
        size = request.GET.get("size", 10)
        if login_type == 4:
            group_id = request.user.allgroups_set_instructors.get().id
        else:
            group_id = request.user.allgroups_set_instructor_assistants.get().id
        allList = Business.objects.filter(
            Q(Q(target_company__group=group_id) | Q(target_part__company__group=group_id)) & Q(
                officeItem__in=request.user.instructorItems.all()))
        paginator = Paginator(allList, size)
        try:
            allBusiness = paginator.page(page)
        except EmptyPage:
            allBusiness = paginator.page(1)
        if allBusiness:
            results = [{
                'business_id': b.id,
                'business_name': b.name,
                'create_company': b.target_company.name if b.target_company else b.target_part.company.name,
                'create_time': b.create_time.strftime('%Y-%m-%d %H:%M:%S'),
                'business_status': u'已完成',
                'members': [{
                    'business_id': b.id,
                    'user_id': member.user_id,
                    'name': member.user.username,
                    'company': '' if not member.user.tcompany else member.user.tcompany.name if not member.user.tcompany.is_default else '',
                    'part': member.user.tposition.parts.name if member.user.tposition else '',
                    'position': member.user.tposition.name if member.user.tposition else '',
                    'role': member.business_role.name,
                    'value': BusinessEvaluation.objects.get(business_id=b.id,
                                                            user_id=member.user_id).value if BusinessEvaluation.objects.filter(
                        business_id=b.id, user_id=member.user_id).first() else '',
                    'comment': BusinessEvaluation.objects.get(business_id=b.id,
                                                              user_id=member.user_id).comment if BusinessEvaluation.objects.filter(
                        business_id=b.id, user_id=member.user_id).first() else '',
                    'node_evaluation': [{
                        'alloc_id': alloc.id,
                        'node_name': alloc.node.name,
                        'node_comment': BusinessEvaluation.objects.get(
                            role_alloc_id=alloc.id).comment if BusinessEvaluation.objects.filter(
                            role_alloc_id=alloc.id).first() else '',
                    } for alloc in BusinessRoleAllocation.objects.filter(role_id=member.business_role_id, no=member.no)]
                } for member in BusinessTeamMember.objects.filter(business_id=b.id)]
            } for b in allBusiness]
        else:
            results = []

        paging = {
            'count': paginator.count,
            'has_previous': allBusiness.has_previous(),
            'has_next': allBusiness.has_next(),
            'num_pages': paginator.num_pages,
            'cur_page': allBusiness.number,
            'page_size': size
        }
        resp = code.get_msg(code.SUCCESS)
        resp['d'] = {'results': results, 'paging': paging}
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

    except Exception as e:
        logger.exception('get_own_group Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


def api_business_save_evaluation(request):
    resp = auth_check(request, "POST")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
    try:
        business_id = request.POST.get("business_id", None)
        user_id = request.POST.get("user_id", None)
        value = request.POST.get("value", None)
        comment = request.POST.get("comment", None)
        node_evaluation = eval(request.POST.get("node_evaluation", None))

        totalEvaluation = BusinessEvaluation.objects.filter(business_id=business_id, user_id=user_id)

        if totalEvaluation.first() is None:
            BusinessEvaluation.objects.create(
                business_id=business_id,
                user_id=user_id,
                comment=comment,
                value=value
            )
        else:
            totalEvaluation.update(comment=comment, value=value)

        for item in node_evaluation:
            nodeEvaluation = BusinessEvaluation.objects.filter(role_alloc_id=item['alloc_id'])
            if nodeEvaluation.first() is None:
                BusinessEvaluation.objects.create(
                    role_alloc_id=item['alloc_id'],
                    comment=item['node_comment'],
                )
            else:
                nodeEvaluation.update(comment=item['node_comment'])

        resp = code.get_msg(code.SUCCESS)
        resp['d'] = {'results': 'success'}
    except Exception as e:
        logger.exception('api_business_send_guider_message Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)

    return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


def api_bill_chapter_list(request):
    resp = auth_check(request, "POST")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
    try:
        business_id = request.POST.get("business_id", None)
        user_id = request.POST.get("user_id", None)
        value = request.POST.get("value", None)
        comment = request.POST.get("comment", None)
        node_evaluation = eval(request.POST.get("node_evaluation", None))

        totalEvaluation = BusinessEvaluation.objects.filter(business_id=business_id, user_id=user_id)

        if totalEvaluation.first() is None:
            BusinessEvaluation.objects.create(
                business_id=business_id,
                user_id=user_id,
                comment=comment,
                value=value
            )
        else:
            totalEvaluation.update(comment=comment, value=value)

        for item in node_evaluation:
            nodeEvaluation = BusinessEvaluation.objects.filter(role_alloc_id=item['alloc_id'])
            if nodeEvaluation.first() is None:
                BusinessEvaluation.objects.create(
                    role_alloc_id=item['alloc_id'],
                    comment=item['node_comment'],
                )
            else:
                nodeEvaluation.update(comment=item['node_comment'])

        resp = code.get_msg(code.SUCCESS)
        resp['d'] = {'results': 'success'}
    except Exception as e:
        logger.exception('api_business_send_guider_message Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)

    return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


#######################################################################################


def getAllBillList(bill_id, show_mode):
    res = []
    bill_name_object = BusinessBillList.objects.filter(id=bill_id).first()
    chapters_objects = bill_name_object.chapters.all().order_by("chapter_number")
    for chapters_object in chapters_objects:
        sections_objects = chapters_object.sections.all().order_by("section_number")
        for sections_object in sections_objects:
            parts_objects = sections_object.parts.all().order_by("part_number")
            for parts_object in parts_objects:
                res_json = {}
                if (show_mode == '1'):
                    res_json["chapter_id"] = chapters_object.id
                    res_json["chapter_number"] = chapters_object.chapter_number
                    res_json["chapter_title"] = chapters_object.chapter_title
                    res_json["chapter_content"] = chapters_object.chapter_content
                    res_json["section_id"] = sections_object.id
                    res_json["section_number"] = sections_object.section_number
                    res_json["section_title"] = sections_object.section_title
                    res_json["section_content"] = sections_object.section_content
                    res_json["part_id"] = parts_object.id
                    res_json["part_number"] = parts_object.part_number
                    res_json["part_title"] = parts_object.part_title
                    res_json["part_content"] = parts_object.part_content
                    res_json["part_reason"] = parts_object.part_reason
                    res.append(res_json)
                    continue
                if (show_mode == '2'):
                    res_json["chapter_id"] = chapters_object.id
                    res_json["chapter_number"] = chapters_object.chapter_number
                    res_json["chapter_title"] = chapters_object.chapter_title
                    res_json["chapter_content"] = chapters_object.chapter_content
                    res_json["section_id"] = sections_object.id
                    res_json["section_number"] = sections_object.section_number
                    res_json["section_title"] = sections_object.section_title
                    res_json["section_content"] = sections_object.section_content
                    res_json["part_id"] = parts_object.id
                    res_json["part_number"] = parts_object.part_number
                    res_json["part_title"] = parts_object.part_title
                    res_json["part_content"] = parts_object.part_content
                    res_json["part_reason"] = parts_object.part_reason
                    res.append(res_json)
                    continue
    return res


def api_bill_name_list(request):
    resp = auth_check(request, "GET")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
    try:
        business_id = request.GET.get("business_id", None)
        show_mode = request.GET.get("show_mode", None)
        bill_name = BusinessBillList.objects.filter(business_id=business_id)
        resp = code.get_msg(code.SUCCESS)
        if (len(bill_name) == 0):
            resp['d'] = {'bill_name': '', 'bill_id': 0, 'bill_data': []}
        else:
            bill_data = getAllBillList(bill_name.first().id, show_mode)
            resp['d'] = {'bill_name': bill_name.first().bill_name, 'bill_id': bill_name.first().id,
                         'bill_data': bill_data}
    except Exception as e:
        logger.exception('api_business_send_guider_message Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)

    return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


def api_bill_update_full(request):
    resp = auth_check(request, "POST")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
    try:
        chapter_id = request.POST.get("chapter_id", None)
        section_id = request.POST.get("section_id", None)
        part_id = request.POST.get("part_id", None)
        chapter_title = request.POST.get("chapter_title", None)
        section_title = request.POST.get("section_title", None)
        part_title = request.POST.get("part_title", None)
        part_content = request.POST.get("part_content", None)
        chapter = BusinessBillChapter.objects.get(id=chapter_id)
        section = BusinessBillSection.objects.get(id=section_id)
        part = BusinessBillPart.objects.get(id=part_id)
        chapter.chapter_title = chapter_title
        section.section_title = section_title
        part.part_title = part_title
        part.part_content = part_content
        chapter.save()
        section.save()
        part.save()
        resp = code.get_msg(code.SUCCESS)
    except Exception as e:
        logger.exception('api_business_send_guider_message Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)

    return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


def api_bill_update_billname(request):
    resp = auth_check(request, "POST")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
    try:
        business_id = request.POST.get("business_id", None)
        bill_name = request.POST.get("bill_name", None)
        BusinessBillList.objects.update_or_create(business_id=business_id, defaults={'bill_name': bill_name})
        resp = code.get_msg(code.SUCCESS)
    except Exception as e:
        logger.exception('api_business_send_guider_message Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)

    return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


def api_bill_part_delete(request):
    resp = auth_check(request, "POST")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
    try:
        section_id = request.POST.get("section_id", None)
        part_id = request.POST.get("part_id", None)
        part = BusinessBillPart.objects.filter(id=part_id).first()
        part_number = part.part_number
        section = BusinessBillSection.objects.filter(id=section_id).first()
        section.parts.remove(part)
        part_docs = part.part_docs.all()
        for part_doc in part_docs:
            part.part_docs.remove(part_doc)
            part_doc.delete()
        part.delete()
        for part_one in section.parts.all():
            if (part_one.part_number > part_number):
                part_one.part_number = part_one.part_number - 1
                part_one.save()
        resp = code.get_msg(code.SUCCESS)
    except Exception as e:
        logger.exception('api_business_send_guider_message Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)

    return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


def api_bill_part_add(request):
    resp = auth_check(request, "POST")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
    try:
        section_id = request.POST.get("section_id", None)
        part_number = request.POST.get("part_number", None)
        part_title = request.POST.get("part_title", None)
        part_content = request.POST.get("part_content", None)
        part_reason = request.POST.get("part_reason", None)
        added_part = BusinessBillPart.objects.create(part_number=int(part_number), part_title=part_title,
                                                     part_content=part_content, part_reason=part_reason)
        section = BusinessBillSection.objects.get(id=section_id)
        section.parts.add(added_part)
        resp = code.get_msg(code.SUCCESS)
    except Exception as e:
        logger.exception('api_business_send_guider_message Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)

    return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


def api_bill_doc_list(request):
    resp = auth_check(request, "GET")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
    try:
        part_id = request.GET.get("part_id", None)
        parts = BusinessBillPart.objects.filter(id=part_id).first()
        part_docs = parts.part_docs.all()
        res = []
        for part_doc in part_docs:
            res_one = {}
            res_one["id"] = part_doc.id
            res_one["doc_id"] = part_doc.doc_id
            res_one["doc_conception"] = part_doc.doc_conception
            res_one["doc_url"] = part_doc.doc_url
            res_one["doc_name"] = part_doc.doc_name
            res.append(res_one)
        resp = code.get_msg(code.SUCCESS)
        resp['d'] = {'doc_data': res}
    except Exception as e:
        logger.exception('api_business_send_guider_message Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)

    return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


def api_bill_doc_delete(request):
    resp = auth_check(request, "POST")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
    try:
        part_id = request.POST.get("part_id", None)
        doc_id = request.POST.get("doc_id", None)
        part = BusinessBillPart.objects.filter(id=part_id).first()
        doc = BusinessBillPartDoc.objects.filter(id=doc_id).first()
        part.part_docs.remove(doc)
        doc.delete()
        resp = code.get_msg(code.SUCCESS)
    except Exception as e:
        logger.exception('api_business_send_guider_message Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)

    return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


def api_bill_doc_upload(request):
    resp = auth_check(request, "POST")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
    try:
        doc_id = request.POST.get("doc_id", None)
        doc_url = request.POST.get("doc_url", None)
        doc_conception = request.POST.get("doc_conception", None)
        part_id = request.POST.get("part_id", None)
        doc_name = doc_url.split("/")[-1]
        added_doc = BusinessBillPartDoc.objects.create(doc_id=int(doc_id), doc_url=doc_url, doc_name=doc_name,
                                                       doc_conception=doc_conception)
        part = BusinessBillPart.objects.filter(id=part_id).first()
        part.part_docs.add(added_doc)
        resp = code.get_msg(code.SUCCESS)
    except Exception as e:
        logger.exception('api_business_send_guider_message Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)

    return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


def api_bill_part_up(request):
    resp = auth_check(request, "GET")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
    try:
        section_id = request.GET.get("section_id", None)
        part_number = request.GET.get("part_number", None)
        if (int(part_number) == 1):
            resp = code.get_msg(code.BUSINESS_BILL_NOT_UP)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
        des_part_number = int(part_number) - 1
        sections = BusinessBillSection.objects.filter(id=section_id).first()
        selected_part = sections.parts.filter(part_number=part_number)[0]
        des_part = sections.parts.filter(part_number=des_part_number)[0]
        selected_part.part_number = des_part_number
        selected_part.save()
        des_part.part_number = part_number
        des_part.save()
        resp = code.get_msg(code.SUCCESS)
    except Exception as e:
        logger.exception('api_business_send_guider_message Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)

    return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


def api_bill_part_down(request):
    resp = auth_check(request, "GET")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
    try:
        section_id = request.GET.get("section_id", None)
        part_number = request.GET.get("part_number", None)
        sections = BusinessBillSection.objects.filter(id=section_id).first()
        selected_part = sections.parts.filter(part_number=part_number)[0]
        if (int(part_number) == len(sections.parts.all())):
            resp = code.get_msg(code.BUSINESS_BILL_NOT_DOWN)
            return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
        des_part_number = int(part_number) + 1
        des_part = sections.parts.filter(part_number=des_part_number)[0]
        selected_part.part_number = des_part_number
        selected_part.save()
        des_part.part_number = part_number
        des_part.save()
        resp = code.get_msg(code.SUCCESS)
    except Exception as e:
        logger.exception('api_business_send_guider_message Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)

    return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


def api_bill_part_insert(request):
    resp = auth_check(request, "POST")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
    try:
        section_id = request.POST.get("section_id", None)
        part_number = request.POST.get("part_number", None)
        part_title = request.POST.get("part_title", None)
        part_content = request.POST.get("part_content", None)
        part_reason = request.POST.get("part_reason", None)
        section = BusinessBillSection.objects.get(id=section_id)

        all_parts = section.parts.all()
        for one_part in all_parts:
            if one_part.part_number >= int(part_number):
                one_part.part_number = one_part.part_number + 1
                one_part.save()
        added_part = BusinessBillPart.objects.create(part_number=int(part_number), part_title=part_title,
                                                     part_content=part_content, part_reason=part_reason)
        section.parts.add(added_part)
        resp = code.get_msg(code.SUCCESS)
    except Exception as e:
        logger.exception('api_business_send_guider_message Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)

    return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")


def api_bill_doc_preview(request):
    document = Document()
    section = document.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    resp = auth_check(request, "GET")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
    try:
        business_id = request.GET.get("business_id", None)
        bill_lists = BusinessBillList.objects.filter(business_id=business_id).first()
        bill_title = bill_lists.bill_name
        style = document.styles['Heading 1']
        font_bill = style.font
        font_bill.name = 'Arial'
        font_bill.size = Pt(40)
        paragraph_bill = document.add_paragraph(bill_title, style='Heading 1')
        paragraph_bill.add_run().add_break(WD_BREAK.LINE)

        chapters_lists = bill_lists.chapters.all()
        for chapter_one in chapters_lists:
            style_chapter = document.styles['List Number']
            font_chapter = style_chapter.font
            font_chapter.name = 'Arial'
            font_chapter.size = Pt(30)
            paragraph_chapter = document.add_paragraph(chapter_one.chapter_title, style='List Number')
            paragraph_chapter.add_run().add_break(WD_BREAK.LINE)

            sections_lists = chapter_one.sections.all()
            for section_one in sections_lists:
                style_section = document.styles['List Number 2']
                font_section = style_section.font
                font_section.name = 'Arial'
                font_section.size = Pt(20)
                paragraph_section = document.add_paragraph(section_one.section_title, style='List Number 2')
                paragraph_section.add_run().add_break(WD_BREAK.LINE)

                parts_lists = section_one.parts.all()
                for part_one in parts_lists:
                    style_part = document.styles['List Number 3']
                    font_part = style_part.font
                    font_part.name = 'Arial'
                    font_part.size = Pt(15)
                    paragraph_part = document.add_paragraph(part_one.part_title, style='List Number 3')
                    paragraph_part.add_run().add_break(WD_BREAK.LINE)

                    style_part_content = document.styles['Body Text']
                    font_part_content = style_part_content.font
                    font_part_content.name = 'Arial'
                    font_part_content.size = Pt(10)
                    paragraph_part_content = document.add_paragraph(part_one.part_content, style='Body Text')
                    paragraph_part_content.add_run().add_break(WD_BREAK.LINE)
                    if ((parts_lists[len(parts_lists) - 1] == part_one) and (
                                sections_lists[len(sections_lists) - 1] == section_one)):
                        if (chapters_lists[len(chapters_lists) - 1] != chapter_one):
                            paragraph_part_content.add_run().add_break(WD_BREAK.PAGE)

        document.add_page_break()
        document.save('demo.docx')
        resp = code.get_msg(code.SUCCESS)
    except Exception as e:
        logger.exception('api_business_send_guider_message Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)

    return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")

def api_bill_save(request):
    resp = auth_check(request, "POST")
    if resp != {}:
        return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
    try:
        bill_data = json.loads(request.POST.get("bill_data", None))
        business_id = request.POST.get("business_id", None)
        bill_name = request.POST.get("bill_name", None)

        chapters_all_origin = []
        sections_all_origin = []
        parts_all_origin = []
        bill_name_origin = BusinessBillList.objects.get(business_id=int(business_id))
        chapters_objects = bill_name_origin.chapters.all().order_by("chapter_number")
        for chapters_object in chapters_objects:
            chapters_all_origin_temp = {}
            chapters_all_origin_temp["chapter_id"] = chapters_object.id
            chapters_all_origin_temp["chapter_number"] = chapters_object.chapter_number
            chapters_all_origin_temp["chapter_title"] = chapters_object.chapter_title
            chapters_all_origin_temp["chapter_content"] = chapters_object.chapter_content
            chapters_all_origin.append(chapters_all_origin_temp)
            sections_objects = chapters_object.sections.all().order_by("section_number")
            for sections_object in sections_objects:
                sections_all_origin_temp = {}
                sections_all_origin_temp["chapter_id"] = chapters_object.id
                sections_all_origin_temp["chapter_number"] = chapters_object.chapter_number
                sections_all_origin_temp["chapter_title"] = chapters_object.chapter_title
                sections_all_origin_temp["chapter_content"] = chapters_object.chapter_content
                sections_all_origin_temp["section_id"] = sections_object.id
                sections_all_origin_temp["section_number"] = sections_object.section_number
                sections_all_origin_temp["section_title"] = sections_object.section_title
                sections_all_origin_temp["section_content"] = sections_object.section_content
                sections_all_origin.append(sections_all_origin_temp)
                parts_objects = sections_object.parts.all().order_by("part_number")
                for parts_object in parts_objects:
                    parts_all_origin_temp = {}
                    parts_all_origin_temp["chapter_id"] = chapters_object.id
                    parts_all_origin_temp["chapter_number"] = chapters_object.chapter_number
                    parts_all_origin_temp["chapter_title"] = chapters_object.chapter_title
                    parts_all_origin_temp["chapter_content"] = chapters_object.chapter_content
                    parts_all_origin_temp["section_id"] = sections_object.id
                    parts_all_origin_temp["section_number"] = sections_object.section_number
                    parts_all_origin_temp["section_title"] = sections_object.section_title
                    parts_all_origin_temp["section_content"] = sections_object.section_content
                    parts_all_origin_temp["part_id"] = parts_object.id
                    parts_all_origin_temp["part_number"] = parts_object.part_number
                    parts_all_origin_temp["part_title"] = parts_object.part_title
                    parts_all_origin_temp["part_content"] = parts_object.part_content
                    parts_all_origin_temp["part_reason"] = parts_object.part_reason
                    parts_all_origin.append(parts_all_origin_temp)

        chapters_all_request = []
        sections_all_request = []
        parts_all_request = []
        for bill_data_request_one in bill_data:
            chapters_one = {}
            chapters_temp={}
            sections_temp = {}
            parts_temp = {}
            chapters_one["chapter_id"] = bill_data_request_one["chapter_id"]
            chapters_one["chapter_number"] = bill_data_request_one["chapter_number"]
            chapters_one["chapter_title"] = bill_data_request_one["chapter_title"]
            chapters_one["chapter_content"] = bill_data_request_one["chapter_content"]
            chapters_temp = copy.copy(chapters_one)
            if not (chapters_temp in chapters_all_request):
                chapters_all_request.append(chapters_temp)
            chapters_one["section_id"] = bill_data_request_one["section_id"]
            chapters_one["section_number"] = bill_data_request_one["section_number"]
            chapters_one["section_title"] = bill_data_request_one["section_title"]
            chapters_one["section_content"] = bill_data_request_one["section_content"]
            sections_temp = copy.copy(chapters_one)
            if not (sections_temp in sections_all_request):
                sections_all_request.append(sections_temp)
            chapters_one["part_id"] = bill_data_request_one["part_id"]
            chapters_one["part_number"] = bill_data_request_one["part_number"]
            chapters_one["part_title"] = bill_data_request_one["part_title"]
            chapters_one["part_content"] = bill_data_request_one["part_content"]
            chapters_one["part_reason"] = bill_data_request_one["part_reason"]
            if ('added_flag' in bill_data_request_one):
                chapters_one['added_flag'] = bill_data_request_one['added_flag']
            else:
                chapters_one['added_flag'] = None
            parts_temp = copy.copy(chapters_one)
            if not (parts_temp in parts_all_request):
                parts_all_request.append(parts_temp)

        # update bill name
        bill_name_list = BusinessBillList.objects.update_or_create(business_id=business_id, defaults={'bill_name': bill_name})[0]
        previous_section = []
        previous_chapter = []
        previous_part = []
        previous_chapter = bill_name_list.chapters.all()
        for previous_one_chapter in previous_chapter:
            previous_section_temp = previous_one_chapter.sections.all()
            for previous_one_section_temp in previous_section_temp:
                previous_section.append(previous_one_section_temp)
                previous_part_temp = previous_one_section_temp.parts.all()
                for previous_one_part_temp in previous_part_temp:
                    previous_part.append(previous_one_part_temp)

        #         UPDATE CHAPTER
        added_chapter = []
        for chapters_one_request in chapters_all_request:
            checksumTemp = 0
            for previous_one_chapter in previous_chapter:
                if (int(chapters_one_request['chapter_number']) == int(previous_one_chapter.chapter_number)):
                    previous_one_chapter.chapter_title = chapters_one_request['chapter_title']
                    previous_one_chapter.chapter_content = chapters_one_request['chapter_content']
                    previous_one_chapter.save()
                    checksumTemp = 1
                    break
            if (checksumTemp == 0):
                added_chapter_item = BusinessBillChapter.objects.create(chapter_number=int(chapters_one_request['chapter_number']),chapter_title=chapters_one_request['chapter_title'],chapter_content="")
                added_chapter.append(added_chapter_item)
                bill_name_list.chapters.add(added_chapter_item)
                # added section
                for sections_one_request in sections_all_request:
                    if (int(added_chapter_item.chapter_number) == int(sections_one_request["chapter_number"])):
                        added_section = BusinessBillSection.objects.create(section_number=int(sections_one_request['section_number']),section_title=sections_one_request['section_title'], section_content="")
                        added_chapter_item.sections.add(added_section)
                        for parts_one_request in parts_all_request:
                            if (int(added_chapter_item.chapter_number) == int(parts_one_request["chapter_number"])):
                                if (int(added_section.section_number) == int(parts_one_request["section_number"])):
                                    added_part = BusinessBillPart.objects.create(part_number=int(parts_one_request['part_number']),part_title=parts_one_request['part_title'],part_content=parts_one_request['part_content'],part_reason=parts_one_request['part_reason'])
                                    added_section.parts.add(added_part)



        #                 update section
        added_section = []
        for sections_one_request in sections_all_request:
            checksumTemp = 0
            chapter_number_request = sections_one_request["chapter_number"]
            for previous_one_chapter in previous_chapter:
                if (int(previous_one_chapter.chapter_number) == int(chapter_number_request)):
                    previous_section = previous_one_chapter.sections.all()
                    for previous_one_section in previous_section:
                        if (int(sections_one_request["section_number"]) == int(previous_one_section.section_number)):
                            previous_one_section.section_title = sections_one_request["section_title"]
                            previous_one_section.section_content = ""
                            previous_one_section.save()
                            checksumTemp = 1
                            break
                    if (checksumTemp == 0):
                        added_section_item = BusinessBillSection.objects.create(
                            section_number=int(sections_one_request['section_number']),
                            section_title=sections_one_request['section_title'], section_content="")
                        added_section.append(added_section_item)
                        previous_one_chapter.sections.add(added_section_item)

                        for parts_one_request in parts_all_request:
                            if (int(previous_one_chapter.chapter_number) == int(parts_one_request["chapter_number"])):
                                if (int(added_section.section_number) == int(parts_one_request["section_number"])):
                                    added_part = BusinessBillPart.objects.create(part_number=int(parts_one_request['part_number']),part_title=parts_one_request['part_title'],part_content=parts_one_request['part_content'],part_reason=parts_one_request['part_reason'])
                                    added_section.parts.add(added_part)

                        break
                else:
                    continue


        # update Parts
        added_part = []
        for parts_one_request in parts_all_request:
            checksumTemp = 0
            chapter_number_request = parts_one_request["chapter_number"]
            section_number_request = parts_one_request["section_number"]
            if ('added_flag' in parts_one_request):
                if (parts_one_request['added_flag'] == '2'):
                    # inserted, added
                    for previous_one_chapter in previous_chapter:
                        if (int(previous_one_chapter.chapter_number) == int(chapter_number_request)):
                            previous_section = previous_one_chapter.sections.all()
                            for previous_one_section in previous_section:
                                if (int(previous_one_section.section_number) == int(section_number_request)):
                                    added_part_item = BusinessBillPart.objects.create(
                                        part_number=int(parts_one_request['part_number']),
                                        part_title=parts_one_request['part_title'],
                                        part_content=parts_one_request['part_content'],
                                        part_reason=parts_one_request['part_reason'])
                                    added_part.append(added_part_item)
                                    previous_one_section.parts.add(added_part_item)
                if (parts_one_request['added_flag'] == '1'):
                    # updated
                    for previous_one_chapter in previous_chapter:
                        if (int(previous_one_chapter.chapter_number) == int(
                                chapter_number_request)):
                            previous_section = previous_one_chapter.sections.all()
                            for previous_one_section in previous_section:
                                if (int(previous_one_section.section_number) == int(
                                        section_number_request)):
                                    previous_part = previous_one_section.parts.all()
                                    for previous_one_part in previous_part:
                                        if (int(previous_one_part.id) == int(parts_one_request['part_id'])):
                                            previous_one_part.part_number = int(parts_one_request['part_number'])
                                            previous_one_part.part_title = parts_one_request['part_title']
                                            previous_one_part.part_content = parts_one_request['part_content']
                                            previous_one_part.part_reason = parts_one_request['part_reason']
                                            previous_one_part.save()

        # deleted Parts

        resp = code.get_msg(code.SUCCESS)
    except Exception as e:
        logger.exception('api_business_send_guider_message Exception:{0}'.format(str(e)))
        resp = code.get_msg(code.SYSTEM_ERROR)

    return HttpResponse(json.dumps(resp, ensure_ascii=False), content_type="application/json")
##############################################
